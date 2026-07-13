"""
Cortex AI IDE — Main Window
Full 3-panel layout: Sidebar | Editor Tabs | AI Chat + Terminal
"""

import ctypes
import ctypes.wintypes
import json
import os
import sys
import time
import platform
import uuid as _uuid
from pathlib import Path
from typing import Any, Optional, cast
from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QHBoxLayout, QVBoxLayout, QSplitter,
    QTabWidget, QLabel, QPushButton, QFileDialog,
    QMenu, QMessageBox, QInputDialog, QTabBar,
    QFrame, QSizePolicy, QApplication, QDialog
)
from PyQt6.QtCore import Qt, pyqtSignal, pyqtSlot, QTimer, QSignalBlocker
from PyQt6.QtGui import QAction, QKeySequence, QIcon, QFont, QCloseEvent, QTextDocument
# Custom title bar removed — using native Windows title bar
from PyQt6.QtWebChannel import QWebChannel
from PyQt6.QtWebEngineCore import QWebEnginePage

from src.config.settings import get_settings
from src.config.theme_manager import get_theme_manager
from src.config.points_manager import get_points_manager, InsufficientPointsError
from src.ui.cursor_split_handle import CursorSplitter
from src.core.project_manager import ProjectManager
from src.core.file_manager import FileManager
from src.core.session_manager import SessionManager
from src.core.codebase_index import get_codebase_index
# Agent bridge - connects Cortex to agent module
try:
    from src.ai.agent_bridge import get_agent_bridge as AIAgent
    HAS_AGENT_BRIDGE = True
except ImportError:
    from src.ai.stub_agent import get_stub_agent as AIAgent  # Fallback stub
    HAS_AGENT_BRIDGE = False

import subprocess as _subprocess
import logging as _logging

_log = _logging.getLogger(__name__)





class CodeAnalyzer:
    """Simple prompt builder for AI code actions."""

    def build_explain_prompt(self, code: str, language: str) -> str:
        return f"Explain this {language} code in detail. Break down what each part does and why:\n\n```{language}\n{code}\n```"

    def build_refactor_prompt(self, code: str, language: str) -> str:
        return f"Refactor this {language} code to be cleaner, more efficient, and follow best practices. Explain your changes:\n\n```{language}\n{code}\n```"

    def build_test_prompt(self, code: str, language: str) -> str:
        return f"Write comprehensive unit tests for this {language} code. Include edge cases and error handling:\n\n```{language}\n{code}\n```"

    def build_debug_prompt(self, code: str, error: str, language: str) -> str:
        return f"Help me debug this {language} code. Error: {error}\n\n```{language}\n{code}\n```\n\nWhat's causing this error and how do I fix it?"

    def build_optimize_prompt(self, code: str, language: str) -> str:
        return f"Optimize this {language} code for performance, readability, and memory efficiency. Explain your optimizations:\n\n```{language}\n{code}\n```"
# from src.ai.file_edit_tracker import FileEditTracker
from src.core.git_manager import GitManager, GitStatus
# Sidebar is now HTML-based (sidebar.html + sidebar_bridge.py) — no Python widget class
# CommandPalette removed - not implemented in AI-first mode
# from src.ui.components.command_palette import CommandPalette
from src.ui.components.editor import CodeEditor  # kept for backward compat (welcome/PDF/image tabs)
from src.ui.components.webview_panel import WebviewPanel
from src.ui.components.xterm_terminal import XTermWidget
from src.ui.components.find_replace import FindReplaceDialog

from src.ui.dialogs.diff_viewer import DiffWindow
from src.utils.icons import make_icon
# Live Server removed in AI-first mode - AI handles code execution
# from src.core.live_server import LiveServer
from src.utils.helpers import detect_language, shorten_path
from src.utils.logger import get_logger
from src.utils.notifications import show_task_complete_notification, show_toast_notification, notify_input_needed, notify_permission_required

log = get_logger("main_window")

try:
    from src.ui.syntax_highlighting_config import (
        UniversalCodeColorizer,
        MarkdownColorizer, 
        DRACULA_COLORS,
        FONTS
    )
    HAS_SYNTAX_HIGHLIGHTING = True
except ImportError:
    HAS_SYNTAX_HIGHLIGHTING = False
    log.debug("Syntax highlighting module not available")



def _resource_path(relative_path: str) -> str:
    """Get absolute path to a bundled resource, works for dev and PyInstaller .exe.

    When frozen (compiled .exe), sys._MEIPASS points to the temp extraction dir.
    When running as script, __file__ gives the dev source location.
    Falls back to os.getcwd() if both fail.
    """
    if getattr(sys, 'frozen', False):
        # PyInstaller bundles everything under sys._MEIPASS
        base = sys._MEIPASS
    else:
        # Dev mode: resolve relative to this file (src/main_window.py -> project root)
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    candidate = os.path.join(base, relative_path)
    if os.path.exists(candidate):
        return candidate

    # Fallback 1: try next to the .exe (onedir mode)
    if getattr(sys, 'frozen', False):
        exe_dir = os.path.dirname(sys.executable)
        candidate = os.path.join(exe_dir, relative_path)
        if os.path.exists(candidate):
            return candidate

    # Fallback 2: try _internal subdir (PyInstaller onedir mode)
    if getattr(sys, 'frozen', False):
        internal_dir = os.path.join(os.path.dirname(sys.executable), '_internal')
        candidate = os.path.join(internal_dir, relative_path)
        if os.path.exists(candidate):
            return candidate

    # Fallback 3: cwd
    candidate = os.path.join(os.getcwd(), relative_path)
    return candidate  # may not exist, caller handles that


class CleanTabBar(QTabBar):
    """Tab bar that draws a clean × close button instead of Qt's default box."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMouseTracking(True)
        self._hovered_tab = -1
        # Inherit theme from theme manager
        from src.config.theme_manager import get_theme_manager
        self._is_dark = get_theme_manager().is_dark

    def set_dark(self, is_dark: bool):
        self._is_dark = is_dark
        self.update()

    def mouseMoveEvent(self, event):
        idx = self.tabAt(event.pos())
        if idx != self._hovered_tab:
            self._hovered_tab = idx
            self.update()
        super().mouseMoveEvent(event)

    def leaveEvent(self, event):
        self._hovered_tab = -1
        self.update()
        super().leaveEvent(event)

    def tabSizeHint(self, index):
        s = super().tabSizeHint(index)
        s.setHeight(34)
        return s

    def paintEvent(self, event):
        """Draw tabs with a clean × button — theme-aware."""
        from PyQt6.QtGui import QPainter, QColor
        from PyQt6.QtCore import QRect, Qt
        painter = QPainter(self)
        try:
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)

            d = self._is_dark
            if d:
                # Dark Theme Tab Colors
                col_sel_bg    = QColor("#181818")
                col_hover_bg  = QColor("#1f1f1f")
                col_normal_bg = QColor("#141414")
                col_accent    = QColor("#228df2")
                col_divider   = QColor("#2a2a2a")
                col_sel_fg    = QColor("#ffffff")
                col_hover_fg  = QColor("#d6d6dd")
                col_normal_fg = QColor("#6d6d6d")
                col_close     = QColor("#f14c4c")
                col_close_hover = QColor("#ff6b6b")
            else:
                # Light Theme Tab Colors (VS Code/Excel palette — green accent)
                col_sel_bg    = QColor("#FFFFFF")
                col_hover_bg  = QColor("#D0D0D0")
                col_normal_bg = QColor("#E0E0E0")
                col_accent    = QColor("#4CAF50")
                col_divider   = QColor("#D0D0D0")
                col_sel_fg    = QColor("#000000")
                col_hover_fg  = QColor("#000000")
                col_normal_fg = QColor("#6D6D6D")
                col_close     = QColor("#F44336")
                col_close_hover = QColor("#E53935")

            for i in range(self.count()):
                rect = self.tabRect(i)
                is_selected = (i == self.currentIndex())
                is_hovered  = (i == self._hovered_tab)

            # Background
            if is_selected:
                painter.fillRect(rect, col_sel_bg)
            elif is_hovered:
                painter.fillRect(rect, col_hover_bg)
            else:
                painter.fillRect(rect, col_normal_bg)

            # Accent top border on active tab
            if is_selected:
                painter.fillRect(QRect(rect.x(), rect.y(), rect.width(), 2), col_accent)

            # Right divider
            painter.fillRect(QRect(rect.right(), rect.y() + 4, 1, rect.height() - 8), col_divider)

            # Tab label
            text = self.tabText(i)
            
            # ── Premium Tab Icon ───────────────────────────────────────────────
            icon_x = rect.x() + 10
            icon_size = 14
            
            # Determine icon based on file extension
            filepath = None
            if hasattr(self.parent(), '_files'):
                filepath = self.parent()._files.get(i)
            
            from src.ui.components.sidebar import _get_icon_name
            icon_name = _get_icon_name(filepath) if filepath else "files"
            if text == "Welcome": icon_name = "ai"
            if "Terminal" in text: icon_name = "terminal"
            
            # Cursor IDE Syntax Colors for File Icons
            colors = {
                "python":   "#83d6c5",  # teal - keyword
                "html":     "#87c3ff",  # light blue - class/tag
                "css":      "#87c3ff",  # light blue - class
                "javascript": "#e394dc", # pink - string
                "typescript": "#87c3ff", # light blue - class
                "markdown": "#d6d6dd", # primary text
                "json":     "#efb080",  # orange - number
                "java":     "#83d6c5",  # teal
                "rust":     "#efb080",  # orange
                "go":       "#87c3ff",  # light blue
                "sql":      "#83d6c5",  # teal
                "ai":       "#228df2",  # accent blue
                "terminal": "#6d6d6d"   # muted
            }
            icon_color = colors.get(icon_name, "#abb2bf")
            
            # Using a fallback if terminal icon not found, or use the shell icon
            icon_pixmap = make_icon(icon_name, icon_color, icon_size).pixmap(icon_size, icon_size)
            painter.drawPixmap(icon_x, rect.y() + (rect.height() - icon_size)//2, icon_pixmap)
            
            # Reserve space for close button (14px + 2px right padding)
            btn_reserved = 16  # 14 + 2px right margin
            
            # Label - leave space for close button
            label_x = icon_x + icon_size + 6
            label_width = rect.width() - (label_x - rect.x()) - btn_reserved - 2
            label_rect = QRect(label_x, rect.y(), max(0, label_width), rect.height())
            
            # Draw label with eliding to prevent overflow
            painter.save()
            fg = col_sel_fg if is_selected else (col_hover_fg if is_hovered else col_normal_fg)
            painter.setPen(fg)
            f = painter.font()
            f.setPointSize(9)
            painter.setFont(f)
            # Use elided text to prevent overflow into button area
            from PyQt6.QtGui import QFontMetrics
            fm = QFontMetrics(f)
            elided_text = fm.elidedText(text, Qt.TextElideMode.ElideRight, label_width)
            painter.drawText(label_rect,
                             Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft,
                             elided_text)
            painter.restore()

            # × close button — ONLY show on hovered or selected tabs
            if is_hovered or is_selected:
                # Fixed position from right edge - 2px padding
                btn_size = 14
                btn_x = rect.right() - btn_size - 2
                btn_y = rect.y() + (rect.height() - btn_size) // 2
                btn_rect = QRect(btn_x, btn_y, btn_size, btn_size)
                
                # Check if cursor is over the × button
                cursor_pos = self.mapFromGlobal(self.cursor().pos())
                is_close_hovered = btn_rect.contains(cursor_pos)

                # Determine colors
                if is_close_hovered:
                    bg_color = QColor("#f14c4c") if d else QColor("#d73a49")  # Red
                    x_color = QColor("#ffffff")  # White X
                else:
                    bg_color = QColor(255, 255, 255, 40) if d else QColor(0, 0, 0, 30)
                    x_color = QColor("#f14c4c") if d else QColor("#d73a49")  # Red X
                
                # CRITICAL: Fill button area with tab background color first
                bg_fill = col_sel_bg if is_selected else col_hover_bg
                painter.fillRect(btn_rect, bg_fill)
                
                # Draw button background
                painter.setBrush(bg_color)
                painter.setPen(Qt.PenStyle.NoPen)
                painter.drawRoundedRect(btn_rect, 2, 2)
                
                # Draw X centered
                painter.setPen(x_color)
                font = painter.font()
                font.setPointSize(10)
                font.setBold(True)
                painter.setFont(font)
                painter.drawText(btn_rect, Qt.AlignmentFlag.AlignCenter, "×")
        finally:
            painter.end()

    def mousePressEvent(self, event):
        """Handle × button click to close the tab."""
        from PyQt6.QtCore import Qt, QRect
        from PyQt6.QtGui import QMouseEvent
        i = self.tabAt(event.pos())
        if i >= 0:
            # Only check if tab is hovered or selected (matching paint logic)
            is_selected = (i == self.currentIndex())
            is_hovered = (i == self._hovered_tab)
            if is_hovered or is_selected:
                rect = self.tabRect(i)
                btn_size = 14
                btn_x = rect.right() - btn_size - 2  # Match paintEvent
                btn_y = rect.y() + (rect.height() - btn_size) // 2
                btn_rect = QRect(btn_x, btn_y, btn_size, btn_size)
                if btn_rect.contains(event.pos()):
                    self.tabCloseRequested.emit(i)
                    return
        super().mousePressEvent(event)



class EditorTabWidget(QTabWidget):
    """Central editor area with tabs."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTabBar(CleanTabBar(self))
        self.setTabsClosable(True)
        self.setMovable(True)
        self.setDocumentMode(True)
        self.tabCloseRequested.connect(self._close_tab)
        self._files: dict[int, str] = {}   # tab_index -> filepath
        self._modified: set[str] = set()

    def open_diff_tab(self, file_path: str, original: str, modified: str, is_dark: bool = True):
        """Open a read-only diff tab next to the file tab. Tab name: '⟷ filename'."""
        import difflib
        from PyQt6.QtWidgets import QTextBrowser
        from PyQt6.QtGui import QFont
        from pathlib import Path as _P

        file_name = _P(file_path).name
        tab_label = f'⟷ {file_name}'

        # If a diff tab for this file already exists, switch to it
        for idx in range(self.count()):
            if self.tabText(idx) == tab_label:
                self.setCurrentIndex(idx)
                return

        # Build unified diff HTML
        diff_lines = list(difflib.unified_diff(
            original.splitlines(keepends=True),
            modified.splitlines(keepends=True),
            fromfile='Original',
            tofile='Modified',
            n=3
        ))

        bg   = '#1e1e1e'
        fg   = '#cccccc'
        add_bg = 'rgba(46,160,67,0.2)'
        add_fg = '#56d364'
        rem_bg = 'rgba(248,81,73,0.2)'
        rem_fg = '#f85149'
        info_fg = '#8b949e'

        def esc(t: str) -> str:
            return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        parts: list[str] = [
            f"<div style='background:{bg};color:{fg};white-space:pre;font-family:\"Cascadia Code\",Consolas,monospace;font-size:13px;line-height:1.5;padding:10px;'>"
        ]

        if not diff_lines:
            parts.append(f"<div style='color:{info_fg};padding:20px;'>No changes detected.</div>")
        else:
            for line in diff_lines:
                line = line.rstrip('\n')
                s = esc(line)
                if line.startswith('+++') or line.startswith('---'):
                    parts.append(f"<div style='color:{info_fg};font-weight:bold;background:rgba(128,128,128,0.1);padding-left:6px;'>{s}</div>")
                elif line.startswith('@@'):
                    parts.append(f"<div style='color:{info_fg};padding-left:6px;margin-top:6px;'>{s}</div>")
                elif line.startswith('+'):
                    parts.append(f"<div style='color:{add_fg};background:{add_bg};padding-left:6px;'>{s}</div>")
                elif line.startswith('-'):
                    parts.append(f"<div style='color:{rem_fg};background:{rem_bg};padding-left:6px;'>{s}</div>")
                else:
                    parts.append(f"<div style='color:{fg};padding-left:6px;'>{s}</div>")
        parts.append('</div>')

        browser = QTextBrowser()
        browser.setReadOnly(True)
        browser.setOpenExternalLinks(False)
        font = QFont('Cascadia Code', 10)
        font.setStyleHint(QFont.StyleHint.Monospace)
        browser.setFont(font)
        browser.setStyleSheet(f'background:{bg};border:none;')
        browser.setHtml(''.join(parts))

        idx = self.addTab(browser, tab_label)
        self.setCurrentIndex(idx)

    def open_file(self, filepath: str, content: str, language: str, is_dark: bool = True) -> int:
        """Open a file in a new tab (or switch to existing)."""
        # Check if already open
        for idx, fp in self._files.items():
            if fp == filepath:
                # File already open - check if content matches, if not update it
                editor = self.widget(idx)
                if isinstance(editor, CodeEditor):
                    current_content = editor.toPlainText()
                    if current_content != content:
                        # Content changed, reload it
                        with QSignalBlocker(editor.document()):
                            editor.set_content(content, language, filepath)
                        _log.info(f"Updated content for already-open file: {filepath}")
                    else:
                        _log.info(f"File already open with same content: {filepath}")
                self.setCurrentIndex(idx)
                return idx
        
        # Create editor
        editor = CodeEditor(language=language)
        
        # Apply current theme to the new editor
        editor.set_theme(is_dark)
        
        # Disconnect the internal document→editor connection temporarily
        # Use blockSignals instead of disconnect to avoid errors
        with QSignalBlocker(editor.document()):
            editor.set_content(content, language, filepath)
        
        # NOW connect OUR handler - anything after this is a user edit
        editor.content_modified.connect(lambda: self._mark_modified(filepath))

        name = Path(filepath).name
        idx = self.addTab(editor, name)
        self._files[idx] = filepath
        self.setCurrentIndex(idx)
        self.setTabToolTip(idx, filepath)
        
        # Set file type icon on tab
        self._set_tab_icon(idx, filepath)
        
        return idx
    
    def _set_tab_icon(self, idx: int, filepath: str):
        """Set the tab icon based on file extension."""
        from PyQt6.QtGui import QIcon, QPixmap
        from pathlib import Path
        
        ext = Path(filepath).suffix.lower()
        icon_map = {
            '.py': 'python.png',
            '.js': 'javascript.png',
            '.ts': 'typescript.png',
            '.jsx': 'javascript.png',
            '.tsx': 'typescript.png',
            '.html': 'html.png',
            '.htm': 'html.png',
            '.css': 'css.png',
            '.json': 'json.png',
            '.md': 'markdown.png',
            '.java': 'java.png',
            '.rs': 'rust.png',
            '.csv': 'csv.png',
            '.env': 'env.png',
        }
        
        icon_file = icon_map.get(ext)
        if icon_file:
            icon_path = _resource_path(os.path.join("src", "assets", "icons", icon_file))
            if os.path.exists(icon_path):
                pixmap = QPixmap(str(icon_path))
                if not pixmap.isNull():
                    # Scale to appropriate size for tab (16x16)
                    scaled = pixmap.scaled(16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
                    self.setTabIcon(idx, QIcon(scaled))

    def _mark_modified(self, filepath: str):
        """Mark file as modified and update tab with white dot."""
        self._modified.add(filepath)
        for idx, fp in self._files.items():
            if fp == filepath:
                name = Path(fp).name
                # Set tab text with white dot (●) prefix
                self.setTabText(idx, f"● {name}")
                # Set tab tooltip to show modified status
                self.setTabToolTip(idx, f"{filepath} (Modified)")
                break
    
    def _mark_saved(self, filepath: str):
        """Mark file as saved and remove dot from tab."""
        self._modified.discard(filepath)
        for idx, fp in self._files.items():
            if fp == filepath:
                name = Path(fp).name
                # Remove dot and restore normal tab text
                self.setTabText(idx, name)
                # Restore normal tooltip
                self.setTabToolTip(idx, filepath)
                break

    def _close_tab(self, index: int):
        """Close tab with save confirmation if modified."""
        widget = self.widget(index)
        filepath = self._files.get(index)
        
        # Check if file has unsaved changes
        if filepath and filepath in self._modified:
            from PyQt6.QtWidgets import QMessageBox
            
            reply = QMessageBox.question(
                self,
                "Save Changes?",
                f"Do you want to save the changes to '{Path(filepath).name}'?",
                QMessageBox.StandardButton.Save | 
                QMessageBox.StandardButton.Discard | 
                QMessageBox.StandardButton.Cancel
            )
            
            if reply == QMessageBox.StandardButton.Cancel:
                return  # User cancelled - don't close tab
            elif reply == QMessageBox.StandardButton.Save:
                # Save the file before closing
                editor = self.current_editor() if index == self.currentIndex() else None
                if editor:
                    content = editor.get_all_text()
                    try:
                        # Normalize line endings to prevent doubled empty lines
                        content = content.replace("\r\n", "\n").replace("\r", "\n")
                        Path(filepath).write_text(content, encoding='utf-8', newline='')
                        self._mark_saved(filepath)
                    except Exception as e:
                        from PyQt6.QtWidgets import QMessageBox
                        QMessageBox.critical(self, "Save Error", f"Failed to save file: {e}")
                        return  # Don't close tab on save error
        
        # Proceed with closing tab
        self.removeTab(index)
        
        # Cleanup files mapping
        new_files = {}
        for old_idx, fp in self._files.items():
            if old_idx == index:
                continue
            new_idx = old_idx if old_idx < index else old_idx - 1
            new_files[new_idx] = fp
        self._files = new_files
        
        if filepath:
            self._modified.discard(filepath)
            
        # If no tabs left, we could show welcome again or just keep it empty
        if self.count() == 0:
            # Maybe tell parent to show welcome? 
            # For now just let it be empty to match VS Code behavior
            pass

        # Notify sidebar search of updated open files
        self._notify_search_of_open_files()

    def current_editor(self) -> CodeEditor | None:
        w = self.currentWidget()
        return w if isinstance(w, CodeEditor) else None

    def current_filepath(self) -> str | None:
        return self._files.get(self.currentIndex())

    def save_current(self, file_manager: FileManager) -> bool:
        """Save current file and remove modified indicator."""
        editor = self.current_editor()
        fp = self.current_filepath()
        if not editor or not fp:
            return False
        content = editor.get_all_text()
        ok = file_manager.write(fp, content)
        if ok:
            # Use _mark_saved to properly update tab text and tooltip
            self._mark_saved(fp)
        return ok
    
    def save_file(self, filepath: str, content: str) -> bool:
        """Save a specific file and update its modified state."""
        from pathlib import Path as _Path
        try:
            # ── SAVE-SIZE GUARD: a code-editor buffer is never legitimately
            # this large. A crashed/looping webview once flushed gigabytes of
            # duplicated junk into a source file — refuse instead of writing.
            if len(content) > 50 * 1024 * 1024:  # 50MB
                _log.error(f"[SAVE-GUARD] BLOCKED save of {len(content):,} chars "
                           f"to {filepath} — editor buffer exceeds 50MB sanity cap")
                return False
            # Normalize line endings to prevent doubled empty lines
            content = content.replace("\r\n", "\n").replace("\r", "\n")
            _Path(filepath).write_text(content, encoding='utf-8', newline='')
            self._mark_saved(filepath)
            return True
        except Exception as e:
            _log.warning(f"Save error: {e}")
            return False

    def get_open_files(self) -> list[str]:
        return list(self._files.values())

    def close_current_tab(self):
        """Close the currently active tab."""
        current_idx = self.currentIndex()
        if current_idx >= 0:
            self._close_tab(current_idx)

    def close_all_tabs(self):
        """Close all open tabs."""
        while self.count() > 0:
            self._close_tab(0)

    def update_theme(self, is_dark: bool):
        # Update tab bar colours
        if isinstance(self.tabBar(), CleanTabBar):
            self.tabBar().set_dark(is_dark)

        # Update individual editor widgets
        for i in range(self.count()):
            w = self.widget(i)
            if isinstance(w, CodeEditor):
                w.set_theme(is_dark)


# ═══════════════════════════════════════════════════════════════
# NOTIFICATION SUMMARY EXTRACTOR (module-level helper)
# ═══════════════════════════════════════════════════════════════

def _extract_notification_summary(text: str) -> str:
    """Extract a clean, minimal task summary from AI response text.

    Skips conversational fluff (greetings, "I think", "Here is", etc.)
    and returns the first meaningful sentence describing what was done.
    """
    import re
    # Truncate to reasonable length before processing
    text = text[:500]
    # Strip markdown formatting
    cleaned = re.sub(r'#{1,6}\s+', '', text)
    cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned)     # bold
    cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)           # italic
    cleaned = re.sub(r'`{1,3}[^`]*`{1,3}', '', cleaned)         # inline + block code
    cleaned = re.sub(r'\|[-:\s|]+\|', '', cleaned)             # table separators
    cleaned = re.sub(r'\|[^|]*\|', '', cleaned)                 # table rows
    cleaned = re.sub(r'[-*_]{3,}', '', cleaned)                  # horizontal rules
    cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)  # links
    cleaned = re.sub(r'[>\-]\s+', '', cleaned)                  # blockquote markers
    cleaned = re.sub(r'\n{2,}', '. ', cleaned)                  # double newlines → period
    cleaned = cleaned.replace('\n', ' ').strip()
    # Collapse multiple spaces
    cleaned = re.sub(r'\s{2,}', ' ', cleaned)
    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', cleaned)
    # Fluff patterns to skip (conversational intros)
    skip_starts = {
        'i think', 'i can', 'i will', 'i would', 'i have', 'i see',
        'here is', 'here are', 'let me', 'let\'s', 'sure', 'of course',
        'great question', 'thank', 'hello', 'hi there', 'hi,', 'hey',
        'certainly', 'absolutely', 'no problem', 'you\'re', 'that\'s a',
        'based on', 'looking at', 'first', 'now', 'next', 'then',
        'to start', 'to begin',
    }
    for s in sentences:
        s_clean = s.strip().strip('.,;:!?()[]"\' ')
        if len(s_clean) < 12:
            continue
        s_lower = s_clean.lower()
        if any(s_lower.startswith(p) for p in skip_starts):
            continue
        # Found a meaningful sentence — cap at 100 chars
        if len(s_clean) > 100:
            # Try to break at a natural boundary
            cut = s_clean[:100].rfind(' ')
            s_clean = s_clean[:cut] if cut > 60 else s_clean[:100]
        return s_clean
    # Fallback: first sentence >= 20 chars
    for s in sentences:
        s_clean = s.strip().strip('.,;:!?()[]"\' ')
        if len(s_clean) >= 20:
            return s_clean[:100]
    return ""




from src.ui.components.sidebar import SidebarWidget

class CortexMainWindow(QMainWindow):

    # Update-check result crossing from the background thread to the GUI
    # thread. Bug history (v2.7.0): this hop used a string-name invokeMethod
    # to a plain Python method that was never registered as a Qt slot, so it
    # ALWAYS failed ("call failed") and the update dialog — including FORCE
    # updates — never appeared. A signal emit is the idiomatic, thread-safe
    # mechanism: Qt auto-queues it onto the GUI thread.
    _update_check_ready = pyqtSignal(object)

    @staticmethod
    def _kill_orphaned_chromium():
        """Kill ALL orphaned QtWebEngineProcess.exe from previous instances.

        When Cortex crashes or is force-killed, Chromium subprocesses
        survive and consume 200-500MB RAM each. On next launch, these
        orphans conflict with the new Chromium processes and cause
        access violations / heap corruption crashes.

        Uses taskkill /F /IM (reliable, always available on Windows)
        instead of WMIC (deprecated on Windows 11+).
        """
        import subprocess, os, time
        import logging as _log
        _logger = _log.getLogger("main_window")
        killed = 0

        # --- Phase 1: Kill orphaned Chromium subprocesses via taskkill ---
        try:
            # taskkill /F /IM kills ALL instances of the named process.
            # We kill ALL of them (not just orphans) because at startup
            # there should be ZERO Chromium processes — any left over are
            # from a previous instance.
            result = subprocess.run(
                ["taskkill", "/F", "/IM", "QtWebEngineProcess.exe"],
                capture_output=True, text=True, timeout=8,
                creationflags=0x08000000  # CREATE_NO_WINDOW
            )
            # taskkill returns 0 on success, 128 if no process found
            if result.returncode == 0:
                # Count killed processes from output ("Successfully terminated process ...")
                killed = result.stdout.count("Successfully terminated process")
                if killed > 0:
                    _logger.info(f"[STARTUP] Killed {killed} Chromium process(es) via taskkill")
            elif "not found" not in result.stderr.lower():
                _logger.debug(f"[STARTUP] taskkill result: rc={result.returncode} stderr={result.stderr.strip()}")
        except FileNotFoundError:
            _logger.warning("[STARTUP] taskkill not found — skipping Chromium cleanup")
        except Exception as e:
            _logger.debug(f"[STARTUP] taskkill failed: {e}")

        # --- Phase 2: Also kill via ctypes for any that survived taskkill ---
        try:
            import ctypes
            my_pid = os.getpid()
            # Use CreateToolhelp32Snapshot to enumerate and kill survivors
            TH32CS_SNAPPROCESS = 0x00000002
            snapshot = ctypes.windll.kernel32.CreateToolhelp32Snapshot(TH32CS_SNAPPROCESS, 0)
            if snapshot:
                class PROCESSENTRY32(ctypes.Structure):
                    _fields_ = [
                        ("dwSize", ctypes.c_ulong), ("cntUsage", ctypes.c_ulong),
                        ("th32ProcessID", ctypes.c_ulong), ("th32DefaultHeapID", ctypes.c_size_t),
                        ("th32ModuleID", ctypes.c_ulong), ("cntThreads", ctypes.c_ulong),
                        ("th32ParentProcessID", ctypes.c_ulong), ("pcPriClassBase", ctypes.c_long),
                        ("dwFlags", ctypes.c_ulong), ("szExeFile", ctypes.c_wchar * 260),
                    ]
                pe = PROCESSENTRY32()
                pe.dwSize = ctypes.sizeof(PROCESSENTRY32)
                if ctypes.windll.kernel32.Process32FirstW(snapshot, ctypes.byref(pe)):
                    while True:
                        if pe.szExeFile.lower() == "qtwebengineprocess.exe" and pe.th32ProcessID != my_pid:
                            try:
                                h = ctypes.windll.kernel32.OpenProcess(0x0001, False, pe.th32ProcessID)
                                if h:
                                    ctypes.windll.kernel32.TerminateProcess(h, 0)
                                    ctypes.windll.kernel32.CloseHandle(h)
                                    killed += 1
                            except Exception:
                                pass
                        if not ctypes.windll.kernel32.Process32NextW(snapshot, ctypes.byref(pe)):
                            break
                ctypes.windll.kernel32.CloseHandle(snapshot)
        except Exception:
            pass

        # --- Phase 3: Wait briefly for OS to release sockets/shared memory ---
        if killed > 0:
            time.sleep(0.5)  # Let the OS reclaim resources before we start Chromium
            _logger.info(f"[STARTUP] Cleaned up {killed} total Chromium process(es)")


    def __init__(self):
        super().__init__()
        import PyQt6.QtCore as _QtCore
        self._start_time = time.time()  # track startup time for crash-prevention guards
        self._warmup_duration = 5  # seconds — Chromium GPU/compositor stabilization (crash zone ≤34s; 5s sufficient after memory flag tuning)
        self._warmup_queued_files: list = []  # (filepath, priority) tuples queued during warmup
        self._warmup_flush_timer = QTimer(self)
        self._warmup_flush_timer.setSingleShot(True)
        self._warmup_flush_timer.timeout.connect(self._flush_warmup_queue)
        self._warmup_flush_timer.start(self._warmup_duration * 1000)
        log.info("MainWindow: __init__ START")

        # NOTE: Orphaned Chromium cleanup is handled ONCE in main.py
        # (before QApplication creation). Calling it here again caused
        # a "double boot" delay — the IDE would pause twice during
        # startup (3s in main.py + 0.5s here) making it appear to
        # "try to open twice". Removed to fix the double-boot issue.
        from src.utils.startup_profiler import checkpoint as _profile
        _profile("main_window_init_start")
        log.info("MainWindow: Initializing managers...")
        self._settings = get_settings()
        self._theme_manager = get_theme_manager()

        # ── CRITICAL ORDER: set theme state BEFORE any widget is built ──
        # Chat panel, tool cards, CleanTabBar, toolbar icons etc. all read
        # their palette at CONSTRUCTION time (tokens.TOKENS / is_dark).
        # Previously the theme was only applied in _apply_initial_theme(),
        # AFTER _build_ui() — so on a light-theme startup every widget was
        # built with dark tokens and only partially re-themed afterwards
        # (the "dark leakage in light mode" bug). set_active_no_qss is
        # state-only (no setStyleSheet), so this is free; the full QSS
        # still applies once in _apply_initial_theme().
        _saved_theme = getattr(self._settings, 'theme', 'dark') or 'dark'
        self._theme_manager.set_active_no_qss(_saved_theme)
        from src.ui import tokens as _tokens
        _tokens.set_theme("dark" if self._theme_manager.is_dark else "light")
        self._project_manager = ProjectManager()
        self._file_manager = FileManager()
        self._session_manager = SessionManager()
        self._live_server: Optional[Any] = None  # built-in HTML Live Server (disabled in AI-first mode)
        # Git manager for source control integration
        self._git_manager = GitManager()
        log.info("[GIT] GitManager initialized")
        _profile("git_manager")
        
        # Agent bridge connects Cortex to agent module
        self._ai_agent = AIAgent(file_manager=self._file_manager)
        if HAS_AGENT_BRIDGE:
            log.info("[AGENT] Agent bridge initialized - full integration active")
        else:
            log.info("[AGENT] Stub agent initialized - bridge not available")
        _profile("agent_bridge")
        
        # Legacy AI components removed - agent_bridge.py is the active runtime
        self._file_tracker = None
        self._diff_window = DiffWindow(self)
        self._changed_files_panel = None  # ChangedFilesPanel not available in dev source
        self._codebase_index = None
        self._inline_edit_context = None
        # Live Server removed in AI-first mode
        self._live_server = None
        
        # Initialize UI components to None to prevent theme application crashes if build fails
        self._toolbar = None
        self._toolbar_sep = None
        self._toolbar_logo = None
        self._toolbar_btns = []
        self._memory_btn = None
        self._settings_btn = None

        try:
            log.info("MainWindow: Building UI...")
            _profile("pre_build_ui")
            self._build_ui()
            _profile("post_build_ui")
            log.info("MainWindow: Building Menu...")
            # Build menu bar for all modes (Codex-style has menu bar)
            self._build_menu()
            log.info("MainWindow: Building Status Bar...")
            self._build_status_bar()
        except Exception as e:
            log.error(f"UI Build Error: {e}", exc_info=True)
            raise  # Re-raise to stop execution

        log.info("MainWindow: Connecting signals...")
        self._connect_signals()
        log.info("MainWindow: Applying initial theme...")
        self._apply_initial_theme()

        # ══════ STARTUP PERFORMANCE: Show window FIRST, restore session AFTER ══════
        # Previously _restore_session() blocked for 9s (DB init + 20 messages + ChatPersist)
        # before the window appeared. Now we show the window immediately and restore
        # in a deferred QTimer so the user sees the UI in ~2s instead of ~12s.
        _profile("signals_and_theme")
        log.info("MainWindow: Initialization complete (session restore deferred).")
        _profile("init_complete_before_show")
        # FIX: Session restore MUST run AFTER _deferred_show (which also uses
        # QTimer.singleShot(0)). Previously both fired on the same tick — session
        # restore ran first (2.4s blocking), preventing the window from showing
        # and starving Chromium of event-loop time. Now we defer restore to 500ms
        # AFTER __init__ returns, guaranteeing _deferred_show runs first.
        QTimer.singleShot(500, self._deferred_session_restore)

        # Enable drag and drop of folders/files onto the main window
        self.setAcceptDrops(True)

        # Native Windows title bar — no frameless, no custom title bar
        self.setWindowTitle("Cortex AI IDE")

        # Set Window Icon (Title Bar + Taskbar) - BEFORE show() to prevent flash
        # Uses pre-generated taskbar_rounded.png (run generate_icons.py once)
        # Uses _resource_path() which handles both dev and PyInstaller frozen modes
        logo_dir = _resource_path(os.path.join("src", "assets", "logo"))

        icon_candidates = [
            os.path.join(logo_dir, "taskbar_rounded.png"),
            os.path.join(logo_dir, "taskbar.png"),
            os.path.join(logo_dir, "taskbar.ico"),
        ]

        icon = QIcon()
        found_icon = False
        for candidate in icon_candidates:
            if os.path.exists(candidate):
                from PyQt6.QtGui import QPixmap
                from PyQt6.QtCore import Qt
                pm = QPixmap(candidate)
                if not pm.isNull():
                    for sz in [16, 32, 48, 64, 128, 256]:
                        icon.addPixmap(pm.scaled(sz, sz, _QtCore.Qt.AspectRatioMode.KeepAspectRatio, _QtCore.Qt.TransformationMode.SmoothTransformation))
                    found_icon = True
                    log.info(f"[ICON] Successfully loaded icon: {candidate}")
                    break

        if not found_icon:
            log.error(f"[ICON] No valid icon found in candidates: {icon_candidates}")
            log.error(f"[ICON] Checked paths: {[os.path.abspath(candidate) for candidate in icon_candidates]}")

        if not icon.isNull():
            self.setWindowIcon(icon)
            # Also set app-level icon for taskbar grouping
            from PyQt6.QtWidgets import QApplication
            QApplication.instance().setWindowIcon(icon)
        else:
            log.error("[ICON] Failed to set window icon: QIcon is null")

        # Window geometry
        w = self._settings.get("window", "width") or 1400
        h = self._settings.get("window", "height") or 900
        self.resize(w, h)
        self.setGeometry(100, 100, w, h)
        self._want_maximized = bool(self._settings.get("window", "maximized"))

        # DEFERRED SHOW — avoids ACCESS VIOLATION during __init__ when
        # child widget tree (spinners, webengine, graphics effects) isn't
        # fully stable yet. QTimer.singleShot(0) fires on the first event
        # loop tick AFTER __init__ returns, giving Qt time to finish
        # internal widget construction.
        log.info("[STARTUP] Deferring window show to event loop...")
        QTimer.singleShot(0, self._deferred_show)

    def _deferred_show(self):
        """Show the window in ONE step — no normal→maximized resize jump.

        Previously this called showNormal() first, then showMaximized() 500ms later.
        That caused a 3-phase visual glitch: partial → dark areas → full.
        Now we go directly to the final state (maximized or normal) in one call.

        The loading overlay (created in _build_ui) covers the entire window
        so the user sees a clean dark splash → full UI transition.

        This avoids ACCESS VIOLATION crashes that happen when show() is called
        during __init__ while child widgets (spinners, webengine, overlays) are
        still being constructed by Qt's internal C++ layer.

        CRITICAL z-order fix: The main window is shown FIRST, THEN the overlay
        is re-raised on top. If we raise the overlay BEFORE showNormal(), the
        DWM z-order recomposition during showNormal() can briefly flash the
        main window content for exactly 1 frame — the "blink" users see.
        Showing the window first (hidden behind overlay) and then asserting
        the overlay's topmost position AFTER composition is complete
        eliminates the race condition entirely.
        """
        try:
            # PHASE 1: Show the main window — it's hidden behind the
            # startup overlay (already visible from _build_ui), so the
            # user sees nothing yet.
            if getattr(self, '_want_maximized', False):
                log.info("[STARTUP] _deferred_show: calling showMaximized() directly...")
                self.showMaximized()
            else:
                log.info("[STARTUP] _deferred_show: calling showNormal()...")
                self.showNormal()

            # PHASE 2: re-assert the overlay above all sibling panels now
            # that the window is composed (child overlay — covers only the
            # Cortex window, never the rest of the desktop).
            if hasattr(self, '_startup_overlay') and self._startup_overlay:
                self._startup_overlay.setGeometry(self.rect())
                self._startup_overlay.raise_()
                self._startup_overlay.show()
                # Second assertion after 1 frame (16ms) as a safety net —
                # some DWM compositions are delayed by 1 frame on Win11.
                QTimer.singleShot(16, self._assert_overlay_topmost)

            log.info("[STARTUP] _deferred_show: window shown")
        except Exception as e:
            log.error(f"[STARTUP] _deferred_show failed: {e}", exc_info=True)
            return
        # Apply dark title bar after the first paint stabilizes
        QTimer.singleShot(200, self._apply_dark_title_bar)
        # Re-sync splitter handles after window is fully rendered
        # (handles may not have correct dimensions until after first show)
        QTimer.singleShot(100, self._force_sync_all_splitter_handles)
        # Yield to event loop so the window paint completes before session restore
        QApplication.processEvents()

    def _assert_overlay_topmost(self):
        """Safety-net: re-raise overlay after 1 frame to catch delayed DWM compositions."""
        if hasattr(self, '_startup_overlay') and self._startup_overlay and self._startup_overlay.isVisible():
            self._startup_overlay.setGeometry(self.rect())
            self._startup_overlay.raise_()

    def _deferred_session_restore(self):
        """Deferred session restore — runs AFTER window is visible.

        This was previously called synchronously in __init__ (blocking 9s).
        Now the window shows immediately and this runs on the next event loop tick.
        """
        _t0 = time.time()
        log.info("[STARTUP] Deferred session restore starting...")
        try:
            self._restore_session()
        except Exception as e:
            log.error(f"[STARTUP] Session restore failed: {e}", exc_info=True)
        elapsed = (time.time() - _t0) * 1000
        log.info(f"[STARTUP] Deferred session restore done in {elapsed:.0f}ms")
        # Yield to event loop so Chromium can process pending load events
        QApplication.processEvents()
        # Check for Cortex IDE updates (deferred 5s after startup to avoid blocking UI)
        QTimer.singleShot(5000, self._check_for_updates)

    def _hide_startup_overlay(self):
        """Hide the full-window startup overlay.

        Triggered by any of: sidebar.html loaded, first user interaction
        (mouse/keyboard), or 15s safety timeout — whichever fires first.
        """
        if not hasattr(self, '_startup_overlay') or not self._startup_overlay:
            return
        # Stop the dot animation timer
        if hasattr(self, '_startup_dot_timer'):
            self._startup_dot_timer.stop()
        # Stop the safety timer
        if hasattr(self, '_startup_safety_timer'):
            self._startup_safety_timer.stop()
        # Remove app-level event filter to avoid overhead after startup
        try:
            app = QApplication.instance()
            if app:
                app.removeEventFilter(self)
        except Exception:
            pass
        # Instant hide
        self._startup_overlay.hide()
        self._startup_overlay.deleteLater()
        self._startup_overlay = None
        log.info("[STARTUP] Startup overlay hidden — full UI visible")

    def _apply_dark_title_bar(self):
        """Backward-compat wrapper — applies the title bar for the CURRENT
        theme (the QTimer startup callsite predates light mode)."""
        self._apply_title_bar_theme(self._theme_manager.is_dark)

    def showEvent(self, a0):
        """Re-assert the title bar theme EVERY time the window is shown.

        Bug history: the title bar was applied once via a one-shot 200ms
        timer after startup. Qt can (re)create the native window handle
        (DPI moves, window-flag changes, minimize/restore edge cases) —
        the DWM attribute belongs to the OLD handle, and the state guard
        then refused to re-apply, leaving a light title bar on a dark IDE
        (or vice versa) until the next full restart.
        """
        super().showEvent(a0)
        try:
            self._apply_title_bar_theme(self._theme_manager.is_dark)
        except Exception:
            pass

    def _apply_title_bar_theme(self, is_dark: bool):
        """Apply the native Windows title bar theme via DWM API.

        Uses DWMWA_USE_IMMERSIVE_DARK_MODE only — avoids capsule fragment
        artifacts caused by DWMWA_CAPTION_COLOR/DWMWA_BORDER_COLOR.
        Bug history 1: this was dark-only with a run-once guard, so in light
        mode the window kept a black native title bar above a light UI.
        Bug history 2: the guard only tracked the THEME, not the native
        handle — if Qt recreated the window (new HWND), the attribute was
        gone but the guard said "already applied" and skipped. The guard
        now re-applies whenever the HWND changed.
        Bug history 3: Windows does NOT reliably repaint the caption of an
        already-visible window when this attribute changes — the title bar
        visually kept the OLD theme until a minimize/restore. The
        SetWindowPos(SWP_FRAMECHANGED) nudge forces the repaint now.
        """
        try:
            hwnd = int(self.winId())
            if not hwnd:
                return
            if (getattr(self, '_title_bar_is_dark', None) == is_dark
                    and getattr(self, '_title_bar_hwnd', None) == hwnd):
                return  # same theme on the same native window — nothing to do
            if not hasattr(ctypes, 'windll') or not hasattr(ctypes.windll, 'dwmapi'):
                return
            dwmapi = ctypes.windll.dwmapi
            flag = 1 if is_dark else 0

            # DWMWA_USE_IMMERSIVE_DARK_MODE = 20 (Win10 20H1+ / Win11)
            val20 = ctypes.c_int(flag)
            dwmapi.DwmSetWindowAttribute(hwnd, 20, ctypes.byref(val20), ctypes.sizeof(val20))

            # DWMWA_USE_IMMERSIVE_DARK_MODE = 19 (Win10 1809-1909)
            val19 = ctypes.c_int(flag)
            dwmapi.DwmSetWindowAttribute(hwnd, 19, ctypes.byref(val19), ctypes.sizeof(val19))

            # NOTE: DWMWA_CAPTION_COLOR (34) and DWMWA_BORDER_COLOR (35) are
            # NOT used — they cause capsule fragment artifacts near window controls.

            # Force the caption to repaint NOW (frame-changed nudge) — without
            # this, a visible window keeps painting the old title bar color.
            SWP_NOSIZE, SWP_NOMOVE, SWP_NOZORDER = 0x0001, 0x0002, 0x0004
            SWP_NOACTIVATE, SWP_FRAMECHANGED = 0x0010, 0x0020
            try:
                ctypes.windll.user32.SetWindowPos(
                    hwnd, 0, 0, 0, 0, 0,
                    SWP_NOMOVE | SWP_NOSIZE | SWP_NOZORDER | SWP_NOACTIVATE | SWP_FRAMECHANGED,
                )
            except Exception:
                pass  # nudge is best-effort; attribute is already set

            self._title_bar_is_dark = is_dark
            self._title_bar_hwnd = hwnd
            log.info(f"[DWM] {'Dark' if is_dark else 'Light'} title bar applied to HWND={hwnd}")
        except Exception as e:
            log.warning(f"[DWM] Title bar theme failed: {e}")

    def _on_new_chat(self):
        """Handle new chat request from sidebar, File menu, or keyboard shortcut.

        Routes through summarization so conversation is saved to MEMORY.md
        before clearing — same path as the header '+ New Chat' button.
        """
        if hasattr(self, '_ai_chat') and self._ai_chat:
            self._on_new_chat_requested()
        else:
            log.warning("AI chat not ready, cannot create new chat")
    
    def _set_theme(self, theme: str):
        """Switch theme at RUNTIME — no QApplication.setStyleSheet() call.

        Called from the memory manager (Settings → Appearance → Theme picker)
        and from internal theme toggle shortcuts.

        ── WHY THIS DOESN'T CALL QApplication.setStyleSheet() ──
        That call forces Qt to re-polish every widget in the app, including
        4 embedded QWebEngineView panels (sidebar, editor/chat, terminal,
        memory manager). Measured on a RAM-constrained machine: 75+ REAL
        SECONDS of total IDE freeze — confirmed via [THEME-AUDIT] logs even
        after cutting live chat widgets from 126 to 37, proving widget count
        wasn't the dominant cost. Instead, each panel re-themes itself
        independently and cheaply (JS data-theme pushes, per-widget restyle
        via findChildren) — proven fast in the same logs (sub-millisecond to
        ~500ms). Trade-off: native Qt chrome (menus, toolbars, splitters)
        keeps its old-theme colors until the next restart, where
        _apply_initial_theme() applies the full QSS once, before the window
        is shown, so it's never mid-work.

        ── AUDIT LOGGING ──
        Every batch logs its wall-clock duration and RAM at entry/exit.
        Tag: [THEME-AUDIT] — grep for this tag in cortex.log to see the full trace.
        """
        import time as _time
        import os as _os
        from PyQt6.QtWidgets import QApplication as _App
        from PyQt6.QtCore import QTimer as _QTimer

        t_total_start = _time.perf_counter()

        # ── RAM helper ──
        try:
            import psutil as _psutil
            def _ram() -> float: return _psutil.virtual_memory().percent
        except ImportError:
            def _ram() -> float: return -1.0

        theme = theme if theme in ("dark", "light", "system") else "dark"

        # ── GUARD: Skip if theme didn't actually change ──
        if self._theme_manager.current == theme:
            log.debug(f"[MainWindow] Theme already {theme} — skipping redundant apply")
            return

        is_dark = self._theme_manager.is_dark
        if hasattr(self, '_settings') and self._settings:
            self._settings.theme = theme

        log.info(
            f"[THEME-AUDIT] _set_theme() START  "
            f"from={self._theme_manager.current}  to={theme}  "
            f"RAM={_ram():.1f}%  pid={_os.getpid()}"
        )

        # ── BATCH 1 (0ms): Update theme state — NO QApplication.setStyleSheet() ──
        def _defer_qss_apply():
            """Update the active theme without the app-wide QSS reapply.

            QApplication.setStyleSheet() forces Qt to re-polish every widget,
            including 4 embedded QWebEngineView panels (sidebar, editor/chat,
            terminal, memory manager). Measured on this machine: 75+ REAL
            SECONDS, total IDE freeze, confirmed via [THEME-AUDIT] logs even
            after cutting chat widget count 126->37 (proves widget count was
            not the dominant cost — the WebEngineView re-polish is).
            Each panel already re-themes itself independently via BATCH 2/3/4
            below (data-theme JS pushes, per-widget restyle) — those are the
            ONLY mechanism now. The full QSS still applies once at next
            startup via _apply_initial_theme(), so native Qt chrome (menus,
            toolbars) picks up the new theme on restart.
            """
            t0 = _time.perf_counter()
            log.info(f"[THEME-AUDIT] BATCH-1 state-only  START  RAM={_ram():.1f}%")
            self._theme_manager.set_active_no_qss(theme)
            dt = (_time.perf_counter() - t0) * 1000
            log.info(f"[THEME-AUDIT] BATCH-1 state-only  DONE  dt={dt:.1f}ms  RAM={_ram():.1f}%")

        # ── BATCH 2 (30ms): Panel updates (chat, sidebar, webview) ──
        def _defer_panel_updates():
            """Update lightweight panels."""
            t0 = _time.perf_counter()
            log.info(f"[THEME-AUDIT] BATCH-2 panels  START  RAM={_ram():.1f}%")
            is_dark_current = self._theme_manager.is_dark

            t_a = _time.perf_counter()
            if hasattr(self, '_ai_chat') and self._ai_chat:
                self._ai_chat.set_theme(is_dark_current)
            dt_chat = (_time.perf_counter() - t_a) * 1000
            log.info(f"[THEME-AUDIT] BATCH-2  chat_panel.set_theme  dt={dt_chat:.1f}ms  RAM={_ram():.1f}%")

            t_b = _time.perf_counter()
            if hasattr(self, '_sidebar') and self._sidebar:
                self._sidebar.set_theme(is_dark_current)
            dt_sidebar = (_time.perf_counter() - t_b) * 1000

            t_c = _time.perf_counter()
            if hasattr(self, '_webview_panel') and self._webview_panel:
                self._webview_panel.set_theme(is_dark_current)
            dt_webview = (_time.perf_counter() - t_c) * 1000

            dt_panels = (_time.perf_counter() - t0) * 1000
            log.info(
                f"[THEME-AUDIT] BATCH-2 panels  DONE  "
                f"total={dt_panels:.1f}ms  "
                f"chat={dt_chat:.1f}  sidebar={dt_sidebar:.1f}  webview={dt_webview:.1f}  "
                f"RAM={_ram():.1f}%"
            )

        # ── BATCH 3 (60ms): Terminal updates + native chrome ──
        def _defer_terminal_updates():
            """Update terminal-related widgets and window chrome."""
            t0 = _time.perf_counter()
            log.info(f"[THEME-AUDIT] BATCH-3 terminals  START  RAM={_ram():.1f}%")
            is_dark_current = self._theme_manager.is_dark

            # Native chrome (menu bar, status bar, toolbar icons, title bar)
            # — scoped widget restyles, NOT app-wide QSS. Without this, a
            # live switch left dark menus/status bar and invisible
            # light-gray toolbar icons on a light background.
            self._apply_chrome_theme(is_dark_current)

            t_a = _time.perf_counter()
            if hasattr(self, '_editor_tabs') and self._editor_tabs:
                self._editor_tabs.update_theme(is_dark_current)
            dt_editor = (_time.perf_counter() - t_a) * 1000

            t_b = _time.perf_counter()
            self._update_terminal_theme(is_dark_current)
            dt_term_panel = (_time.perf_counter() - t_b) * 1000

            t_c = _time.perf_counter()
            term_count = self._terminal_tabs.count()
            for i in range(term_count):
                term = self._terminal_tabs.widget(i)
                if isinstance(term, XTermWidget):
                    term.set_theme(is_dark_current)
            dt_terms = (_time.perf_counter() - t_c) * 1000

            dt_total = (_time.perf_counter() - t0) * 1000
            log.info(
                f"[THEME-AUDIT] BATCH-3 terminals  DONE  "
                f"total={dt_total:.1f}ms  "
                f"editor_tabs={dt_editor:.1f}  term_panel={dt_term_panel:.1f}  "
                f"xterm_count={term_count}  xterms={dt_terms:.1f}  "
                f"RAM={_ram():.1f}%"
            )

        # ── BATCH 4 (90ms): Memory manager sync (deferred, non-blocking) ──
        def _defer_memory_manager_sync():
            """Sync memory manager dialog — runJavaScript is async, never blocks.

            Push the RESOLVED appearance (dark/light), never the raw
            "system" value — data-theme="system" matches no CSS rule in
            memory_management.css, so it silently fell back to whatever the
            default look was, independent of actual OS preference. That was
            the "System picks dark even when Windows is in light mode" bug.
            """
            t0 = _time.perf_counter()
            resolved = "dark" if self._theme_manager.is_dark else "light"
            self._push_theme_to_memory_manager(resolved)
            dt = (_time.perf_counter() - t0) * 1000

            dt_total = (_time.perf_counter() - t_total_start) * 1000
            log.info(
                f"[THEME-AUDIT] BATCH-4 memory-mgr  DONE  dt={dt:.1f}ms  "
                f"TOTAL_THEME_SWITCH={dt_total:.1f}ms  "
                f"RAM_final={_ram():.1f}%  "
                f"{'⚠️ SLOW (>2000ms)' if dt_total > 2000 else '✓ ACCEPTABLE'}"
            )

        # Chain the deferred operations with event loop yields between each
        _QTimer.singleShot(0, _defer_qss_apply)
        _QTimer.singleShot(30, _defer_panel_updates)
        _QTimer.singleShot(60, _defer_terminal_updates)
        _QTimer.singleShot(90, _defer_memory_manager_sync)
    

    def _build_ui(self):
        """Build AI-First UI Layout - Codex-style with 2-panel and 4-panel states."""
        self.setWindowTitle("Cortex AI IDE")
        central = QWidget()
        self._central = central
        self.setCentralWidget(central)

        # === FOCUS FRAME SUPPRESSION (theme-neutral, applied at widget level) ===
        # NOTE: All QMenuBar/QMenu/QDialog color styles are now in dark.qss / light.qss
        # theme files loaded via ThemeManager.apply(). This widget-level stylesheet
        # previously contained hardcoded dark colors that conflicted with global QSS
        # during live theme switching, causing UI freezes when app.setStyleSheet()
        # and widget.setStyleSheet() fought over the same selectors.
        self.setStyleSheet("""
            QMenuBar::item:focus { background: transparent; outline: none; border: none; }
            QMenu::item:focus { background: transparent; outline: none; border: none; }
            QWidget#panelToggleBar QPushButton:focus {
                outline: none;
                border: none;
                background: transparent;
            }
        """)

        # === STATE MANAGEMENT ===
        # Always show 4-panel layout with chat ready
        self._chat_started = True

        # === CODEX-STYLE LAYOUT WITH SPLITTERS ===
        main_splitter = CursorSplitter(Qt.Orientation.Horizontal)
        main_splitter.setChildrenCollapsible(False)
        main_splitter.setHandleWidth(6)
        main_splitter.setOpaqueResize(False)  # ← CRITICAL: prevents real-time repaint of 3 QWebEngineViews during drag

        # Create all panels
        # Panel 1: Left Sidebar — HTML-based (sidebar.html + SidebarBridge)
        self._sidebar = SidebarWidget(
            file_manager=self._file_manager,
            git_manager=self._git_manager,
        )
        self._sidebar._main_window = self
        self._sidebar.setMinimumWidth(225)
        self._sidebar.setMaximumWidth(350)
        main_splitter.addWidget(self._sidebar)

        # Panel 2: Chat Panel - Main AI conversation (ALWAYS VISIBLE, flexible)
        self._chat_panel = self._create_chat_panel()
        self._chat_panel.setMinimumWidth(320)
        main_splitter.addWidget(self._chat_panel)

        # Panel 3: Webview Code Editor (Monaco) + Terminal — vertical split
        self._editor_terminal_splitter = CursorSplitter(Qt.Orientation.Vertical)
        self._editor_terminal_splitter.setChildrenCollapsible(False)
        self._editor_terminal_splitter.setHandleWidth(6)
        self._editor_terminal_splitter.setOpaqueResize(False)  # ← same: no real-time QWebEngineView repaint during drag

        # Editor (Monaco webview) — top portion
        self._webview_panel = WebviewPanel()
        self._webview_panel.setMinimumWidth(320)

        # Live Preview — renders a local HTML file inside Cortex's own
        # embedded Chromium (same tech as the Monaco webview), side-by-side
        # with the code. Hidden by default; toggled via View menu / Ctrl+Shift+V.
        from src.ui.components.live_preview_panel import LivePreviewPanel
        self._live_preview_panel = LivePreviewPanel()
        self._live_preview_panel.setMinimumWidth(280)
        self._live_preview_panel.closed.connect(lambda: self._toggle_live_preview(show=False))
        self._live_preview_hidden = True

        self._editor_preview_splitter = CursorSplitter(Qt.Orientation.Horizontal)
        self._editor_preview_splitter.setChildrenCollapsible(False)
        self._editor_preview_splitter.setHandleWidth(6)
        self._editor_preview_splitter.setOpaqueResize(False)
        self._editor_preview_splitter.addWidget(self._webview_panel)
        self._editor_preview_splitter.addWidget(self._live_preview_panel)
        self._live_preview_panel.setVisible(False)
        self._editor_preview_splitter.setSizes([1000, 0])  # preview hidden on startup

        self._editor_terminal_splitter.addWidget(self._editor_preview_splitter)

        # Terminal Tab Widget — holds multiple terminal tabs, hidden by default
        # (toggle via Ctrl+J or toolbar button; new tabs via Ctrl+Shift+`)
        self._terminal_tabs = QTabWidget()
        self._terminal_tabs.setTabBar(CleanTabBar(self._terminal_tabs))
        self._terminal_tabs.setTabsClosable(True)
        self._terminal_tabs.setDocumentMode(True)
        self._terminal_tabs.setMovable(True)
        self._terminal_tabs.setVisible(False)
        # FIX: minimum height must be 0 at build time. Setting 120 while
        # hidden + splitter size 0 causes a layout conflict on first
        # showMaximized() — Qt briefly allocates 120px to the terminal
        # (stealing from editor), causing a visible window-wide flicker.
        # _toggle_terminal_panel(show=True) sets it to 120 when actually shown.
        self._terminal_tabs.setMinimumHeight(0)
        self._terminal_tabs.tabCloseRequested.connect(self._close_terminal_tab)
        
        # Hide the PyQt6 tab bar completely — terminal.html provides
        # the header with terminal name, + New, Kill, Clear, Restart buttons.
        self._terminal_tabs.tabBar().setVisible(False)
        self._terminal_tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 0px;
                background: #1e1e1e;
                top: 0px;
            }
            QTabWidget::tab-bar {
                height: 0px;
                max-height: 0px;
                spacing: 0px;
            }
        """)
        
        self._editor_terminal_splitter.addWidget(self._terminal_tabs)
        self._editor_terminal_splitter.setSizes([500, 0])  # Terminal hidden on startup

        # Terminal tabs — first terminal is created lazily on first show
        # Use _current_terminal() to get the active terminal widget

        main_splitter.addWidget(self._editor_terminal_splitter)
        for idx in range(main_splitter.count()):
            main_splitter.setCollapsible(idx, False)
        for idx in range(self._editor_terminal_splitter.count()):
            self._editor_terminal_splitter.setCollapsible(idx, False)

        # Set initial sizes (proportions) — 3-panel layout
        main_splitter.setSizes([300, 500, 500])

        # Stretch factors: sidebar fixed, chat+editor share extra space equally
        main_splitter.setStretchFactor(0, 0)  # sidebar — doesn't grow
        main_splitter.setStretchFactor(1, 1)  # chat — grows proportionally
        main_splitter.setStretchFactor(2, 1)  # editor — grows proportionally

        # Add splitter directly to main layout (native title bar handles close/max/min)
        root_layout = QVBoxLayout(central)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)
        root_layout.addWidget(main_splitter, 1)

        # Store reference for later access
        self._main_splitter = main_splitter
        self._sync_splitter_handles(main_splitter)
        self._sync_splitter_handles(self._editor_terminal_splitter)
        self._sync_splitter_handles(self._editor_preview_splitter)
        # Auto re-sync handles whenever splitter moves
        # REMOVED: splitterMoved → _sync_splitter_handles during drag
        # This caused a relayout storm on every pixel of drag movement.
        # CursorSplitHandle manages its own size; _sync only needed on collapse/expand.
        # main_splitter.splitterMoved.connect(lambda: self._sync_splitter_handles(main_splitter))
        self._editor_terminal_splitter.splitterMoved.connect(self._on_editor_term_moved)

        # Panel toggle state tracking
        self._left_sidebar_hidden = False
        self._chat_panel_hidden = False
        self._code_panel_hidden = False
        self._terminal_panel_hidden = True  # Hidden by default — user toggles open via Ctrl+J or toolbar
        self._summary_panel_hidden = False
        self._git_panel_hidden = False

        # Store minimum widths for panel toggle restore
        self._left_sidebar_min_width = 300
        self._chat_panel_min_width = 320
        self._code_panel_min_width = 500
        self._terminal_panel_min_height = 150

        # Keep old components for backward compatibility
        self._ai_splitter = None  # Replaced by 4-panel layout
        self._editor_tabs = EditorTabWidget()
        self._editor_tabs.setMinimumSize(200, 150)
        self._editor_tabs.hide()  # Hidden in Codex mode, shown in editor mode

        # Find/Replace Dialog
        self._find_replace_dialog = FindReplaceDialog(self)
        self._find_replace_dialog.find_requested.connect(self._on_find_requested)
        self._find_replace_dialog.replace_requested.connect(self._on_replace_requested)
        self._find_replace_dialog.replace_all_requested.connect(self._on_replace_all_requested)

        # ════════════════════════════════════════════════════════════════
        # STARTUP OVERLAY — child widget covering the main window
        # ════════════════════════════════════════════════════════════════
        # Bug history: this was a TOP-LEVEL frameless always-on-top window
        # sized to the ENTIRE SCREEN. In compiled builds it appeared as a
        # separate black "app" before the Cortex window existed, and it
        # blacked out the whole desktop (covering OTHER apps) until first
        # user interaction or a 15s timeout. A child widget masks only the
        # Cortex window itself; raise_() keeps it above the panels (the old
        # 1-frame z-order blink is handled by re-raising after first show).
        self._startup_overlay = QWidget(self)
        self._startup_overlay.setStyleSheet("background: #141414; border: none;")
        _ov_layout = QVBoxLayout(self._startup_overlay)
        _ov_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        _ov_layout.setSpacing(20)

        # Brand wordmark
        _ov_brand = QLabel("C O R T E X")
        _ov_brand.setStyleSheet(
            "color: #5B8CFF; font-size: 22px; font-weight: 700;"
            "font-family: 'Segoe UI', sans-serif; letter-spacing: 8px;"
            "background: transparent;"
        )
        _ov_brand.setAlignment(Qt.AlignmentFlag.AlignCenter)
        _ov_layout.addWidget(_ov_brand)

        # Loading status label
        self._startup_status = QLabel("Starting...")
        self._startup_status.setStyleSheet(
            "color: #555; font-size: 13px; font-family: 'Segoe UI', sans-serif;"
            "background: transparent;"
        )
        self._startup_status.setAlignment(Qt.AlignmentFlag.AlignCenter)
        _ov_layout.addWidget(self._startup_status)

        # Animated dots (zero CPU — text-only)
        self._startup_dot_timer = QTimer(self)
        self._startup_dot_count = 0
        def _animate_startup_dots():
            self._startup_dot_count = (self._startup_dot_count + 1) % 4
            self._startup_status.setText(f"Starting{'.' * self._startup_dot_count}")
        self._startup_dot_timer.timeout.connect(_animate_startup_dots)
        self._startup_dot_timer.start(500)

        # Cover the (not-yet-shown) main window; resizeEvent keeps it sized.
        self._startup_overlay.setGeometry(self.rect())
        self._startup_overlay.raise_()
        self._startup_overlay.show()

        # Safety timeout: force-hide after 15s regardless of interaction
        self._startup_safety_timer = QTimer(self)
        self._startup_safety_timer.setSingleShot(True)
        self._startup_safety_timer.timeout.connect(self._hide_startup_overlay)
        self._startup_safety_timer.start(15000)

        # Install app-wide event filter to detect first user interaction
        # (mouse move, click, keypress) — that's when we hide the overlay
        self._startup_overlay_dismissing = False
        app = QApplication.instance()
        if app:
            app.installEventFilter(self)

    # ------------------------------------------------------------------
    # Codex-Style 4-Panel Layout Methods
    # ------------------------------------------------------------------

    def _create_chat_panel(self) -> QWidget:
        """Create Chat Panel (flexible) - Main AI conversation."""
        panel = QWidget()
        panel.setObjectName("chatPanel")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        bg_color = "#1e1e1e"

        panel.setStyleSheet(f"""
            QWidget#chatPanel {{
                background-color: {bg_color};
            }}
        """)

        # ── Native chat backend (no Chromium, no webview) ──
        from src.ui.chat_panel import ChatPanel as NativeChatPanel
        self._ai_chat = NativeChatPanel()
        self._is_native_chat = True

        # NOTE: Chat persistence is set up in _on_project_opened() AFTER
        # the project path is known. Calling it here would use os.getcwd()
        # which is the Cortex Engine directory, not the user's project.

        layout.addWidget(self._ai_chat, 1)

        return panel

    # NOTE: _create_review_panel, _on_summary_tab_clicked, _on_review_tab_clicked,
    # _update_git_summary, _on_git_status_ready, _set_no_git_status, _check_github_cli,
    # _update_review_file_list, _update_review_file_list_from_stats, _create_file_diff_item,
    # _create_file_diff_item_from_stats, _get_file_diff_stats — REMOVED.
    # Git review/summary UI is now HTML-based in sidebar.html.

    def eventFilter(self, obj, event):
        """Event filter for keyboard shortcuts + startup overlay detection."""
        from PyQt6.QtCore import QEvent
        # Hide startup overlay on first user interaction (mouse/keyboard)
        if (getattr(self, '_startup_overlay', None)
                and not getattr(self, '_startup_overlay_dismissing', False)):
            if event.type() in (QEvent.Type.MouseMove, QEvent.Type.MouseButtonPress,
                                 QEvent.Type.KeyPress, QEvent.Type.Wheel):
                self._startup_overlay_dismissing = True
                QTimer.singleShot(100, self._hide_startup_overlay)
        return super().eventFilter(obj, event)

    def _on_chat_hidden_changed(self, hidden: bool):
        """Handle chat panel visibility change."""


    # ------------------------------------------------------------------
    # Menu
    # ------------------------------------------------------------------
    def _build_menu(self):
        mb = self.menuBar()

        # File
        file_menu = mb.addMenu("File")
        self._add_action(file_menu, "New Chat", self._on_new_chat, "Ctrl+N")
        file_menu.addSeparator()
        self._add_action(file_menu, "Open File...", self._open_file_dialog, "Ctrl+Shift+O")
        self._add_action(file_menu, "Open Folder...", self._open_folder_dialog, "Ctrl+O")
        file_menu.addSeparator()
        self._add_action(file_menu, "Save", self._save_current, "Ctrl+S")
        self._add_action(file_menu, "Save All", self._save_all, "Ctrl+Shift+S")
        file_menu.addSeparator()
        self._add_action(file_menu, "Exit", self.close, "Alt+F4")

        # Edit
        edit_menu = mb.addMenu("Edit")
        self._add_action(edit_menu, "Undo", self._undo, "Ctrl+Z")
        self._add_action(edit_menu, "Redo", self._redo, "Ctrl+Y")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Cut", lambda: self._current_editor_action("cut"), "Ctrl+X")
        self._add_action(edit_menu, "Copy", lambda: self._current_editor_action("copy"), "Ctrl+C")
        self._add_action(edit_menu, "Paste", lambda: self._current_editor_action("paste"), "Ctrl+V")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Find...", self._show_find, "")
        self._add_action(edit_menu, "Find and Replace...", self._show_find_replace, "")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Rename...", self._rename_file, "F2")
        edit_menu.addSeparator()
        self._add_action(edit_menu, "Go to Line...", self._go_to_line, "")

        # View
        view_menu = mb.addMenu("View")
        self._add_action(view_menu, "Toggle Sidebar", self._toggle_sidebar, "Ctrl+B")
        view_menu.addSeparator()
        self._add_action(view_menu, "Toggle AI Chat", lambda: self._toggle_ai_chat_panel(getattr(self, '_chat_panel_hidden', False)), "")
        self._add_action(view_menu, "Toggle Code Editor", lambda: self._toggle_code_panel(getattr(self, '_code_panel_hidden', False)), "")
        self._add_action(view_menu, "Toggle Live Preview", self._toggle_live_preview_shortcut, "Ctrl+Shift+V")
        view_menu.addSeparator()
        self._add_action(view_menu, "Toggle Full Screen", self._toggle_fullscreen, "F11")

        # AI
        ai_menu = mb.addMenu("AI")
        self._add_action(ai_menu, "Explain Code", lambda: self._ai_action("explain"), "Ctrl+Shift+E")
        self._add_action(ai_menu, "Refactor Code", lambda: self._ai_action("refactor"), "Ctrl+Shift+R")
        self._add_action(ai_menu, "Write Tests", lambda: self._ai_action("tests"), "Ctrl+Shift+U")
        self._add_action(ai_menu, "Debug Help", lambda: self._ai_action("debug"), "Ctrl+Shift+H")
        ai_menu.addSeparator()
        
        # Phase 1, 2, 3 Integration: Agent Mode submenu
        mode_menu = ai_menu.addMenu("Agent Mode")
        self._add_action(mode_menu, "Build Mode", lambda: self._set_agent_mode("build"), "")
        self._add_action(mode_menu, "Explore Mode", lambda: self._set_agent_mode("explore"), "")
        self._add_action(mode_menu, "Debug Mode", lambda: self._set_agent_mode("debug"), "")
        self._add_action(mode_menu, "Plan Mode", lambda: self._set_agent_mode("plan"), "")
        ai_menu.addSeparator()
        
        # Phase 3 Integration: Skills and MCP
        self._add_action(ai_menu, "Browse Skills...", self._show_skills_browser, "")
        ai_menu.addSeparator()
        
        # Phase 4 Integration: TODO, Permission, and GitHub
        todo_menu = ai_menu.addMenu("Tasks & TODOs")
        self._add_action(todo_menu, "View Tasks...", self._show_todo_manager, "")
        self._add_action(todo_menu, "Add Task...", self._add_todo_task, "")
        todo_menu.addSeparator()
        self._add_action(todo_menu, "Complete Task", self._complete_todo_task, "")
        
        self._add_action(ai_menu, "Permission Settings...", self._show_permission_settings, "")
        self._add_action(ai_menu, "Memory Manager...", self._show_memory_manager, "Ctrl+Shift+M")
        ai_menu.addSeparator()
        
        self._add_action(ai_menu, "AI Chat Focus", self._focus_ai_chat, "Ctrl+Shift+A")
        ai_menu.addSeparator()
        if not getattr(self, '_is_native_chat', False):
            self._add_action(ai_menu, "Clear Chat", self._ai_chat.clear_chat, "")

        # Terminal
        term_menu = mb.addMenu("Terminal")
        self._add_action(term_menu, "New Terminal", lambda: self._new_terminal(show_panel=True), "Ctrl+Shift+`")
        self._add_action(term_menu, "Kill Terminal", self._kill_current_terminal, "")
        term_menu.addSeparator()
        self._add_action(term_menu, "Toggle Terminal Panel", self._toggle_terminal, "Ctrl+J")

        # Window
        window_menu = mb.addMenu("Window")
        self._add_action(window_menu, "Minimize", self._minimize_window, "Ctrl+M")
        self._add_action(window_menu, "Zoom", self._zoom_window, "")
        # No accelerator on window-Close: Ctrl+F4 used to be bound here and
        # silently quit the WHOLE IDE mid-work (Ctrl+F4 is "close tab" in every
        # other Windows editor — pure muscle-memory data loss). Alt+F4 / the X
        # button still close the window via the OS.
        self._add_action(window_menu, "Close", self._close_window, "")

        # Close tab / close all tabs — QShortcuts (not menu items, to avoid Ctrl+W
        # colliding with Monaco editor's own Ctrl+W close-tab command).
        from PyQt6.QtGui import QShortcut
        self._s_close_tab = QShortcut(QKeySequence("Ctrl+W"), self)
        self._s_close_tab.activated.connect(self._close_current_tab)
        # Ctrl+F4 = close current tab (standard on Windows), same as Ctrl+W
        self._s_close_tab_f4 = QShortcut(QKeySequence("Ctrl+F4"), self)
        self._s_close_tab_f4.activated.connect(self._close_current_tab)
        self._s_close_all = QShortcut(QKeySequence("Ctrl+Shift+W"), self)
        self._s_close_all.activated.connect(self._close_all_tabs)
        # Next/previous tab navigation
        self._s_next_tab = QShortcut(QKeySequence("Ctrl+Tab"), self)
        self._s_next_tab.activated.connect(self._next_tab)
        self._s_prev_tab = QShortcut(QKeySequence("Ctrl+Shift+Tab"), self)
        self._s_prev_tab.activated.connect(self._prev_tab)

        # Format Code — Monaco-native or legacy fallback
        self._s_format_code = QShortcut(QKeySequence("Shift+Alt+F"), self)
        self._s_format_code.activated.connect(self._format_code)

        # Debug Console — toggle terminal panel
        self._s_debug_console = QShortcut(QKeySequence("Ctrl+Alt+D"), self)
        self._s_debug_console.activated.connect(lambda: self._toggle_terminal_panel(
            show=not getattr(self, '_terminal_panel_hidden', False)))

        # Help
        help_menu = mb.addMenu("Help")
        self._add_action(help_menu, "Cortex Documentation", self._open_documentation, "")
        self._add_action(help_menu, "What's New", self._show_whats_new, "")
        self._add_action(help_menu, "Automations", self._show_automations, "")
        self._add_action(help_menu, "Local Environments", self._show_local_envs, "")
        self._add_action(help_menu, "Worktrees", self._show_worktrees, "")
        self._add_action(help_menu, "Skills", self._show_skills_help, "")
        self._add_action(help_menu, "Model Context Protocol", self._show_mcp_help, "")
        self._add_action(help_menu, "Troubleshooting", self._show_troubleshooting, "")
        help_menu.addSeparator()
        self._add_action(help_menu, "Send Feedback", self._send_feedback, "")
        self._add_action(help_menu, "Start Trace Recording", self._start_trace, "")
        help_menu.addSeparator()
        self._add_action(help_menu, "Keyboard Shortcuts", self._show_keyboard_shortcuts, "F1")
        help_menu.addSeparator()
        self._add_action(help_menu, "About Cortex", self._show_about, "")

        # ═══════════════════════════════════════════════════════════════════════
        # Panel Toggle Button Group — right corner of menu bar
        # ═══════════════════════════════════════════════════════════════════════
        self._panel_toggle_bar = self._build_panel_toggle_bar()

        # QMenuBar doesn't support custom widgets in the same way, so we add
        # them as a QAction with a QWidget. We use a spacer trick:
        # Create a spacer action to push buttons right.
        spacer_action = QAction(self)
        spacer_action.setVisible(False)  # Won't show but allows layout control
        # Instead, we'll place buttons in a QWidget that sits on the right
        # by using a custom approach: a QWidget placed next to menu bar via layout.
        # For simplicity, add directly to menuBar's cornerWidget.
        self.menuBar().setCornerWidget(self._panel_toggle_bar, Qt.Corner.TopRightCorner)
        self._panel_toggle_bar.setFocusPolicy(Qt.FocusPolicy.NoFocus)

    def _add_action(self, menu, text, slot, shortcut=""):
        action = QAction(text, self)
        if shortcut:
            action.setShortcut(QKeySequence(shortcut))
        action.triggered.connect(slot)
        menu.addAction(action)
        return action

    # ------------------------------------------------------------------
    # Toolbar
    # ------------------------------------------------------------------
    # Toolbar removed in AI-first mode - replaced by TopNavBar component
    # def _build_toolbar(self):


    # ------------------------------------------------------------------
    # Panel Toggle Buttons
    # ------------------------------------------------------------------
    def _panel_toggle_bar_qss(self, is_dark: bool) -> str:
        """Stylesheet for the panel-toggle bar — hover tints must match the
        theme (white tint is invisible-to-wrong on a light background)."""
        hover = "rgba(255, 255, 255, 0.08)" if is_dark else "rgba(0, 0, 0, 0.08)"
        pressed = "rgba(255, 255, 255, 0.15)" if is_dark else "rgba(0, 0, 0, 0.15)"
        return f"""
            QWidget#panelToggleBar {{
                background: transparent;
                padding: 0px 8px;
            }}
            QPushButton {{
                background: transparent;
                border: none;
                border-radius: 4px;
                padding: 2px;
                margin: 1px 0px;
                outline: none;
            }}
            QPushButton:hover {{
                background: {hover};
            }}
            QPushButton:pressed {{
                background: {pressed};
            }}
            QPushButton:focus {{
                outline: none;
                border: none;
                background: transparent;
            }}
        """

    def _build_panel_toggle_bar(self) -> QWidget:
        """Build a horizontal bar of 4 toggle buttons — one per panel group."""
        bar = QWidget()
        bar.setObjectName("panelToggleBar")
        bar.setFixedHeight(30)
        is_dark = self._theme_manager.is_dark
        bar.setStyleSheet(self._panel_toggle_bar_qss(is_dark))
        self._panel_toggle_bar = bar

        layout = QHBoxLayout(bar)
        layout.setContentsMargins(0, 0, 4, 0)
        layout.setSpacing(1)

        from src.utils.icons import make_icon

        icon_size = 22
        btn_size = 26

        # Icon color is theme-dependent and refreshable: light-gray icons on
        # a light toolbar were invisible (the "Run/Hide-Sidebar buttons
        # disappear in light mode" bug). _toolbar_icon_refreshers lets
        # _apply_chrome_theme() re-tint every icon on a live theme switch.
        self._toolbar_icon_color = "#c8c8c8" if is_dark else "#3c3c3c"
        self._toolbar_icon_refreshers = []

        def _make_toggle(visible_icon: str, hidden_icon: str, tooltip_v: str, tooltip_h: str,
                         is_visible_getter, toggle_fn):
            """Single toggle button — switches icon/tooltip when panel visibility changes."""
            btn = QPushButton()
            btn.setFixedSize(btn_size, btn_size)
            btn.setFlat(True)
            btn.setFocusPolicy(Qt.FocusPolicy.NoFocus)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            # WA_NoSystemBackground REMOVED — causes ACCESS VIOLATION on Windows
            btn.setAttribute(Qt.WidgetAttribute.WA_MacShowFocusRect, False)

            _visible = is_visible_getter()

            def _refresh_icon():
                btn.setIcon(make_icon(visible_icon if _visible else hidden_icon,
                                      self._toolbar_icon_color, icon_size))

            _refresh_icon()
            btn.setToolTip(tooltip_v if _visible else tooltip_h)
            self._toolbar_icon_refreshers.append(_refresh_icon)

            def on_click():
                nonlocal _visible
                _visible = not _visible
                _refresh_icon()
                btn.setToolTip(tooltip_v if _visible else tooltip_h)
                toggle_fn(_visible)

            btn.clicked.connect(on_click)
            return btn

        # 1. Left Sidebar toggle
        layout.addWidget(_make_toggle(
            "panel-left-sidebar-visible", "panel-left-sidebar-hidden",
            "Hide Left Sidebar", "Show Left Sidebar",
            lambda: not getattr(self, '_left_sidebar_hidden', False),
            lambda v: self._toggle_left_sidebar(v)
        ))

        # 2. AI Chat toggle
        layout.addWidget(_make_toggle(
            "panel-ai-chat-visible", "panel-ai-chat-hidden",
            "Hide AI Chat", "Show AI Chat",
            lambda: not getattr(self, '_chat_panel_hidden', False),
            lambda v: self._toggle_ai_chat_panel(v)
        ))

        # 3. Code Editor toggle
        layout.addWidget(_make_toggle(
            "panel-code-visible", "panel-code-hidden",
            "Hide Code Editor", "Show Code Editor",
            lambda: not getattr(self, '_code_panel_hidden', False),
            lambda v: self._toggle_code_panel(v)
        ))

        # 4. Terminal panel toggle (bottom panel in editor split)
        layout.addWidget(_make_toggle(
            "panel-terminal-visible", "panel-terminal-hidden",
            "Hide Terminal", "Show Terminal",
            lambda: not getattr(self, '_terminal_panel_hidden', False),
            lambda v: self._toggle_terminal_panel(v)
        ))

        # Play/Run button — runs the active file (HTML → Live Server, Python/JS → terminal)
        play_btn = QPushButton()
        play_btn.setFixedSize(btn_size, btn_size)
        play_btn.setFlat(True)
        play_btn.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        play_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        play_btn.setAttribute(Qt.WidgetAttribute.WA_MacShowFocusRect, False)

        def _refresh_play_icon():
            play_btn.setIcon(make_icon("play", self._toolbar_icon_color, icon_size))

        _refresh_play_icon()
        self._toolbar_icon_refreshers.append(_refresh_play_icon)
        play_btn.setToolTip("Run File (Ctrl+F5)")
        play_btn.clicked.connect(self._run_file)
        layout.addWidget(play_btn)

        # 5. Review/Summary/Git panel toggle (all 3 tabs share one panel)
        layout.addWidget(_make_toggle(
            "panel-review-visible", "panel-review-hidden",
            "Hide Review Panel", "Show Review Panel",
            lambda: not getattr(self, '_review_panel_hidden', False),
            lambda v: self._toggle_review_panel(v)
        ))

        return bar

    def _make_spacer(self) -> QWidget:
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        return spacer

    # ------------------------------------------------------------------
    # Status Bar
    # ------------------------------------------------------------------
    def _restyle_status_bar(self, is_dark: bool):
        """Apply theme colors to the status bar.

        Bug history: _build_status_bar hardcoded a DARK widget-level
        stylesheet (#1e1e1e / #cccccc). Widget stylesheets OVERRIDE the app
        QSS, so even a light-theme restart could never restyle the status
        bar — it stayed dark forever. Light backgrounds must use dark fonts.
        """
        sb = self.statusBar()
        if is_dark:
            bg, fg, border = "#1e1e1e", "#cccccc", "#2a2a2a"
            lsp_color, ver_color = "#ff5555", "#6272a4"
        else:
            # Warm Anthropic/Claude palette — MATCHES editor.html,
            # memory_management.css and sidebar.html light mode exactly
            # (bg #ECE9E0/#E4E1D8, text #1A1814, muted #6B6860, terracotta
            # accent #C96A3E). A cool-gray status bar looked like a
            # different app bolted onto the rest of the light theme.
            bg, fg, border = "#E4E1D8", "#1A1814", "#CCC9C0"
            lsp_color, ver_color = "#B83232", "#6B6860"
        sb.setStyleSheet(f"""
            QStatusBar {{
                background-color: {bg};
                color: {fg};
                border-top: 1px solid {border};
                font-size: 12px;
            }}
            QLabel {{
                color: {fg};
                padding: 0 6px;
            }}
        """)
        if getattr(self, '_status_lsp_lbl', None):
            self._status_lsp_lbl.setStyleSheet(
                f"color: {lsp_color}; font-size: 11px; padding: 0 6px;")
        if getattr(self, '_status_version_lbl', None):
            self._status_version_lbl.setStyleSheet(
                f"color: {ver_color}; font-size: 11px;")

    def _build_status_bar(self):
        sb = self.statusBar()
        self._status_file = QLabel("  No file open")
        self._status_cursor = QLabel("Ln 1, Col 1")
        self._status_lang = QLabel("Plain Text")
        self._status_ai = QLabel("AI: Ready")

        for lbl in [self._status_file, self._status_cursor, self._status_lang, self._status_ai]:
            sb.addWidget(lbl)

        # LSPs are disabled (system removed)
        self._status_lsp_lbl = QLabel("  LSPs: disabled  ")
        sb.addWidget(self._status_lsp_lbl)

        # ── Version ──
        _app_version = QApplication.instance().applicationVersion() or "2.8.1"
        self._status_version_lbl = QLabel(f"  Cortex AI IDE v{_app_version} ")
        sb.addPermanentWidget(self._status_version_lbl)

        self._restyle_status_bar(self._theme_manager.is_dark)
        self._update_status_cursor(1, 1)

    def _update_status_cursor(self, line: int, col: int):
        self._status_cursor.setText(f"Ln {line}, Col {col}")

    def _update_status_file(self, filepath: str = None):
        if filepath:
            self._status_file.setText(f"  {shorten_path(filepath)}")
            self._status_lang.setText(detect_language(filepath).title())
        else:
            self._status_file.setText("  No file open")
            self._status_lang.setText("Plain Text")

    def _sync_window_title(self):
        """Sync window title — prepend ● when files are modified (VS Code-style)."""
        modified = getattr(self, '_modified_files', set())
        project_name = self._project_manager.root.name if self._project_manager.root else ""
        if modified and project_name:
            self.setWindowTitle(f"● Cortex AI IDE — {project_name}")
        elif project_name:
            self.setWindowTitle(f"Cortex AI IDE — {project_name}")
        else:
            self.setWindowTitle("Cortex AI IDE")

    def _restyle_menu_bar(self, is_dark: bool):
        """Scoped stylesheet for the menu bar + its popup menus.

        At startup the global QSS covers this; at RUNTIME theme switches we
        deliberately never call QApplication.setStyleSheet() (75s freeze),
        so the menu bar needs its own cheap widget-level restyle. Light
        background gets near-black text — never light-hash gray.
        """
        if is_dark:
            bg, fg, hover, border = "#1e1e1e", "#cccccc", "#2d2d2d", "#3c3c3c"
            popup_bg, sel_bg, disabled = "#252526", "#094771", "#6a6a6a"
        else:
            bg, fg, hover, border = "#e8e8e8", "#1f2328", "#d8d8d8", "#d0d0d0"
            popup_bg, sel_bg, disabled = "#ffffff", "#cce4ff", "#9a9a9a"
        qss = f"""
            QMenuBar {{
                background-color: {bg};
                color: {fg};
                border-bottom: 1px solid {border};
            }}
            QMenuBar::item {{
                background: transparent;
                color: {fg};
                padding: 4px 10px;
            }}
            QMenuBar::item:selected {{ background: {hover}; }}
            QMenuBar::item:disabled {{ color: {disabled}; }}
            QMenuBar::item:focus {{ background: transparent; outline: none; border: none; }}
            QMenu {{
                background-color: {popup_bg};
                color: {fg};
                border: 1px solid {border};
            }}
            QMenu::item {{ color: {fg}; }}
            QMenu::item:selected {{ background-color: {sel_bg}; color: {fg}; }}
            QMenu::item:disabled {{ color: {disabled}; }}
            QMenu::item:focus {{ background: transparent; outline: none; border: none; }}
            QMenu::separator {{ height: 1px; background: {border}; margin: 4px 8px; }}
        """
        # Palette, not just QSS: Qt's native Windows styles (windowsvista /
        # windows11 — the default on this platform) partially fall back to
        # QPalette::WindowText/Text for menu item text even when a QSS
        # `color:` rule is present, especially for pre-existing QMenu popup
        # objects that were polished under the OLD theme. That is exactly
        # the bug reported: Window > Minimize/Zoom/Close rendered washed-out
        # gray after a LIVE switch, but crisp and correct after a restart
        # (when the menu is freshly built with the current QSS already the
        # only one ever applied). Setting the palette explicitly is a
        # belt-and-suspenders fix that does not depend on the native style
        # fully honoring QSS text color.
        from PyQt6.QtGui import QColor as _QColor, QPalette as _QPalette
        pal = self.menuBar().palette()
        pal.setColor(_QPalette.ColorRole.Window, _QColor(bg))
        pal.setColor(_QPalette.ColorRole.WindowText, _QColor(fg))
        pal.setColor(_QPalette.ColorRole.Text, _QColor(fg))
        pal.setColor(_QPalette.ColorRole.Base, _QColor(popup_bg))
        pal.setColor(_QPalette.ColorGroup.Disabled, _QPalette.ColorRole.WindowText, _QColor(disabled))
        pal.setColor(_QPalette.ColorGroup.Disabled, _QPalette.ColorRole.Text, _QColor(disabled))
        pal.setColor(_QPalette.ColorRole.Highlight, _QColor(sel_bg))
        pal.setColor(_QPalette.ColorRole.HighlightedText, _QColor(fg))

        self.menuBar().setStyleSheet(qss)
        self.menuBar().setPalette(pal)
        # Each top-level QMenu ("File", "Window", ...) — and nested submenus
        # like "Agent Mode" inside "AI" — is a separate top-level POPUP
        # widget once shown. Bug history: setting the stylesheet on
        # menuBar() alone left already-constructed QMenu popups showing
        # washed-out/disabled-looking text after a LIVE switch (correct
        # only after restart, when menus are freshly built under the QSS
        # already in effect) — Qt does not reliably re-cascade a stylesheet
        # into a popup window's existing menu objects. Re-apply directly to
        # every QMenu found under the menu bar, including nested ones.
        for menu in self.menuBar().findChildren(QMenu):
            menu.setStyleSheet(qss)
            menu.setPalette(pal)

    def _apply_chrome_theme(self, is_dark: bool):
        """Re-theme the native window chrome: menu bar, status bar, panel
        toggle bar (+ its icons), and the OS title bar.

        These are a handful of individual widgets — cheap widget-level
        restyles, NOT an app-wide setStyleSheet. This is what makes menus,
        status bar and toolbar icons follow a LIVE theme switch (they used
        to keep dark colors / invisible light-gray icons until restart).
        """
        try:
            self._restyle_menu_bar(is_dark)
            self._restyle_status_bar(is_dark)
            if getattr(self, '_panel_toggle_bar', None):
                self._panel_toggle_bar.setStyleSheet(self._panel_toggle_bar_qss(is_dark))
            self._toolbar_icon_color = "#c8c8c8" if is_dark else "#3c3c3c"
            for _refresh in getattr(self, '_toolbar_icon_refreshers', []):
                try:
                    _refresh()
                except RuntimeError:
                    pass  # button was deleted with its parent panel
            self._apply_title_bar_theme(is_dark)
        except Exception as e:
            log.warning(f"[THEME] Chrome restyle failed: {e}")

    def _apply_initial_theme(self):
        """Apply the saved (or default) theme to all panels at startup."""
        from PyQt6.QtWidgets import QApplication as _App
        saved = self._settings.theme if hasattr(self._settings, 'theme') else "dark"
        # Support both dark and light themes. This is the ONLY place
        # QApplication.setStyleSheet() runs now — once at startup, before
        # the window is shown, so its cost (measured 75+s under RAM
        # pressure) delays first paint but can never freeze an active
        # session. Runtime switches use _set_theme() -> set_active_no_qss().
        self._theme_manager.apply(saved, _App.instance(), freeze_widget=self)
        is_dark = self._theme_manager.is_dark
        log.info(f"[MainWindow] Initial theme: {saved} (is_dark={is_dark})")

        # Native chrome: menu bar, status bar, toolbar icons, OS title bar
        self._apply_chrome_theme(is_dark)

        # Propagate to all panels
        if not getattr(self, '_is_native_chat', False):
            self._ai_chat.set_theme(is_dark)
        self._sidebar.set_theme(is_dark)
        
        # Apply to webview editor
        if hasattr(self, '_webview_panel'):
            self._webview_panel.set_theme(is_dark)
        if hasattr(self, '_live_preview_panel'):
            self._live_preview_panel.set_theme(is_dark)
        # Also apply to legacy editor tabs (welcome/PDF/image)
        self._editor_tabs.update_theme(is_dark)
        
        # Apply to terminal tab bar
        if isinstance(self._terminal_tabs.tabBar(), CleanTabBar):
            self._terminal_tabs.tabBar().set_dark(is_dark)
        
        # Style the legacy tab widget panels
        self._update_terminal_theme(is_dark)
        
        # Apply to all terminal widgets
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term.set_theme(is_dark)
        
        # Apply global syntax highlighting fonts and colors
        self._apply_syntax_highlighting_fonts()

    def _update_terminal_theme(self, is_dark: bool):
        """Ensure terminal tab bar stays hidden (terminal.html provides its own header).
        
        NOTE: QTabWidget::pane background/border is now handled by the global QSS
        theme files (dark.qss / light.qss). Widget-level setStyleSheet here would
        conflict with the global QSS during live theme switching, causing repaint storms.
        """
        # Keep terminal tab bar hidden — terminal.html provides the header
        self._terminal_tabs.tabBar().setVisible(False)

    def _push_theme_to_memory_manager(self, theme: str):
        """Push theme to the memory manager dialog if it is open.

        Must run on the GUI thread: topLevelWidgets(), isVisible() and
        QtWebEngine's runJavaScript() are all main-thread-only Qt calls —
        invoking them from a background thread crashed the app on theme
        switch. runJavaScript() is internally asynchronous, so this call
        never blocks the UI anyway.
        """
        try:
            from PyQt6.QtWidgets import QApplication as _App
            from src.ui.dialogs.memory_manager import MemoryManagerDialog
            app = _App.instance()
            if not app:
                return
            for widget in app.topLevelWidgets():
                if isinstance(widget, MemoryManagerDialog) and widget.isVisible():
                    js = f"document.documentElement.setAttribute('data-theme', '{theme}');"
                    widget._view.page().runJavaScript(js)
                    log.debug(f"[MainWindow] Queued theme '{theme}' update for memory manager")
                    return
        except Exception:
            pass  # Non-critical; dialog will pick up theme on next open

    def _apply_syntax_highlighting_fonts(self):
        """Apply Dracula-themed fonts and colors globally to code displays."""
        if not HAS_SYNTAX_HIGHLIGHTING:
            return
        
        try:
            # Initialize global code colorizer
            self._code_colorizer = UniversalCodeColorizer()
            
            # Apply monospace font (for code editors, terminal)
            mono_fonts = FONTS['mono']
            
            # Apply sans-serif font (for UI elements)
            sans_fonts = FONTS['sans']
            
            # Apply to all code editors
            if hasattr(self, '_code_editor') and self._code_editor:
                code_font = QFont()
                code_font.setFamily(mono_fonts[0])  # Use first preferred monospace font
                code_font.setPointSize(10)
                code_font.setFixedPitch(True)
                self._code_editor.setFont(code_font)
                log.info(f"[FONTS] Applied monospace font to editor: {mono_fonts[0]}")
            
            # Apply to terminal if available
            if hasattr(self, '_terminal') and self._terminal:
                term_font = QFont()
                term_font.setFamily(mono_fonts[0])
                term_font.setPointSize(9)
                term_font.setFixedPitch(True)
                self._terminal.setFont(term_font)
                log.info(f"[FONTS] Applied terminal font: {mono_fonts[0]}")
            
            # Log Dracula theme application
            log.info(f"[COLORS] Dracula theme applied with {len(DRACULA_COLORS)} color definitions")
            log.info(f"[LANGUAGES] 100+ programming languages supported")
            log.info(f"[FRAMEWORKS] React, Vue, Angular, Django, Flask, FastAPI, and 20+ frameworks")
            log.info(f"[MARKDOWN] Blue headings (#0047AB), White text (#ffffff)")
            
            # Display font stack information
            mono_stack = " -> ".join(mono_fonts)
            sans_stack = " -> ".join(sans_fonts)
            log.info(f"[FONT_STACK_MONO] {mono_stack}")
            log.info(f"[FONT_STACK_SANS] {sans_stack}")
            
        except Exception as e:
            log.warning(f"Error applying syntax highlighting fonts: {e}")
    
    def colorize_code(self, code, language='plaintext'):
        """
        Public method to colorize code with Dracula theme.
        Every language gets colors - NO white text fallback.
        """
        if not HAS_SYNTAX_HIGHLIGHTING or not hasattr(self, '_code_colorizer'):
            return code
        
        try:
            return self._code_colorizer.colorize(code, language)
        except Exception as e:
            log.warning(f"Error colorizing code for language '{language}': {e}")
            return code
    
    def colorize_markdown(self, markdown_text):
        """
        Public method to colorize Markdown with blue headings and white text.
        """
        if not HAS_SYNTAX_HIGHLIGHTING:
            return markdown_text
        
        try:
            return MarkdownColorizer.colorize(markdown_text)
        except Exception as e:
            log.warning(f"Error colorizing markdown: {e}")
            return markdown_text

    # ------------------------------------------------------------------
    # Signal Connections
    # ------------------------------------------------------------------
    def _connect_signals(self):
        # Sidebar signals
        self._sidebar.file_opened.connect(self._open_file)
        self._sidebar.live_preview_requested.connect(self.open_live_preview_for_file)
        self._sidebar.open_folder_requested.connect(self._open_folder_dialog)
        self._sidebar.file_search_opened.connect(self._open_file_at_line)
        self._sidebar.ai_action_requested.connect(self._ai_action)
        self._sidebar.file_renamed.connect(self._on_sidebar_file_renamed)
        self._sidebar.file_renamed.connect(lambda old, new: self._webview_panel.rename_file(old, new))
        self._sidebar.file_deleted.connect(self._on_sidebar_file_deleted)
        # Hide startup overlay as soon as sidebar.html finishes loading
        # (much faster than waiting for user interaction or 60s safety timer)
        if hasattr(self._sidebar, 'page_loaded'):
            self._sidebar.page_loaded.connect(self._hide_startup_overlay)
        # DIRECT backup connection — bypasses SidebarWidget forwarding layer
        if hasattr(self._sidebar, '_bridge') and self._sidebar._bridge:
            self._sidebar._bridge.file_deleted.connect(self._on_sidebar_file_deleted)
            _log.info("Direct sidebar._bridge.file_deleted connection established")
        
        # Changed files panel signals (guarded — panel may not be available in compiled builds)
        if self._changed_files_panel is not None:
            try:
                self._changed_files_panel.file_accepted.connect(self._on_accept_file_edit)
                self._changed_files_panel.file_rejected.connect(self._on_reject_file_edit)
                self._changed_files_panel.accept_all_requested.connect(self._on_accept_all_files)
                self._changed_files_panel.reject_all_requested.connect(self._on_reject_all_files)
                self._changed_files_panel.file_opened.connect(self._open_file)
            except Exception as e:
                _log.warning(f"ChangedFilesPanel signals failed: {e}")

        # Sidebar footer gear button → Memory Manager
        self._sidebar.settings_requested.connect(self._show_memory_manager)

        # Chat history panel signals (inside sidebar, index 4)
        self._sidebar.chat_selected.connect(self._on_chat_selected)
        self._sidebar.chat_renamed.connect(self._on_chat_renamed)
        self._sidebar.chat_delete_requested.connect(self._on_chat_delete_requested)
        self._sidebar.new_chat_requested.connect(self._on_new_chat_requested)

        # AI context sync: when user switches chats, restore conversation context for AI
        if not getattr(self, '_is_native_chat', False):
            self._ai_chat.switch_chat_context.connect(self._on_switch_chat_context)

        # Project manager
        self._project_manager.project_opened.connect(self._on_project_opened)
        self._project_manager.project_closed.connect(self._on_project_closed)


        # Webview editor active file changes
        self._webview_panel.active_file_changed.connect(self._on_webview_file_changed)
        self._webview_panel.file_content_changed.connect(self._on_webview_content_changed)
        self._webview_panel.file_save_requested.connect(self._on_webview_save_requested)
        self._webview_panel.file_closed.connect(self._on_webview_file_closed)
        self._webview_panel.cursor_position_changed.connect(self._update_status_cursor)
        # Legacy editor tab changes (welcome/PDF/image)
        self._editor_tabs.currentChanged.connect(self._on_tab_changed)
        
        # File manager undo/redo signals
        if hasattr(self, '_file_manager'):
            self._file_manager.file_deleted.connect(self._on_file_deleted_for_undo)
            self._file_manager.file_restored.connect(self._on_file_restored_for_redo)

        # AI chat - ONLY connect signals here to avoid duplicates
        # Native chat backend has its own signal wiring (AgentSignals); skip webview connections
        if not getattr(self, '_is_native_chat', False):
            self._ai_chat.message_sent.connect(self._on_ai_chat_message)
            self._ai_chat.vision_history_sync.connect(self._ai_agent.inject_vision_history)
            self._ai_agent.response_chunk.connect(self._ai_chat.on_chunk)
            self._ai_agent.response_complete.connect(self._ai_chat.on_complete)
            self._ai_agent.response_complete.connect(self._on_ai_task_complete)
            # Chat persistence: save timeline after each AI response
            if hasattr(self, '_chat_save_callback') and self._chat_save_callback:
                self._ai_agent.response_complete.connect(self._chat_save_callback)
                log.info("[ChatPersist] Connected save callback to response_complete")
            self._ai_agent.request_error.connect(self._ai_chat.on_error)
            self._ai_agent.file_generated.connect(self._on_agent_file_generated)
            # TEMPORARILY DISABLED - file_tracker was part of deleted agentic code
            # self._ai_agent.file_edited_diff.connect(self._file_tracker.add_edit)
            self._ai_agent.file_edited_diff.connect(self._on_file_edited_diff_for_js)
            self._ai_agent.file_edited_diff.connect(self._on_inline_edit_diff)
            # NOTE: on_file_edited_diff + show_diff_card REMOVED — NativeChatBridge
            # already handles DiffCard creation via file_edited_diff → _on_file_diff →
            # signals.file_diff → chat_panel.on_file_diff. The old direct connections
            # created a SECOND DiffCard for the same edit, causing duplication after
            # IDE restart (serialize captured both cards, restore showed edit twice).
            # File operation cards — animated create/edit cards
            self._ai_agent.file_creating_started.connect(self._on_file_creating_started)
            self._ai_agent.file_editing_started.connect(self._on_file_editing_started)
            self._ai_agent.file_operation_completed.connect(self._on_file_operation_completed)
            self._ai_agent.tool_activity.connect(self._ai_chat.show_tool_activity)
            self._ai_agent.plan_created.connect(self._ai_chat._on_plan_created)
            self._ai_chat.build_plan_requested.connect(self._ai_agent.handle_build_plan)
            self._ai_agent.directory_contents.connect(self._ai_chat.show_directory_contents)
            self._ai_agent.directory_contents.connect(self._on_directory_contents_for_tree)
            self._ai_agent.thinking_started.connect(self._ai_chat.show_thinking)
            self._ai_agent.thinking_stopped.connect(self._ai_chat.hide_thinking)
            self._ai_agent.todos_updated.connect(self._ai_chat.update_todos)
            self._ai_agent.task_progress_update.connect(self._on_task_progress_update)
            self._ai_agent.tool_summary_ready.connect(self._ai_chat.show_tool_summary)
            # Recovery: context compaction status + turn-limit continuation
            self._ai_agent.agent_status_update.connect(self._ai_chat.on_agent_status_update)
            self._ai_agent.turn_limit_hit.connect(self._ai_chat.on_turn_limit_hit)
            # Token budget: real-time context usage bar
            self._ai_agent.context_budget_update.connect(self._ai_chat.on_context_budget_update)
            # Permission gate: agent → chat UI shows card; user response → agent continues
            self._ai_agent.permission_requested.connect(self._ai_chat._on_permission_request)
            self._ai_chat.permission_decided.connect(self._ai_agent.on_permission_respond)
            # Auto-approval toggle in chat toolbar → backend permission behavior
            self._ai_chat.always_allow_changed.connect(self._ai_agent.set_always_allowed)
            # Refresh editor when user accepts an AI edit in chat panel
            self._ai_chat.edit_accepted.connect(self._on_accept_file_edit)
        
        # Model selection connection for BOTH native and webview chat modes.
        # Native chat emits a single string (model_id); webview emits 3 args.
        # The slot handles both via default parameters for perf/cost.
        self._ai_chat.model_changed.connect(self._on_model_changed)

        # Mode selection (Agent/Ask/Plan) — controls tool access
        if hasattr(self._ai_chat, 'input_area') and hasattr(self._ai_chat.input_area, 'mode_changed'):
            self._ai_chat.input_area.mode_changed.connect(self._on_mode_changed)
            log.info("[MAIN] mode_changed signal connected")

        # Active signal connections (webview-only)
        if not getattr(self, '_is_native_chat', False):
            self._ai_chat.generate_plan_requested.connect(self._on_generate_plan)
            self._ai_chat.open_file_requested.connect(self._open_file)
            self._ai_chat.open_file_at_line_requested.connect(self._open_file_at_line)
            log.info(f"[Diff-Debug] Connecting show_diff_requested to _on_show_diff. Signal exists: {hasattr(self._ai_chat, 'show_diff_requested')}")
            self._ai_chat.show_diff_requested.connect(self._on_show_diff)
            self._ai_chat.answer_question_requested.connect(self._ai_agent.user_responded)
            self._ai_chat.smart_paste_check_requested.connect(self._on_smart_paste_check)
            
            # Todo toggle (logs only - state managed by bridge/JS)
            self._ai_chat.toggle_todo_requested.connect(self._on_toggle_todo)
        
        # Interactive questions from agent to user
        # user_question_requested is handled by native_chat_bridge → signals.question → chat_panel.on_question
        # Do NOT also connect to _on_ai_question_requested — that creates duplicate question cards.

        # ========== CortexDiffBridge: wire accept/reject signals ==========
        # This connects useDiffInIDE.py's CortexDiffBridge to Cortex's FEC card
        # accept/reject signals so the agent can await user confirmation of edits.
        if not getattr(self, '_is_native_chat', False):
            try:
                import importlib as _il
                try:
                    _diff_ide_mod = _il.import_module("src.agent.src.hooks.useDiffInIDE")
                except ImportError:
                    _diff_ide_mod = _il.import_module("agent.src.hooks.useDiffInIDE")
                _cdb = _diff_ide_mod.CortexDiffBridge.instance()
                _cdb.register_accept_signal(self._ai_chat.accept_file_edit_requested)
                _cdb.register_reject_signal(self._ai_chat.reject_file_edit_requested)
                log.info("[CortexDiffBridge] Accept/Reject signals wired")
            except Exception as _cdb_err:
                log.warning(f"[CortexDiffBridge] Signal wiring skipped: {_cdb_err}")

        # ── Native chat backend: wire agent_bridge → ChatPanel via adapter ──
        if getattr(self, '_is_native_chat', False):
            try:
                from src.ui.native_chat_bridge import NativeChatBridge
                self._native_bridge = NativeChatBridge(self._ai_agent, parent=self)
                self._ai_chat.bind(self._native_bridge.signals, bridge=self._ai_agent)
                # Refresh editor when user accepts an AI edit in chat panel
                self._ai_chat.edit_accepted.connect(self._on_accept_file_edit)
                # Open new files in editor when AI creates them
                self._ai_agent.file_generated.connect(self._on_agent_file_generated)
                # Auto-open and refresh editor when AI edits a file (auto-apply mode)
                self._ai_agent.file_edited_diff.connect(self._on_agent_file_edited)
                # Wire header "+ New Chat" button → summarize to MEMORY.md + fresh chat
                self._ai_chat.new_chat_requested.connect(self._on_new_chat_requested)
                # Wire user input: ChatPanel.send_requested → agent
                self._ai_chat.input_area.send_requested.connect(self._on_ai_chat_message)
                # Wire stop: ChatPanel.stop_requested → agent cancel
                self._ai_chat.input_area.stop_requested.connect(self._on_ai_stop_requested)
                # Wire image paste → agent vision
                self._ai_chat.input_area._image_callback = self._on_image_pasted
                # Wire project root for file matching on paste
                if hasattr(self, '_project_manager') and self._project_manager.root:
                    self._ai_chat.input_area._project_root = self._project_manager.root
                # Auto-apply deferred edits when AI response completes (native chat path)
                self._ai_agent.response_complete.connect(self._on_ai_task_complete)
                # Register rendering-done callback: notification fires AFTER chat renders
                self._ai_chat._rendering_done_cb = self._on_chat_rendering_done
                # Chat persistence: save timeline after each AI response (native chat path)
                if hasattr(self, '_chat_save_callback') and self._chat_save_callback:
                    self._ai_agent.response_complete.connect(self._chat_save_callback)
                    log.info("[NATIVE-CHAT][ChatPersist] Connected save callback to response_complete")
                log.info("[NATIVE-CHAT] AgentBridge → ChatPanel wired via NativeChatBridge")
            except Exception as _ncb_err:
                log.warning(f"[NATIVE-CHAT] Bridge wiring failed: {_ncb_err}")

        # Close editor tabs for files the agent moved to the Recycle Bin
        # (works for both native and webview chat backends).
        if hasattr(self._ai_agent, 'files_deleted_by_agent'):
            try:
                self._ai_agent.files_deleted_by_agent.connect(self._on_agent_files_deleted)
            except Exception as _fd_err:
                log.warning(f"[AGENT] files_deleted_by_agent wiring failed: {_fd_err}")

        # LivePreview agent tool — open/read/console/close on the UI thread
        if hasattr(self._ai_agent, 'live_preview_tool_request'):
            try:
                self._ai_agent.live_preview_tool_request.connect(self._on_agent_live_preview)
            except Exception as _lp_err:
                log.warning(f"[AGENT] live_preview_tool_request wiring failed: {_lp_err}")

        # Terminal tab changes
        self._terminal_tabs.currentChanged.connect(self._on_terminal_tab_changed)

        # ========== PERMISSION SYSTEM CONNECTION (NEW) ==========
        # Connect AI agent to UI for permission dialogs
        self._ai_agent.set_ui_parent(self)
        log.info("Permission system initialized and connected to main window")

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------
    def _new_file(self):
        """Create a new untitled file in the webview editor."""
        is_dark = self._theme_manager.is_dark
        self._webview_panel.open_file("untitled.py", "", "python")
        self._webview_panel.set_theme(is_dark)

    def _open_file_dialog(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open File",
                                               str(self._project_manager.root or Path.home()))
        if path:
            self._open_file(path)
    
    def _find_file_in_project(self, filename: str) -> str | None:
        """Search for a file by name in the project directory (recursive)."""
        if not self._project_manager.root:
            return None
        
        # Extract just the filename (remove any directory components)
        from pathlib import Path as PathLib
        clean_filename = PathLib(filename).name
        
        root = Path(self._project_manager.root)
        
        # Search recursively for the file
        try:
            for file_path in root.rglob(clean_filename):
                if file_path.is_file():
                    return str(file_path)
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')
        
        return None
    
    def _open_folder_dialog(self):
        """Open a folder as the active project — fully resets IDE state.
            
        Closes all editor tabs, clears AI chat, resets git/diff state,
        then opens the selected folder with a clean sidebar and fresh context.
        """
        folder = QFileDialog.getExistingDirectory(
            self, "Open Folder", str(Path.home())
        )
        if not folder:
            return
    
        log.info(f"[OpenFolder] Switching to project: {folder}")
    
        # ── 1. Close ALL open editor files (both webview AND legacy tabs) ──
        if hasattr(self, '_webview_panel'):
            self._webview_panel.close_all_files()
        if hasattr(self, '_editor_tabs'):
            self._editor_tabs.close_all_tabs()
    
        # ── 2. Clear AI chat completely (fresh conversation for new project) ──
        if not getattr(self, '_is_native_chat', False) and hasattr(self, '_ai_chat') and self._ai_chat:
            try:
                self._ai_chat.clear_chat()
                log.info("[OpenFolder] AI chat cleared")
            except Exception as e:
                log.warning(f"[OpenFolder] Chat clear failed: {e}")
    
        # ── 3. Reset Git state (clear old repo before switching) ──
        if hasattr(self, '_git_manager'):
            self._git_manager._repo_path = None
            log.info("[OpenFolder] Git repo path cleared")
    
        # ── 4. Clear review/summary panels to remove stale diff data ──
        if hasattr(self, '_review_panel'):
            try:
                from PyQt6.QtWebEngineWidgets import QWebEngineView
                for child in self._review_panel.findChildren(QWebEngineView):
                    child.setHtml("<html><body style='background:#1e1e1e'></body></html>")
                log.info("[OpenFolder] Review panel cleared")
            except Exception as e:
                log.debug(f"[OpenFolder] Review clear failed: {e}")
    
        # ── 5. Open the new project folder (sidebar, git reset) ──
        self._open_folder_programmatic(folder)
    
        # ── 6. Update AI chat with new project context ──
        if hasattr(self, '_ai_chat') and self._ai_chat:
            try:
                project_name = Path(folder).name
                self._ai_chat.set_project_info(project_name, folder, "")
                log.info(f"[OpenFolder] Chat project info set: {project_name}")
            except Exception as e:
                log.warning(f"[OpenFolder] Chat project info failed: {e}")
    
        log.info(f"[OpenFolder] \u2713 Project reset complete \u2014 ready at: {folder}")
    
    def _open_folder_programmatic(self, folder: str):
        """Open a folder as the active project (no dialog, usable from argv/drag-drop)."""
        self._project_manager.open(folder)

        # Populate left sidebar file tree
        self._sidebar.set_project(folder)

        # Initialize Git repository
        if hasattr(self, '_git_manager'):
            self._git_manager.set_repository(folder)
            log.info(f"[GIT] Repository set to: {folder}")

    def _new_project(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder for New Project",
                                                   str(Path.home()))
        if folder:
            # For now, just open it as a project. 
            # Could add a template/scaffolding step here later.
            self._project_manager.open(folder)

            # Populate left sidebar file tree
            self._sidebar.set_project(folder)

            # Initialize Git repository
            if hasattr(self, '_git_manager'):
                self._git_manager.set_repository(folder)
                log.info(f"[GIT] Repository set to: {folder}")

    def _on_agent_file_edited(self, filepath: str, _original: str = "", _new: str = ""):
        """Slot for file_edited_diff in auto-apply mode — reload editor with fresh disk content.
        
        ONLY refreshes files that are ALREADY open in the editor.
        Does NOT create new tabs for files the user didn't open — the AI
        writes to disk silently and the user sees changes when they click
        the file in the sidebar.
        """
        try:
            filepath = os.path.normpath(filepath)
            
            # Auto-commit if enabled
            if hasattr(self, '_git_manager') and self._git_manager:
                try:
                    self._git_manager.auto_commit_file(filepath, "modified")
                except Exception:
                    pass  # Non-critical
            # Invalidate file_manager cache so stale content is not returned
            try:
                resolved = str(Path(filepath).resolve())
                if hasattr(self._file_manager, '_file_cache'):
                    with self._file_manager._file_cache._lock:
                        self._file_manager._file_cache.cache.pop(resolved, None)
                if hasattr(self._file_manager, '_hash_cache'):
                    self._file_manager._hash_cache.pop(resolved, None)
            except Exception:
                pass
            
            # Only touch the editor if the file is ALREADY open.
            # Otherwise, disk content is already updated — user sees it
            # when they click the file in the sidebar.
            _is_open = filepath in self._webview_panel._open_files
            if not _is_open:
                log.info(f"[AUTO-EDIT] File not open in editor — disk updated silently: {filepath}")
                return
            
            # Use _new content directly from signal (already written to disk — avoids cache)
            if _new:
                language = detect_language(filepath)
                self._webview_panel.force_reload_file(filepath, _new, language)
                log.info(f"[AUTO-EDIT] Reloaded editor: {filepath} ({len(_new)} chars)")
            else:
                self._open_file(filepath, priority=False, reload=True)
                # Only switch to the file if the user is already looking at it.
                _active = self._webview_panel.get_active_file()
                if (hasattr(self, '_webview_panel') and self._webview_panel
                        and filepath in self._webview_panel._open_files
                        and filepath == _active):
                    self._webview_panel.switch_to_file(filepath)
        except Exception as e:
            log.warning(f"[AUTO-EDIT] Failed to reload {filepath}: {e}")

    def _on_agent_file_generated(self, filepath: str, _content: str = ""):
        """Slot for agent file_generated signal — routes through throttled pump (priority=False).

        Agent Write tools can create many files in a single turn. Opening them all with
        priority=True would flood QWebChannel with simultaneous model.setValue() IPC
        calls, crashing Chromium's render process on Windows 25H2.

        Using priority=False routes them through the warmup throttle
        (10s spacing during first 60s, 1.5s after), preventing the crash.

        IMPORTANT: reload=True ensures that when the AI modifies an already-open file,
        the Monaco editor gets fresh content from disk instead of showing stale cached text.

        ONLY refreshes files that are ALREADY open in the editor.
        Does NOT create new tabs for files the user didn't open.
        """
        try:
            filepath = os.path.normpath(filepath)
            
            # Auto-commit if enabled
            if hasattr(self, '_git_manager') and self._git_manager:
                try:
                    self._git_manager.auto_commit_file(filepath, "created")
                except Exception:
                    pass  # Non-critical
        except Exception:
            pass
        
        # Only touch the editor if the file is ALREADY open.
        _is_open = filepath in self._webview_panel._open_files if hasattr(self, '_webview_panel') else False
        if not _is_open:
            log.info(f"[AUTO-GENERATE] File not open in editor — disk updated silently: {filepath}")
            return
        
        self._open_file(filepath, priority=False, reload=True)
        # Only flush + switch if user is already viewing this file
        _active = self._webview_panel.get_active_file() if hasattr(self, '_webview_panel') else None
        if (hasattr(self, '_webview_panel') and self._webview_panel
                and filepath in self._webview_panel._open_files
                and filepath == _active):
            self._webview_panel.switch_to_file(filepath)

    def _open_file(self, filepath: str, *, priority: bool = True, reload: bool = False):
        # Normalize path (convert forward slashes to backslashes on Windows)
        filepath = os.path.normpath(filepath)
        path = Path(filepath)
        
        # If already tracked by webview panel, handle reload vs switch.
        # switch_to_file() avoids QWebChannel model.setValue() IPC which
        # crashes Chromium on Windows 25H2 during the startup warmup phase.
        # BUT on accept/reject of AI edits, we MUST reload the content so
        # the editor reflects the latest disk state (like VS Code/Cursor).
        if reload:
            # Always re-read from disk on reload — the file may have been
            # closed from JS between the edit and the Accept click.
            try:
                content = self._file_manager.read(filepath, lazy_load=False, use_cache=False)
                if content is not None:
                    # Check for another pending deferred edit after accept/reject
                    _deferred = None
                    if hasattr(self, '_ai_agent') and self._ai_agent:
                        try:
                            _deferred = self._ai_agent.get_deferred_edit(filepath)
                        except Exception:
                            pass
                    if _deferred is not None:
                        log.info(f"[Deferred] Reload showing deferred edit: {filepath} "
                                 f"({len(_deferred)} chars)")
                        content = _deferred
                    language = detect_language(filepath)
                    self._webview_panel.open_file(filepath, content, language, priority=priority)
                    log.info(f"File reloaded in editor after accept/reject: {filepath} ({len(content)} chars)")
                else:
                    log.error(f"Failed to read content on reload: {filepath}")
            except Exception as e:
                log.error(f"Failed to reload file content: {e}")
            self._update_status_file(filepath)
            self._notify_search_of_open_files()
            return
        if hasattr(self, '_webview_panel') and filepath in self._webview_panel._open_files:
            # Check if there's a pending deferred edit — if so, force reload with deferred content
            _deferred = None
            if hasattr(self, '_ai_agent') and self._ai_agent:
                try:
                    _deferred = self._ai_agent.get_deferred_edit(filepath)
                except Exception:
                    pass
            if _deferred is not None:
                log.info(f"File already open but has deferred edit — forcing reload: {filepath} "
                         f"({len(_deferred)} chars)")
                language = detect_language(filepath)
                self._webview_panel.open_file(filepath, _deferred, language, priority=priority)
            else:
                # Always refresh content from disk when user explicitly opens a file
                # This fixes stale cache issues where editor shows wrong content
                try:
                    content = self._file_manager.read(filepath, lazy_load=False, use_cache=False)
                    if content is not None:
                        language = detect_language(filepath)
                        self._webview_panel.open_file(filepath, content, language, priority=priority)
                        log.info(f"File refreshed from disk: {filepath} ({len(content)} chars)")
                    else:
                        log.info(f"File already open in webview, switching: {filepath}")
                        self._webview_panel.switch_to_file(filepath)
                except Exception as e:
                    log.warning(f"Failed to refresh file content: {e}")
                    self._webview_panel.switch_to_file(filepath)
            self._update_status_file(filepath)
            self._notify_search_of_open_files()
            return

        # ── Startup warmup guard: queue BULK opens only ──────────────────
        # Bulk session-restore opens (priority=False) crash Chromium on
        # Windows 25H2 if sent through QWebChannel during startup. Single
        # user-initiated clicks (priority=True) are safe and pass through.
        elapsed = time.time() - self._start_time
        if elapsed < self._warmup_duration and not priority:
            # Deduplicate
            for fp, _ in self._warmup_queued_files:
                if fp == filepath:
                    return
            self._warmup_queued_files.append((filepath, priority))
            count = len(self._warmup_queued_files)
            remaining = self._warmup_duration - int(elapsed)
            log.info(f"[Warmup {elapsed:.0f}s] Queued: {os.path.basename(filepath)} ({count} total)")
            if hasattr(self, '_status_file') and self._status_file:
                self._status_file.setText(f"  \u23f3 Editor warming up ({remaining}s left, {count} file(s) queued)")
            return

        log.info(f"Opening file: {filepath}")

        # If file doesn't exist, try to find it in the project
        if not path.exists() or not path.is_file():
            # Try searching in project directory
            found_path = self._find_file_in_project(path.name)
            if found_path:
                log.info(f"Found file in project: {found_path}")
                path = Path(found_path)
                filepath = str(found_path)
            else:
                log.warning(f"File skip (not found or dir): {filepath}")
                return
        
        # Check file extension for images and documents
        file_ext = path.suffix.lower()
        
        # Handle image files
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp', '.ico', '.tiff', '.tif'}
        if file_ext in image_extensions:
            self._open_image_file(filepath)
            return
        
        # Handle PDF files
        if file_ext == '.pdf':
            self._open_pdf_file(filepath)
            return
        
        # Handle Office documents
        office_extensions = {'.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx'}
        if file_ext in office_extensions:
            self._open_office_file(filepath)
            return
            
        # Initialize file snapshots dict for diff generation
        if not hasattr(self, '_file_snapshots'):
            self._file_snapshots = {}
        
        if self._file_manager.is_binary(filepath):
            log.info(f"File skip (binary): {filepath}")
            QMessageBox.information(self, "Binary File",
                                    f"'{path.name}' is a binary file and cannot be edited.")
            return
            
        try:
            # Force full read (lazy_load=False) — the editor needs the complete
            # file content, not a 100-line viewport. Large files use openFileFromUri
            # in webview_panel._open_file_js which XHR-loads the full file from disk.
            content = self._file_manager.read(filepath, lazy_load=False, use_cache=False)
            if content is None:
                log.error(f"Failed to read content: {filepath}")
                return
            
            # ── Deferred edit check ──
            # If AI has made edits to this file that are pending user acceptance,
            # show the deferred content (not disk) so the user sees the latest AI state.
            _deferred_content = None
            if hasattr(self, '_ai_agent') and self._ai_agent:
                try:
                    _deferred_content = self._ai_agent.get_deferred_edit(filepath)
                except Exception:
                    pass
            if _deferred_content is not None:
                log.info(f"[Deferred] Showing deferred edit for: {filepath} "
                         f"(disk={len(content)} chars, deferred={len(_deferred_content)} chars)")
                content = _deferred_content
            
            # Store original snapshot for diff generation (only if not already stored)
            # This preserves the FIRST version opened, allowing multiple diffs
            if filepath not in self._file_snapshots:
                self._file_snapshots[filepath] = content
                log.info(f"[Snapshot] Stored initial snapshot: {filepath} ({len(content)} chars)")
            else:
                log.info(f"[Snapshot] Keeping existing snapshot for: {filepath}")
                
            log.info(f"Content read ({len(content)} chars). Detecting language...")
            language = detect_language(filepath)
            log.info(f"Language detected: {language}. Opening index in tabs...")
            
            # Get current theme state and pass it to the editor
            is_dark = self._theme_manager.is_dark
            self._webview_panel.open_file(filepath, content, language, priority=priority)
            self._webview_panel.set_theme(is_dark)

            self._update_status_file(filepath)
            log.info(f"File opened successfully: {filepath}")

            # SAFETY NET: Force renderTabs() after a short delay to guarantee
            # the tab header is visible. _safe_run_js defers JS execution via
            # QTimer.singleShot(0), so openFile() may not have run yet when we
            # return here. This deferred renderTabs() ensures the tab appears
            # even if the primary JS call was delayed.
            QTimer.singleShot(100, lambda: self._webview_panel._safe_run_js("renderTabs();"))

            # Notify sidebar search of currently open files
            self._notify_search_of_open_files()

        except Exception as e:
            log.error(f"Error opening file {filepath}: {e}", exc_info=True)

    def _notify_search_of_open_files(self):
        """Gather all currently open editor files and push to sidebar search panel."""
        open_paths = []
        # Collect from webview panel (primary editor)
        if hasattr(self, '_webview_panel'):
            open_paths = list(self._webview_panel._open_files.keys())
        # Fallback: collect from legacy tab widget
        if not open_paths and hasattr(self, '_editor_tabs'):
            open_paths = self._editor_tabs.get_open_files()
        if hasattr(self, '_sidebar'):
            self._sidebar.set_opened_files(open_paths)

    def _flush_warmup_queue(self):
        """Open all files queued during the Chromium startup warmup period.

        Chromium WebEngine on Windows 25H2 crashes its render process if
        ANY QWebChannel IPC (model.setValue()) reaches Monaco during the
        first ~60s after launch. This method fires after the warmup timer
        expires, opening all queued files through the normal pipeline.
        """
        elapsed = time.time() - self._start_time
        count = len(self._warmup_queued_files)
        
        # ── Skip warmup file opens if AI is actively processing ──
        # Opening files mid-turn shows stale disk content and confuses
        # the user — they didn't click Accept, the file just appears.
        _ai_active = False
        if hasattr(self, '_ai_agent') and self._ai_agent:
            try:
                _ai_active = getattr(self._ai_agent, '_agentic_turn_active', False)
            except Exception:
                pass
        if _ai_active:
            log.info(
                f"[Warmup] Suppressed flush — AI is actively processing "
                f"({count} queued file(s) skipped)"
            )
            self._warmup_queued_files.clear()
            return
        
        log.info(f"[Warmup] Flush after {elapsed:.1f}s — opening {count} queued file(s)")
        for filepath, priority in self._warmup_queued_files:
            if Path(filepath).exists():
                # Re-enter _open_file — warmup guard is now bypassed (elapsed >= duration)
                self._open_file(filepath, priority=priority)
        self._warmup_queued_files.clear()
        if hasattr(self, '_status_file') and self._status_file:
            self._status_file.setText("  No file open")
        log.info("[Warmup] Queue flushed — Chromium should now be stable")

    def _open_image_file(self, filepath: str):
        """Open an image file as a visual preview in the webview editor panel."""
        import base64
        import mimetypes
        import json as _json

        try:
            log.info(f"Opening image file: {filepath}")
            path = Path(filepath)
            file_size = path.stat().st_size
            if file_size < 1024:
                size_str = f"{file_size} B"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / (1024 * 1024):.1f} MB"

            mime_type = mimetypes.guess_type(filepath)[0] or "image/png"
            with open(filepath, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")

            data_uri = f"data:{mime_type};base64,{b64}"
            safe_path = _json.dumps(filepath)
            safe_uri = _json.dumps(data_uri)
            safe_size = _json.dumps(size_str)
            self._webview_panel._safe_run_js(
                f"showImagePreview({safe_path}, {safe_uri}, {safe_size});"
            )
            self._update_status_file(filepath)
            log.info(f"Image opened in viewer: {filepath} ({size_str})")

        except Exception as e:
            log.error(f"Error opening image file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open image: {e}")

    def _open_pdf_file(self, filepath: str):
        """Open a PDF file by rendering pages as images."""
        from PyQt6.QtWidgets import QLabel, QScrollArea, QVBoxLayout, QWidget
        from PyQt6.QtCore import Qt
        from PyQt6.QtGui import QPixmap, QImage
        import fitz  # PyMuPDF
        
        try:
            log.info(f"Opening PDF file: {filepath}")
            path = Path(filepath)
            
            # Open PDF with PyMuPDF
            doc = fitz.open(filepath)
            
            # Store page count immediately
            total_pages = doc.page_count
            
            if total_pages == 0:
                QMessageBox.warning(self, "Error", "PDF has no pages")
                return
            
            # Create scrollable container for all pages
            container = QWidget()
            layout = QVBoxLayout(container)
            layout.setSpacing(10)
            layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Get available width
            tab_width = self._editor_tabs.width() - 60
            
            # Render each page as an image
            for page_num in range(min(total_pages, 50)):  # Limit to 50 pages
                page = doc[page_num]
                
                # Calculate zoom to fit width
                zoom = tab_width / page.rect.width
                mat = fitz.Matrix(zoom, zoom)
                
                # Render page to pixmap
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to QImage
                img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888)
                pixmap = QPixmap.fromImage(img)
                
                # Create label for page
                label = QLabel()
                label.setPixmap(pixmap)
                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                label.setStyleSheet("background-color: #1e1e1e; border: 1px solid #3d3d3d;")
                
                layout.addWidget(label)
            
            doc.close()
            
            # Create scroll area
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setWidget(container)
            scroll.setStyleSheet("background-color: #f0f0f0;")
            
            # Add to tabs
            idx = self._editor_tabs.addTab(scroll, path.name)
            self._editor_tabs.setTabToolTip(idx, f"{filepath} ({total_pages} pages)")
            self._editor_tabs.setCurrentIndex(idx)
            
            self._update_status_file(filepath)
            log.info(f"PDF file opened successfully: {filepath} ({total_pages} pages)")
            
        except ImportError:
            log.error("PyMuPDF (fitz) not installed")
            QMessageBox.warning(self, "Error", "PyMuPDF not installed. Run: pip install PyMuPDF")
        except Exception as e:
            log.error(f"Error opening PDF file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open PDF: {e}")

    def _open_office_file(self, filepath: str):
        """Open Office documents (Word, Excel, PowerPoint) inside the IDE as formatted text."""
        from PyQt6.QtWidgets import QTextEdit
        from PyQt6.QtCore import Qt
        
        try:
            log.info(f"Opening Office file: {filepath}")
            path = Path(filepath)
            file_ext = path.suffix.lower()
            
            # Create text viewer
            text_viewer = QTextEdit()
            text_viewer.setReadOnly(True)
            text_viewer.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            
            content_html = ""
            
            if file_ext in ['.docx']:
                # Read Word document
                try:
                    from docx import Document
                    doc = Document(filepath)
                    
                    content_html = f"<h2>{path.name}</h2><hr>"
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            # Check if it's a heading based on style
                            if para.style.name.startswith('Heading'):
                                content_html += f"<h3>{para.text}</h3>"
                            else:
                                content_html += f"<p>{para.text}</p>"
                    
                    # Add tables
                    for table in doc.tables:
                        content_html += "<table border='1' cellpadding='5'>"
                        for row in table.rows:
                            content_html += "<tr>"
                            for cell in row.cells:
                                content_html += f"<td>{cell.text}</td>"
                            content_html += "</tr>"
                        content_html += "</table><br>"
                    
                except ImportError:
                    content_html = f"<p style='color:red'>Error: python-docx library not installed.<br>Install with: pip install python-docx</p>"
                    log.error("python-docx library not installed")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading document: {e}</p>"
                    log.error(f"Error reading docx: {e}")
            
            elif file_ext in ['.xlsx', '.xls']:
                # Read Excel document
                try:
                    if file_ext == '.xlsx':
                        import openpyxl
                        wb = openpyxl.load_workbook(filepath, data_only=True)
                        sheetnames = wb.sheetnames
                        
                        content_html = f"<h2>{path.name}</h2><hr>"
                        
                        for sheet_name in sheetnames:
                            sheet = wb[sheet_name]
                            content_html += f"<h3>Sheet: {sheet_name}</h3>"
                            content_html += "<table border='1' cellpadding='5' style='border-collapse:collapse'>"
                            
                            # Read first 100 rows max
                            row_count = 0
                            for row in sheet.iter_rows(max_row=100):
                                content_html += "<tr>"
                                for cell in row:
                                    value = cell.value if cell.value is not None else ""
                                    content_html += f"<td>{value}</td>"
                                content_html += "</tr>"
                                row_count += 1
                                if row_count >= 100:
                                    content_html += "<tr><td colspan='100'>... (showing first 100 rows)</td></tr>"
                                    break
                            
                            content_html += "</table><br>"
                    
                    elif file_ext == '.xls':
                        import xlrd
                        wb = xlrd.open_workbook(filepath)
                        
                        content_html = f"<h2>{path.name}</h2><hr>"
                        
                        for sheet_idx in range(wb.nsheets):
                            sheet = wb.sheet_by_index(sheet_idx)
                            content_html += f"<h3>Sheet: {sheet.name}</h3>"
                            content_html += "<table border='1' cellpadding='5' style='border-collapse:collapse'>"
                            
                            # Read first 100 rows max
                            for row_idx in range(min(sheet.nrows, 100)):
                                content_html += "<tr>"
                                for col_idx in range(sheet.ncols):
                                    value = sheet.cell_value(row_idx, col_idx)
                                    content_html += f"<td>{value}</td>"
                                content_html += "</tr>"
                            
                            if sheet.nrows > 100:
                                content_html += "<tr><td colspan='100'>... (showing first 100 rows)</td></tr>"
                            
                            content_html += "</table><br>"
                    
                except ImportError as e:
                    content_html = f"<p style='color:red'>Error: Required library not installed.<br>Install with: pip install openpyxl xlrd</p>"
                    log.error(f"Library not installed: {e}")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading spreadsheet: {e}</p>"
                    log.error(f"Error reading xlsx/xls: {e}")
            
            elif file_ext == '.doc':
                # Old Word format - try to extract text
                try:
                    # Try using textract if available
                    import textract
                    text = textract.process(filepath).decode('utf-8', errors='ignore')
                    content_html = f"<h2>{path.name}</h2><hr>"
                    content_html += f"<pre style='white-space:pre-wrap;font-family:Arial,sans-serif'>{text}</pre>"
                except ImportError:
                    content_html = f"""<p style='color:orange'>
                        <b>Old Word Format (.doc)</b><br><br>
                        This file uses the older .doc format which requires additional libraries.<br><br>
                        <b>Options:</b><br>
                        1. Convert to .docx format (open in Word and Save As .docx)<br>
                        2. Install textract: <code>pip install textract</code><br>
                        3. Open externally with Microsoft Word
                    </p>"""
                    log.warning(f"Old .doc format not supported without textract: {filepath}")
                except Exception as e:
                    content_html = f"<p style='color:red'>Error reading .doc file: {e}</p>"
                    log.error(f"Error reading doc: {e}")
            
            else:
                content_html = f"<p>File type '{file_ext}' is not supported for internal viewing.</p>"
            
            # Set content with styling
            text_viewer.setHtml(f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; padding: 20px; line-height: 1.6; }}
                    h2 {{ color: #333; border-bottom: 2px solid #0078d4; padding-bottom: 10px; }}
                    h3 {{ color: #555; margin-top: 20px; }}
                    table {{ margin: 10px 0; border: 1px solid #ddd; }}
                    td {{ padding: 8px; border: 1px solid #ddd; }}
                    p {{ margin: 10px 0; }}
                </style>
            </head>
            <body>
                {content_html}
            </body>
            </html>
            """)
            
            # Add to tabs
            idx = self._editor_tabs.addTab(text_viewer, path.name)
            self._editor_tabs.setTabToolTip(idx, filepath)
            self._editor_tabs.setCurrentIndex(idx)
            
            self._update_status_file(filepath)
            log.info(f"Office file opened internally: {filepath}")
            
        except Exception as e:
            log.error(f"Error opening Office file {filepath}: {e}", exc_info=True)
            QMessageBox.warning(self, "Error", f"Could not open document: {e}")


    def _diff_cache_path(self):
        try:
            from pathlib import Path
            return Path.home() / ".cortex" / "diff_cache.json"
        except Exception:
            return None

    def _load_diff_cache(self):
        if hasattr(self, '_diff_cache'):
            return self._diff_cache
        self._diff_cache = {}
        try:
            cache_path = self._diff_cache_path()
            if cache_path and cache_path.exists():
                import json
                self._diff_cache = json.loads(cache_path.read_text(encoding='utf-8')) or {}
        except Exception:
            self._diff_cache = {}
        return self._diff_cache

    def _save_diff_cache(self):
        try:
            cache_path = self._diff_cache_path()
            if not cache_path:
                return
            cache_path.parent.mkdir(parents=True, exist_ok=True)
            import json
            cache_path.write_text(json.dumps(self._diff_cache), encoding='utf-8')
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')

    def _on_show_diff(self, file_path: str):
        """Show diff in Qt dialog window — triggered by Diff button click in chat."""
        log.info(f"[Diff] _on_show_diff called with: {file_path}")
        original, modified = '', ''

        # 1. Try the Python diff data store (most reliable)
        import os
        import subprocess
        from pathlib import Path
        normalized_requested = os.path.normcase(os.path.normpath(file_path))
        
        if hasattr(self, '_diff_data_store'):
            # Direct check
            if file_path in self._diff_data_store:
                original, modified = self._diff_data_store[file_path]
                log.info(f"[Diff] Found in _diff_data_store (direct): {file_path}")
            else:
                # Iterative normalized check
                log.debug(f"[Diff] Checking normalized path for: {normalized_requested}")
                for k, v in self._diff_data_store.items():
                    if os.path.normcase(os.path.normpath(k)) == normalized_requested:
                        original, modified = v
                        log.info(f"[Diff] Found in _diff_data_store (normalized): {k}")
                        break

        # 2b. Fallback to persisted diff cache
        if not modified:
            cache = self._load_diff_cache()
            cached = None
            if cache:
                norm = os.path.normcase(os.path.normpath(file_path))
                cached = cache.get(file_path) or cache.get(norm)
                if not cached:
                    # try to find by normalized keys
                    for k, v in cache.items():
                        if os.path.normcase(os.path.normpath(k)) == norm:
                            cached = v
                            break
            if cached:
                original = cached.get('original', '')
                modified = cached.get('modified', '')
                log.info(f"[Diff] Found in diff cache: {file_path}")
        # 3. Fallback to file_tracker
        if not modified:
            edit_info = self._file_tracker.get_edit(file_path)
            if edit_info:
                original = edit_info.original_content if edit_info.edit_type != 'C' else ''
                modified = edit_info.new_content
                log.info(f"[Diff] Found in file_tracker: {file_path}")

        if not modified:
            # Fallback: try Git diff against HEAD if repository available
            try:
                project_root = getattr(self, '_project_manager', None).root if hasattr(self, '_project_manager') else None
                if project_root and os.path.isdir(os.path.join(project_root, '.git')) and os.path.exists(file_path):
                    rel_path = os.path.relpath(file_path, project_root)
                    # Try to load original from HEAD (tracked files)
                    git_show = subprocess.run(
                        ['git', '-C', project_root, 'show', f'HEAD:{rel_path}'],
                        capture_output=True, text=True,
                        encoding='utf-8', errors='replace',
                        timeout=3,
                        creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                    )
                    if git_show.returncode == 0:
                        original = git_show.stdout
                        modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                        log.info(f"[Diff] Loaded diff via git for: {file_path}")
                    else:
                        # If untracked, treat original as empty
                        git_ls = subprocess.run(
                            ['git', '-C', project_root, 'ls-files', '--error-unmatch', rel_path],
                            capture_output=True, text=True,
                            encoding='utf-8', errors='replace',
                            timeout=3,
                            creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                        )
                        if git_ls.returncode != 0:
                            original = ''
                            modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                            log.info(f"[Diff] Loaded diff for untracked file: {file_path}")
            except Exception as _ge:
                log.debug(f"[Diff] Git fallback failed: {_ge}")

        # Final fallback: Use file snapshot taken when opened
        if not modified and hasattr(self, '_file_snapshots') and file_path in self._file_snapshots:
            original = self._file_snapshots[file_path]
            try:
                modified = Path(file_path).read_text(encoding='utf-8', errors='replace')
                if original != modified:
                    log.info(f"[Diff] Using file snapshot for diff: {file_path}")
                else:
                    log.info(f"[Diff] File unchanged since opened: {file_path}")
                    original = ''
                    modified = ''
            except Exception as e:
                log.warning(f"[Diff] Failed to read current file content: {e}")
                original = ''
                modified = ''

        if not modified:
            log.warning(f"[Diff] No diff data found for {file_path}")
            log.debug(f"[Diff] _diff_data_store keys: {list(getattr(self, '_diff_data_store', {}).keys())}")
            if hasattr(self, '_ai_chat'):
                import os
                filename = os.path.basename(file_path)
                self._ai_chat.add_system_message(f"⚠️ No diff data available for {filename}. It was not edited in this session.")
            return

        log.info(f"[Diff] Opening diff tab for {file_path} (+{len(modified)} chars)")
        is_dark = self._theme_manager.is_dark
        self._editor_tabs.open_diff_tab(file_path, original, modified, is_dark)

    def _on_inline_edit_submitted(self, prompt: str, selection_text: str, line_range: tuple):
        editor = self._editor_tabs.currentWidget()
        if not isinstance(editor, CodeEditor):
            self._ai_chat.add_system_message("Open a file to use inline edit.")
            return

        file_path = self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
        if not file_path:
            self._ai_chat.add_system_message("Open a file to use inline edit.")
            return

        start_line, end_line = line_range
        if start_line == end_line:
            line_range_text = f"{start_line}"
        else:
            line_range_text = f"{start_line}-{end_line}"

        rel_path = file_path
        project_root = self._project_manager.root
        if project_root:
            try:
                rel_path = str(Path(file_path).resolve().relative_to(Path(project_root).resolve()))
            except Exception:
                rel_path = file_path

        self._inline_edit_context = {
            "file_path": os.path.normpath(file_path),
            "relative_path": rel_path,
            "line_range": line_range,
            "selection_text": selection_text,
            "editor": editor,
            "diff": None,
        }

        self._ai_chat.add_system_message(
            f"Inline edit: `{rel_path}` lines {line_range_text}"
        )
        self._ai_chat._add_ai_bubble_streaming()

        inline_prompt = (
            "Inline edit request.\n"
            f"File (project-relative): {rel_path}\n"
            f"Selection lines: {line_range_text}\n"
            "Selected code:\n"
            "```\n"
            f"{selection_text}\n"
            "```\n"
            "Instruction:\n"
            f"{prompt}\n\n"
            "Constraints:\n"
            "- Apply changes only within the selection unless absolutely required.\n"
            "- Use file editing tools (prefer replace_lines) to update the file.\n"
            "- Do not modify other files.\n"
            "- Use the project-relative path above.\n"
        )
        self._ai_agent.chat(inline_prompt)

    def _on_inline_edit_cancelled(self):
        self._inline_edit_context = None

    def _on_inline_diff_requested(self):
        context = self._inline_edit_context or {}
        file_path = context.get("file_path")
        if not file_path:
            return

        diff_pair = context.get("diff")
        if diff_pair:
            original, modified = diff_pair
            is_dark = self._theme_manager.is_dark
            self._editor_tabs.open_diff_tab(file_path, original, modified, is_dark)
            return

        self._on_show_diff(file_path)

    def _on_inline_edit_diff(self, file_path: str, original: str, new_content: str):
        context = self._inline_edit_context
        if not context:
            return

        target = context.get("file_path")
        if not target:
            return

        if os.path.normcase(os.path.normpath(file_path)) != os.path.normcase(os.path.normpath(target)):
            return

        editor = context.get("editor")
        if not isinstance(editor, CodeEditor):
            return

        import difflib

        diff_lines = list(difflib.unified_diff(
            original.splitlines(),
            new_content.splitlines(),
            fromfile="Original",
            tofile="Modified",
            lineterm=""
        ))
        if not diff_lines:
            diff_text = "No changes detected."
        else:
            max_lines = 200
            if len(diff_lines) > max_lines:
                diff_lines = diff_lines[:max_lines]
                diff_lines.append("... (diff truncated)")
            diff_text = "\n".join(diff_lines)

        context["diff"] = (original, new_content)
        editor.show_inline_diff(diff_text)

    def _on_file_edited_diff_for_js(self, file_path: str, original: str, new_content: str):
        """Store diff data in Python dict for the Qt dialog viewer."""
        # Ignore no-op edits to prevent false "modified" cards/files in UI.
        if (original or "") == (new_content or ""):
            log.info(f"[Diff] No-op edit ignored for: {file_path}")
            return

        if not hasattr(self, '_diff_data_store'):
            self._diff_data_store = {}
        norm_path = os.path.normcase(os.path.normpath(file_path))
        # CRITICAL: Only store the TRUE original (first edit before any AI changes).
        # If the AI edits the same file multiple times (A→B→C), we must preserve
        # original=A so Reject always reverts to the pre-AI state, not mid-chain.
        if file_path not in self._diff_data_store:
            self._diff_data_store[file_path] = (original, new_content)
        else:
            # Keep original, update modified to latest
            true_original = self._diff_data_store[file_path][0]
            self._diff_data_store[file_path] = (true_original, new_content)
        if norm_path not in self._diff_data_store:
            self._diff_data_store[norm_path] = (original, new_content)
        else:
            true_original = self._diff_data_store[norm_path][0]
            self._diff_data_store[norm_path] = (true_original, new_content)

        # ── Invalidate file_manager cache so _open_file reads fresh content ──
        try:
            resolved = str(Path(file_path).resolve())
            self._file_manager._file_cache.put(resolved, new_content)
            self._file_manager._hash_cache[resolved] = self._file_manager._compute_hash(new_content)
            if resolved in self._file_manager._open_files:
                self._file_manager._open_files[resolved] = new_content
            log.debug(f"[Diff] Updated file_manager cache for: {file_path} ({len(new_content)} chars)")
        except Exception as e:
            log.debug(f"[Diff] Cache update failed: {e}")
        # Persist to diff cache for cross-session diff viewing
        cache = self._load_diff_cache()
        cache[file_path] = {
            'original': original,
            'modified': new_content,
            'ts': int(__import__('time').time())
        }
        # Prune cache to last 100 entries
        if len(cache) > 100:
            items = sorted(cache.items(), key=lambda kv: kv[1].get('ts', 0), reverse=True)
            self._diff_cache = dict(items[:100])
        else:
            self._diff_cache = cache
        self._save_diff_cache()
        log.info(f"[Diff] Stored diff data for: {file_path} (original: {len(original)} chars, new: {len(new_content)} chars)")
        log.debug(f"[Diff] _diff_data_store now has {len(self._diff_data_store)} entries")

        # Update sidebar changed files panel
        try:
            edit_type = "C" if not original else "M"
            if self._changed_files_panel:
                self._changed_files_panel.add_file(file_path, edit_type)
        except Exception as e:
            log.debug(f"Sidebar update skipped: {e}")

    # ============================================================
    # FILE OPERATION CARDS (Create/Edit with animation)
    # ============================================================
    
    def _on_file_creating_started(self, file_path: str):
        """Show 'Creating file...' card with pulse animation."""
        try:
            card_id = self._ai_chat.show_file_creating_card(file_path)
            if not hasattr(self, '_file_op_cards'):
                self._file_op_cards = {}
            self._file_op_cards[file_path] = card_id
            log.debug(f"[FileOp] Started creating card ({card_id}) for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to show creating card: {e}")

    def _on_file_editing_started(self, file_path: str):
        """Show 'Editing file...' card with pulse animation."""
        try:
            if not hasattr(self, '_file_op_cards'):
                self._file_op_cards = {}
            # If there's already an active card for this file, remove/complete it first
            old_card_id = self._file_op_cards.get(file_path)
            if old_card_id:
                self._ai_chat.dismiss_file_op_card(old_card_id)
                self._file_op_cards.pop(file_path, None)
            card_id = self._ai_chat.show_file_editing_card(file_path)
            self._file_op_cards[file_path] = card_id
            log.debug(f"[FileOp] Started editing card ({card_id}) for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to show editing card: {e}")

    def _on_file_operation_completed(self, _unused_card_id: str, file_path: str, content: str, op_type: str):
        """Transform operation card to show completed file."""
        try:
            # Use the card_id stored when the card was first created (from show_file_*_card)
            # The card_id from agent_bridge is different — it was generated before the JS card was made.
            real_card_id = getattr(self, '_file_op_cards', {}).get(file_path)
            if not real_card_id:
                log.debug(f"[FileOp] No stored card_id for {file_path}, skipping completion")
                return
            if op_type == "create":
                self._ai_chat.complete_file_creating_card(real_card_id, file_path, content)
            else:
                original = ""
                if hasattr(self, '_diff_data_store'):
                    if file_path in self._diff_data_store:
                        original, _ = self._diff_data_store[file_path]
                    else:
                        norm_path = os.path.normcase(os.path.normpath(file_path))
                        if norm_path in self._diff_data_store:
                            original, _ = self._diff_data_store[norm_path]
                # If edit produced no textual change, dismiss stale card instead of
                # showing a misleading "modified" file entry.
                if (original or "") == (content or ""):
                    self._ai_chat.dismiss_file_op_card(real_card_id)
                    self._file_op_cards.pop(file_path, None)
                    log.info(f"[FileOp] Suppressed no-op edit card for: {file_path}")
                    return
                self._ai_chat.complete_file_editing_card(real_card_id, file_path, original, content)
            # Clean up stored card_id
            self._file_op_cards.pop(file_path, None)
            log.debug(f"[FileOp] Completed {op_type} card for: {file_path}")
        except Exception as e:
            log.debug(f"[FileOp] Failed to complete operation card: {e}")

    def _on_ai_question_requested(self, question_payload: dict):
        """Handle AI asking a question that requires user response in chat."""
        log.info(f"AI requested user input: {question_payload.get('text', question_payload.get('question', ''))[:50]}...")
        # Structuring the question info for the JS UI
        # CRITICAL: Use permission_request_id if available (for permission cards),
        # otherwise fall back to tool_call_id (for general questions)
        request_id = question_payload.get("permission_request_id", question_payload.get("id", str(_uuid.uuid4())))
        raw_choices = question_payload.get("choices", [])
        # Normalize choices to list of dicts
        normalized = []
        for c in raw_choices:
            if isinstance(c, dict):
                normalized.append(c)
            elif isinstance(c, str):
                normalized.append({"label": c, "value": c})
        info = {
            "id": request_id,
            "text": question_payload["text"],
            "type": question_payload.get("type", "text"),
            "choices": normalized,
            "default": question_payload.get("default", ""),
        }
        self._ai_chat.on_question(
            info.get('id', ''),
            info.get('text', ''),
            info.get('type', 'text'),
            info.get('choices', []),
            info.get('default', '')
        )
        try:
            preview = (info.get('text') or '')
            preview = preview.replace('\n', ' ').strip()
            # Permission cards get their own notification type (always unfocused-only)
            is_permission = bool(question_payload.get("permission_request_id"))
            if is_permission:
                from src.utils.notifications import notify_permission_required
                notify_permission_required(preview)
            else:
                from src.utils.notifications import notify_input_needed
                notify_input_needed(preview)
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')

    def _open_file_at_line(self, file_path: str, line_number: int):
        """Open file and navigate to specific line."""
        try:
            # First open the file — this is the critical part
            self._open_file(file_path)
            
            # Navigate to line in webview panel (Monaco editor)
            # Use QTimer to ensure the file is fully loaded before navigating
            if hasattr(self, '_webview_panel') and self._webview_panel:
                QTimer.singleShot(300, lambda: self._webview_panel._safe_run_js(
                    f"if (typeof goToLine === 'function') goToLine({line_number});"
                ))
                # Extra safety: force renderTabs() to guarantee tab header shows
                QTimer.singleShot(150, lambda: self._webview_panel._safe_run_js("renderTabs();"))
            
            # Legacy: also try the old editor tabs (CodeEditor) if available
            try:
                editor = self._editor_tabs.currentWidget() if hasattr(self, '_editor_tabs') else None
                if isinstance(editor, CodeEditor):
                    line_index = max(0, line_number - 1)
                    editor.setCursorPosition(line_index, 0)
                    editor.ensureLineVisible(line_index)
                    log.info(f"Navigated to line {line_number} in {file_path}")
            except Exception:
                pass  # Old editor not available — webview panel handles it
        except Exception as e:
            log.error(f"[OpenAtLine] Failed to open {file_path} at line {line_number}: {e}")
            # Fallback: try opening without line navigation
            try:
                self._open_file(file_path)
            except Exception as e2:
                log.error(f"[OpenAtLine] Fallback also failed: {e2}")

    def _on_accept_file_edit(self, file_path: str):
        """Accept AI edit — the file is already written to disk, just acknowledge."""
        if file_path == "__ALL__":
            self._on_accept_all_files()
            return
        file_path = os.path.normpath(file_path)
        log.info(f"[Accept] User accepted edit: {file_path}")

        # Tell the AI the file state changed so it re-reads before editing again
        if hasattr(self, '_ai_agent') and self._ai_agent:
            try:
                self._ai_agent.queue_accept_nudge(file_path)
            except Exception:
                pass

        # Clear stale snapshot so the next open stores the fresh disk content
        if hasattr(self, '_file_snapshots') and file_path in self._file_snapshots:
            del self._file_snapshots[file_path]

        # Invalidate file_manager cache so _open_file reads fresh content from disk
        try:
            resolved = str(Path(file_path).resolve())
            if hasattr(self._file_manager, '_file_cache'):
                with self._file_manager._file_cache._lock:
                    self._file_manager._file_cache.cache.pop(resolved, None)
            if hasattr(self._file_manager, '_hash_cache'):
                self._file_manager._hash_cache.pop(resolved, None)
        except Exception:
            pass

        # Open/refresh the file in the editor so the user sees the accepted state
        # Use force_reload_file to bypass throttle and guarantee content delivery
        try:
            content = self._file_manager.read(file_path, lazy_load=False, use_cache=False)
            if content is not None:
                language = detect_language(file_path)
                self._webview_panel.force_reload_file(file_path, content, language)
                log.info(f"File force-reloaded in editor after accept: {file_path} ({len(content)} chars)")
            else:
                self._open_file(file_path, reload=True)
        except Exception as e:
            log.error(f"Failed to force-reload file: {e}")
            self._open_file(file_path, reload=True)
        # Flash the editor to draw attention
        if hasattr(self, '_webview_panel') and self._webview_panel._page_loaded:
            self._webview_panel.flash_editor()
        self.statusBar().showMessage(f"✓ Accepted changes to {Path(file_path).name}", 3000)

        # Sync saved content — the file was just written to disk by the AI,
        # so the editor content matches disk. Without this, the debounced
        # onContentChanged would compare against stale _saved_content and
        # falsely show the white dot as modified.
        if hasattr(self, '_saved_content'):
            cached = self._file_manager.get_cached_content(file_path)
            if cached is not None:
                self._saved_content[file_path] = cached
        if hasattr(self, '_modified_files'):
            self._modified_files.discard(file_path)
        self._sync_window_title()
        # Clear white dot in editor tab
        if hasattr(self, '_webview_panel') and self._webview_panel._page_loaded:
            self._webview_panel.mark_modified(file_path, False)

        # Clean up tracking state
        if hasattr(self, '_diff_data_store'):
            norm = os.path.normcase(file_path)
            self._diff_data_store.pop(file_path, None)
            self._diff_data_store.pop(norm, None)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.pop(file_path, None)

        if self._changed_files_panel:
            self._changed_files_panel.remove_file(file_path)

        # Refresh sidebar file tree to show new/modified file in real-time
        if hasattr(self, '_sidebar'):
            self._sidebar.refresh()

    def _on_reject_file_edit(self, file_path: str):
        """Reject AI edit — write original content back to disk and reload editor."""
        if file_path == "__ALL__":
            self._on_reject_all_files()
            return
        file_path = os.path.normpath(file_path)
        log.info(f"[Reject] User rejected edit: {file_path}")

        original = self._get_original_content(file_path)

        if original is not None:
            try:
                # Normalize line endings to prevent doubled empty lines
                original = original.replace("\r\n", "\n").replace("\r", "\n")
                Path(file_path).write_text(original, encoding='utf-8', newline='')
                log.info(f"[Reject] Reverted {file_path} ({len(original)} chars)")
            except Exception as e:
                log.error(f"[Reject] Failed to revert {file_path}: {e}")
                self.statusBar().showMessage(f"✗ Revert failed for {Path(file_path).name}: {e}", 5000)
                return

            # Reload in editor so editor shows the reverted content
            self._open_file(file_path, reload=True)
            self.statusBar().showMessage(f"↩ Reverted {Path(file_path).name} to original", 3000)
        else:
            # No original found in snapshots — try git checkout as last resort
            reverted_via_git = False
            try:
                import subprocess
                project_root = str(getattr(self, '_project_root', ''))
                if project_root and os.path.isabs(file_path):
                    rel = os.path.relpath(file_path, project_root)
                    if not rel.startswith('..'):
                        result = subprocess.run(
                            ['git', 'checkout', '--', rel],
                            capture_output=True, text=True,
                            cwd=project_root, timeout=3,
                            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
                        )
                        if result.returncode == 0:
                            log.info(f"[Reject] Git checkout reverted: {rel}")
                            reverted_via_git = True
                            self._open_file(file_path, reload=True)
                            self.statusBar().showMessage(
                                f"↩ Reverted {Path(file_path).name} via git", 3000
                            )
            except Exception as e:
                log.debug(f"[Reject] Git checkout fallback failed: {e}")

            if not reverted_via_git:
                log.warning(f"[Reject] No original content for {file_path} — opening for review")
                self._open_file(file_path, reload=True)
                self.statusBar().showMessage(
                    f"⚠ No original content found for {Path(file_path).name} — review manually", 5000
                )

        # Clean up tracking state
        if hasattr(self, '_diff_data_store'):
            norm = os.path.normcase(file_path)
            self._diff_data_store.pop(file_path, None)
            self._diff_data_store.pop(norm, None)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.pop(file_path, None)

        if self._changed_files_panel:
            self._changed_files_panel.remove_file(file_path)

        # Refresh sidebar file tree to show reverted/restored state
        if hasattr(self, '_sidebar'):
            self._sidebar.refresh()

    def _get_original_content(self, file_path: str) -> Optional[str]:
        """
        Look up the original (pre-AI-edit) content for a file.
        Priority: _diff_data_store → _file_snapshots → diff_cache.json → git HEAD
        """
        norm = os.path.normcase(os.path.normpath(file_path))

        # 1. In-memory diff store (most reliable — set by _on_file_edited_diff_for_js)
        if hasattr(self, '_diff_data_store'):
            for key, (original, _modified) in self._diff_data_store.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    if original:  # empty string means new file
                        return original

        # 2. File snapshot taken when file was first opened
        if hasattr(self, '_file_snapshots'):
            for key, content in self._file_snapshots.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    return content

        # 3. Persisted diff cache
        try:
            cache = self._load_diff_cache()
            for key, entry in cache.items():
                if os.path.normcase(os.path.normpath(key)) == norm:
                    orig = entry.get('original', '')
                    if orig:
                        return orig
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')

        # 4. Git HEAD fallback — get the last committed version
        try:
            import subprocess
            project_root = str(getattr(self, '_project_root', ''))
            if project_root and os.path.isabs(file_path):
                rel = os.path.relpath(file_path, project_root)
                if not rel.startswith('..'):
                    result = subprocess.run(
                        ['git', 'show', f'HEAD:{rel}'],
                        capture_output=True, text=True,
                        cwd=project_root, timeout=3,
                        creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0
                    )
                    if result.returncode == 0 and result.stdout:
                        log.info(f"[Reject] Got original from git HEAD for: {rel}")
                        return result.stdout
        except Exception as e:
            log.debug(f"[Reject] Git fallback failed: {e}")

        return None
    
    def _on_load_full_chat_requested(self, conversation_id: str):
        """Load full chat messages from SQLite database."""
        if getattr(self, '_is_native_chat', False):
            log.info(f"[native] Load chat requested: {conversation_id} — not yet wired")
            return
        log.info(f"Loading full chat: {conversation_id}")
        
        try:
            # Load from SQLite via bridge
            full_chat_json_str = self._ai_chat.load_full_chat_from_sqlite(conversation_id)
            
            if full_chat_json_str and full_chat_json_str != "[]":
                log.info(f"Loaded {len(full_chat_json_str)} chars of chat data")
                # Send back to JavaScript
                self._ai_chat._view.page().runJavaScript(
                    f"window.handleFullChatLoad('{conversation_id}', {full_chat_json_str});"
                )
                self.statusBar().showMessage(f"✓ Loaded chat history", 2000)
            else:
                log.warning(f"No chat data found for: {conversation_id}")
                self._ai_chat._view.page().runJavaScript(
                    f"window.handleFullChatLoad('{conversation_id}', null);"
                )
        except Exception as e:
            log.error(f"Failed to load full chat: {e}")
            self._ai_chat._view.page().runJavaScript(
                f"window.handleFullChatLoad('{conversation_id}, null);"
            )

    def _on_chat_selected(self, conversation_id: str):
        """User clicked a chat in the sidebar history panel → load it."""
        if not hasattr(self, '_current_displayed_chat_id'):
            self._current_displayed_chat_id = None
        if self._current_displayed_chat_id == str(conversation_id):
            log.debug(f"[ChatHistory] Chat {conversation_id} already displayed — skipping reload")
            return
        self._current_displayed_chat_id = str(conversation_id)
        log.info(f"[ChatHistory] Selected chat: {conversation_id}")

        if getattr(self, '_is_native_chat', False):
            # Native backend: rebuild widgets from timeline data
            self._ai_chat.load_chat(conversation_id)
        else:
            # Webview backend: load via JavaScript
            js_code = f"window.loadChat('{conversation_id}');"
            self._ai_chat._view.page().runJavaScript(js_code)

    def _on_chat_renamed(self, conversation_id: str, new_title: str):
        """User renamed a chat in the sidebar — update store."""
        log.info(f"[ChatHistory] Renamed chat {conversation_id} → {new_title}")
        if not getattr(self, '_is_native_chat', False):
            safe_title = new_title.replace("'", "\\'").replace('"', '\\"')
            js_code = f"if(window.updateChatTitle) window.updateChatTitle('{conversation_id}', '{safe_title}');"
            self._ai_chat._view.page().runJavaScript(js_code)

    def _on_chat_delete_requested(self, conversation_id: str):
        """User requested delete from sidebar → show confirmation dialog."""
        reply = QMessageBox.question(
            self, "Delete Chat",
            "Are you sure you want to permanently delete this chat?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        log.info(f"[ChatHistory] Deleting chat: {conversation_id}")
        try:
            from src.core.chat_history import get_chat_history
            history = get_chat_history()
            history.delete_conversation(conversation_id)
        except Exception as e:
            log.error(f"[ChatHistory] Delete failed: {e}")
        if hasattr(self, '_current_displayed_chat_id') and self._current_displayed_chat_id == str(conversation_id):
            self._current_displayed_chat_id = None
        if not getattr(self, '_is_native_chat', False):
            self._ai_chat._view.page().runJavaScript(
                f"if(window.removeChatFromList) window.removeChatFromList('{conversation_id}');"
            )

    def _on_new_chat_requested(self):
        """User confirmed '+ New Chat' → summarize current chat into MEMORY.md
        FIRST, then clear the UI and start a fresh session.

        Shows spinner overlay IMMEDIATELY so the user sees feedback, then
        kicks off the background summarize thread. On completion (or timeout),
        the spinner hides and the fresh chat appears.
        """
        log.info("[NewChat] Requested — summarizing current chat to MEMORY.md first")

        if getattr(self, '_new_chat_saving', False):
            log.info("[NewChat] Already saving — ignoring duplicate request")
            return
        self._new_chat_saving = True

        # Show spinner IMMEDIATELY so user sees feedback right away
        if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, '_spinner_overlay'):
            self._ai_chat._spinner_overlay.show_overlay(
                message="Summarizing & compacting chat…",
                spinner_key='saving_memory',
            )

        if (getattr(self, '_is_native_chat', False)
                and hasattr(self, '_ai_agent') and self._ai_agent
                and hasattr(self._ai_agent, 'save_session_to_memory')):
            from PyQt6.QtCore import QTimer

            # Sync bridge conv_id with chat panel's so DB query finds messages
            if hasattr(self, '_ai_chat') and self._ai_chat:
                _panel_conv = self._ai_chat.conversation_id()
                if _panel_conv and _panel_conv != getattr(self._ai_agent, '_current_conversation_id', None):
                    self._ai_agent._current_conversation_id = _panel_conv
                    log.info(f"[NewChat] Synced bridge conv_id to {_panel_conv}")

            _finalized = False
            _save_started_at = time.monotonic()

            def _do_finalize():
                self._finalize_new_chat()

            def _finish(success: bool):
                nonlocal _finalized
                if _finalized:
                    return
                _finalized = True
                self._new_chat_saving = False
                try:
                    self._ai_agent.session_saved_to_memory.disconnect(_finish)
                except Exception:
                    pass
                if success:
                    log.info("[NewChat] MEMORY.md saved — clearing chat")
                else:
                    log.warning("[NewChat] Memory save incomplete — clearing anyway")
                # Keep spinner visible for at least 800ms so user sees it
                elapsed_ms = int((time.monotonic() - _save_started_at) * 1000)
                remaining = max(0, 800 - elapsed_ms)
                QTimer.singleShot(remaining, _do_finalize)

            def _timeout_fallback():
                nonlocal _finalized
                if _finalized:
                    return
                _finalized = True
                self._new_chat_saving = False
                log.warning("[NewChat] Timeout — session_saved_to_memory did not fire; "
                            "clearing chat anyway")
                try:
                    self._ai_agent.session_saved_to_memory.disconnect(_finish)
                except Exception:
                    pass
                _do_finalize()

            try:
                self._ai_agent.session_saved_to_memory.connect(_finish)
                self._ai_agent.save_session_to_memory()
                # Increased timeout: LLM summarization can take 30+ seconds
                QTimer.singleShot(45000, _timeout_fallback)
                return
            except Exception as e:
                self._new_chat_saving = False
                log.warning(f"[NewChat] memory-save failed to start, clearing anyway: {e}")
        else:
            self._new_chat_saving = False

        self._finalize_new_chat()

    def _finalize_new_chat(self):
        """Clear the chat UI and reset to a brand-new conversation that PERSISTS.

        The new (empty) conversation becomes the project's active conversation, so
        after closing/reopening the IDE the old chat does NOT reload — the fresh
        chat stays fresh. Context is preserved only in MEMORY.md (already saved).
        """
        log.info("[NewChat] Finalizing — clearing UI and creating fresh session")

        # Hide spinner overlay
        if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, '_spinner_overlay'):
            self._ai_chat._spinner_overlay.hide_overlay()

        if hasattr(self, '_current_displayed_chat_id'):
            self._current_displayed_chat_id = None

        new_id = None
        try:
            if hasattr(self, '_ai_agent') and self._ai_agent and hasattr(self._ai_agent, 'start_fresh_session'):
                new_id = self._ai_agent.start_fresh_session()
        except Exception as e:
            log.warning(f"[NewChat] start_fresh_session failed: {e}")
        if not new_id:
            import uuid
            new_id = str(uuid.uuid4())

        # Clear the UI
        if getattr(self, '_is_native_chat', False):
            self._ai_chat.new_chat()
            self._ai_chat.set_conversation_id(new_id)
        else:
            self._ai_chat._view.page().runJavaScript("window.newChat();")

        # Persist this empty conversation as the project's ACTIVE one so restart
        # restores it (fresh) instead of the old conversation that has timeline.
        self._set_active_conversation(new_id)
        log.info(f"[NewChat] Fresh chat ready (active conv={new_id}) — context in MEMORY.md")

    def _active_conv_key(self) -> str:
        """Settings key for the project's active conversation id."""
        root = ''
        if hasattr(self, '_project_manager') and self._project_manager.root:
            root = str(self._project_manager.root)
        root = root or getattr(self, '_current_project_path', '') or os.getcwd()
        import hashlib
        return 'active_' + hashlib.md5(os.path.normpath(root).encode('utf-8')).hexdigest()[:12]

    def _set_active_conversation(self, conv_id: str):
        """Create the conversation row + remember it as the project's active chat."""
        try:
            from src.core.chat_history import get_chat_history
            history = get_chat_history()
            project_path = os.path.normpath(
                (str(self._project_manager.root) if getattr(self, '_project_manager', None)
                 and self._project_manager.root else '')
                or getattr(self, '_current_project_path', '') or os.getcwd()
            )
            # Ensure the (empty) conversation exists so restart can resolve it.
            try:
                history.create_conversation(project_path, conversation_id=conv_id)
            except Exception:
                pass
            # Persist the active-conversation pointer.
            if hasattr(self, '_settings') and self._settings:
                self._settings.set('chat_active', self._active_conv_key(), conv_id)
            log.info(f"[NewChat] Active conversation set to {conv_id} for {project_path}")
        except Exception as e:
            log.warning(f"[NewChat] _set_active_conversation failed: {e}")

    def _on_switch_chat_context(self, conversation_id: str):
        """User switched to a different chat — restore AI conversation context.
        
        This ensures the AI has the correct conversation history loaded so
        follow-up questions reference the right context after a crash/reopen
        or when switching between chats.
        """
        try:
            if hasattr(self, '_ai_agent') and self._ai_agent:
                self._ai_agent.restore_conversation_context(conversation_id)
                log.info(f"[MainWindow] AI context switched to conversation: {conversation_id}")
        except Exception as e:
            log.warning(f"[MainWindow] Failed to switch AI context: {e}")

    def _setup_chat_persistence(self):
        """Wire ChatPanel save/restore to SQLite chat_history."""
        try:
            from src.core.chat_history import get_chat_history
            history = get_chat_history()
            # Use _current_project_path set by _on_project_opened, fallback to _project_path
            project_path = getattr(self, '_current_project_path', '') or getattr(self, '_project_path', '') or os.getcwd()
            project_path = os.path.normpath(project_path)  # Normalize path for consistent matching
            log.info(f"[ChatPersist] Setting up for project_path='{project_path}'")
            # Prefer the explicitly-set ACTIVE conversation (set by New Chat) so a
            # fresh chat stays fresh on restart instead of reloading the old chat.
            _active_id = None
            try:
                if hasattr(self, '_settings') and self._settings:
                    _active_id = self._settings.get('chat_active', self._active_conv_key(), default=None)
            except Exception:
                _active_id = None
            conv_id = history.get_or_create_conversation(project_path, conversation_id=_active_id)
            self._ai_chat.set_conversation_id(conv_id)
            # Sync bridge so save_session_to_memory queries the right conversation
            if hasattr(self, '_ai_agent') and self._ai_agent:
                self._ai_agent._current_conversation_id = conv_id
                log.info(f"[ChatPersist] Bridge conv_id synced to {conv_id}")
            # Restore last timeline if any (async — non-blocking, shows spinner)
            try:
                # Tell the DB to skip flush_write_queue during restore —
                # no writes are pending and the flush would block the UI.
                if hasattr(history, 'db') and history.db:
                    history.db._restoring = True
                timeline_data = history.get_timeline(conv_id)
                # DB read done — re-enable flush_write_queue
                if hasattr(history, 'db') and history.db:
                    history.db._restoring = False
                log.info(f"[ChatPersist] Timeline data type: {type(timeline_data)}, length: {len(timeline_data) if timeline_data else 0}")
                if timeline_data and isinstance(timeline_data, dict) and timeline_data.get("messages"):
                    # Suppress sidebar refreshes during restore to prevent WebEngine flashing
                    self._sidebar.set_suppress_refresh(True)
                    log.info(f"[ChatPersist] Calling load_timeline_async with {len(timeline_data.get('messages', []))} messages")
                    self._ai_chat.load_timeline_async(timeline_data)
                    log.info(f"[ChatPersist] Queued async restore of {len(timeline_data.get('messages', []))} messages for {conv_id} (project={project_path})")
                    # Re-enable sidebar refreshes after restore completes (safety: 35s)
                    QTimer.singleShot(35000, lambda: self._sidebar.set_suppress_refresh(False))
                elif timeline_data and isinstance(timeline_data, list) and len(timeline_data) > 0:
                    # Legacy format: timeline is a list of message dicts
                    self._sidebar.set_suppress_refresh(True)
                    log.info(f"[ChatPersist] Calling load_timeline_async (legacy) with {len(timeline_data)} messages")
                    self._ai_chat.load_timeline_async({"conversation_id": conv_id, "messages": timeline_data})
                    log.info(f"[ChatPersist] Queued async restore of {len(timeline_data)} messages (legacy list) for {conv_id}")
                    QTimer.singleShot(35000, lambda: self._sidebar.set_suppress_refresh(False))
                else:
                    # No history — show empty state (ring logo)
                    log.info(f"[ChatPersist] No chat history for project={project_path} — showing fresh state")
            except Exception as restore_err:
                log.warning(f"[ChatPersist] Restore failed (will still save new messages): {restore_err}")
                import traceback
                traceback.print_exc()
                # Ensure DB flush is re-enabled even on error
                if hasattr(history, 'db') and history.db:
                    history.db._restoring = False

            # ── CRASH RECOVERY: Check if previous session crashed mid-turn ──
            try:
                from src.core.crash_persistence import get_crash_store
                crash_store = get_crash_store()
                if crash_store.was_crash_detected(conv_id):
                    unsaved = crash_store.get_unsaved_turns(conv_id)
                    if unsaved:
                        # get_unsaved_turns returns the WHOLE crash log (up to
                        # 500 msgs), which duplicated the timeline restore on
                        # screen after every crash and inflated the widget tree
                        # each session (126 widgets → theme switch repolishes
                        # them all → freeze → kill → even more next time).
                        # Drop messages already restored from the timeline.
                        already_restored = set()
                        try:
                            if timeline_data and isinstance(timeline_data, dict):
                                for _m in timeline_data.get("messages", []):
                                    already_restored.add(
                                        (_m.get("role"), (_m.get("content") or "").strip())
                                    )
                        except Exception:
                            pass
                        unsaved = [
                            m for m in unsaved
                            if (m.get("role"), (m.get("content") or "").strip())
                            not in already_restored
                        ]
                    if unsaved:
                        log.warning(
                            f"[CrashRecovery] IDE crashed last session — "
                            f"recovering {len(unsaved)} messages for {conv_id[:8]}.."
                        )
                        # Inject recovered messages into chat panel
                        self._sidebar.set_suppress_refresh(True)
                        self._ai_chat.load_recovered_messages(unsaved)
                        QTimer.singleShot(35000, lambda: self._sidebar.set_suppress_refresh(False))
                    else:
                        log.info("[CrashRecovery] All crash-log messages already in timeline — nothing to inject")
                # Mark clean shutdown for this conversation
                crash_store.mark_clean_shutdown(conv_id)
            except Exception as crash_err:
                log.debug(f"[CrashRecovery] Check skipped: {crash_err}")

            # ── Restore AI conversation context so it remembers last session ──
            # Without this, the AI's _conversation_history starts EMPTY on each
            # IDE restart. Chat messages display in the UI but the AI doesn't
            # have them in its working context — so it can't answer "what was
            # the last thing we worked on?" without re-reading MEMORY.md.
            try:
                if hasattr(self, '_ai_agent') and self._ai_agent:
                    self._ai_agent.restore_conversation_context(conv_id)
                    log.info(f"[ChatPersist] AI conversation context restored for {conv_id[:8]}..")
            except Exception as ctx_err:
                log.warning(f"[ChatPersist] Failed to restore AI context: {ctx_err}")

            # Save after each turn — connect to response_complete signal
            def _save_on_turn_done(*_args, **_kwargs):
                # Defer heavy serialization to next event loop iteration so
                # response_complete.emit() returns immediately and UI stays responsive.
                _fn = _do_save_turn  # capture local function
                QTimer.singleShot(0, _fn)

            def _do_save_turn():
                try:
                    # NEVER save while the panel is restoring (project switch):
                    # the conversation id already points at the NEW project but
                    # the widgets may still hold the OLD project's chat —
                    # saving that snapshot cross-contaminates conversations.
                    if getattr(self._ai_chat, '_restoring', False):
                        log.info("[ChatPersist] Save skipped — chat restore in progress")
                        return
                    # Serialize on the main thread (Qt widget access must happen here)
                    tl = self._ai_chat.get_timeline()
                    _msg_count = len(tl.get("messages", []))

                    # Run DB write in background thread to avoid blocking UI
                    import threading as _threading
                    def _do_save():
                        try:
                            history.save_timeline(self._ai_chat.conversation_id(), tl)
                            log.info(f"[ChatPersist] Saved {_msg_count} messages for {self._ai_chat.conversation_id()}")
                        except Exception as e:
                            log.warning(f"[ChatPersist] Save failed: {e}", exc_info=True)
                    _threading.Thread(target=_do_save, daemon=True, name="ChatPersistSave").start()

                    # AUTO-SAVE TO MEMORY.md — lightweight (no LLM call), runs in background.
                    if hasattr(self, '_ai_agent') and self._ai_agent:
                        def _do_memory_autosave():
                            try:
                                self._ai_agent._auto_save_to_memory_light()
                            except Exception as _e:
                                pass
                        _threading.Thread(
                            target=_do_memory_autosave, daemon=True,
                            name="MemoryAutoSave"
                        ).start()
                except Exception as e:
                    log.warning(f"[ChatPersist] Save failed: {e}", exc_info=True)
            # Disconnect the PREVIOUS project's save callback first — this
            # setup runs on every project switch, and stacked connections
            # meant one response_complete fired multiple saves (including
            # during the switch window → cross-project chat contamination).
            _old_cb = getattr(self, '_chat_save_callback', None)
            self._chat_save_callback = _save_on_turn_done
            # Connect to response_complete signal if agent is available
            if hasattr(self, '_ai_agent') and self._ai_agent:
                if _old_cb is not None:
                    try:
                        self._ai_agent.response_complete.disconnect(_old_cb)
                    except (TypeError, RuntimeError):
                        pass  # never connected / already gone
                try:
                    self._ai_agent.response_complete.connect(self._chat_save_callback)
                    log.info("[ChatPersist] Connected save callback to response_complete")
                except Exception as conn_err:
                    log.debug(f"[ChatPersist] Signal connection deferred: {conn_err}")
            log.info(f"[ChatPersist] Persistence wired for conv_id={conv_id} (project={project_path})")
        except Exception as e:
            log.warning(f"[ChatPersist] Setup failed: {e}", exc_info=True)
            # Still set a fallback callback so saves work even if setup partially failed
            if not hasattr(self, '_chat_save_callback') or not self._chat_save_callback:
                self._chat_save_callback = lambda *a, **kw: None

    def _on_accept_all_files(self):
        """Accept all pending AI edits — files already on disk, just clean up state."""
        log.info("[Accept All] User accepted all file edits")
        # Deduplicate keys — _diff_data_store has both raw and normcase keys per file
        seen = set()
        for file_path in list(getattr(self, '_diff_data_store', {}).keys()):
            try:
                norm = os.path.normcase(os.path.normpath(file_path))
                if norm in seen:
                    continue
                seen.add(norm)
                if os.path.isfile(norm):
                    # Tell the AI each accepted file changed
                    if hasattr(self, '_ai_agent') and self._ai_agent:
                        try:
                            self._ai_agent.queue_accept_nudge(norm)
                        except Exception:
                            pass
                    self._open_file(norm, reload=True)
                    # Sync saved content to prevent false white dot
                    if hasattr(self, '_saved_content'):
                        cached = self._file_manager.get_cached_content(norm)
                        if cached is not None:
                            self._saved_content[norm] = cached
                    if hasattr(self, '_modified_files'):
                        self._modified_files.discard(norm)
                    if hasattr(self, '_webview_panel') and self._webview_panel._page_loaded:
                        self._webview_panel.mark_modified(norm, False)
            except Exception as e:
                log.debug(f'[MainWindow] Suppressed error: {e}')
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        self._sync_window_title()
        self.statusBar().showMessage("✓ Accepted all changes", 3000)
        if self._changed_files_panel:
            self._changed_files_panel.clear_files()

    def _on_reject_all_files(self):
        """Reject all pending AI edits — revert each file to its original content."""
        log.info("[Reject All] User rejected all file edits")
        reverted, failed = 0, 0
        # Deduplicate keys — _diff_data_store has both raw and normcase keys per file
        seen = set()
        for file_path in list(getattr(self, '_diff_data_store', {}).keys()):
            try:
                norm = os.path.normcase(os.path.normpath(file_path))
                if norm in seen:
                    continue
                seen.add(norm)
                original = self._get_original_content(norm)
                if original is not None and os.path.isfile(norm):
                    # Normalize line endings to prevent doubled empty lines
                    original = original.replace("\r\n", "\n").replace("\r", "\n")
                    Path(norm).write_text(original, encoding='utf-8', newline='')
                    self._open_file(norm)
                    reverted += 1
                    log.info(f"[Reject All] Reverted {norm}")
                else:
                    failed += 1
            except Exception as e:
                log.error(f"[Reject All] Failed to revert {file_path}: {e}")
                failed += 1
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        msg = f"↩ Reverted {reverted} file(s)"
        if failed:
            msg += f" ({failed} could not be reverted)"
        self.statusBar().showMessage(msg, 5000)
        if self._changed_files_panel:
            self._changed_files_panel.clear_files()

    def _on_code_copied(self, text: str, file_path: str, start_line: int, end_line: int):
        """Store copy metadata so smart paste can use it after focus changes."""
        self._last_copy_info = {
            'text': text,
            'file_path': file_path,
            'start_line': start_line,
            'end_line': end_line,
        }

    def _on_smart_paste_check(self, pasted_text: str):
        """Check if pasted text matches last editor copy. Send result to chat."""
        try:
            import json as _json

            # Use stored copy metadata (captured at Ctrl+C time, before focus changed)
            copy_info = getattr(self, '_last_copy_info', None)
            if not copy_info or not copy_info.get('text'):
                self._ai_chat._view.page().runJavaScript(
                    "handleSmartPasteResult({isMatch: false});"
                )
                return

            def normalize(text):
                return '\n'.join(line.strip() for line in text.strip().split('\n') if line.strip())

            pasted_norm = normalize(pasted_text)
            copied_norm = normalize(copy_info['text'])

            if not copied_norm or not pasted_norm:
                self._ai_chat._view.page().runJavaScript(
                    "handleSmartPasteResult({isMatch: false});"
                )
                return

            is_match = (pasted_norm == copied_norm or
                        pasted_norm in copied_norm or
                        copied_norm in pasted_norm)

            if is_match and copy_info.get('file_path'):
                file_path  = copy_info['file_path']
                start_line = copy_info['start_line']
                end_line   = copy_info['end_line']
                file_name  = os.path.basename(file_path)
                ext        = os.path.splitext(file_path)[1].lstrip('.')
                line_range = str(start_line) if start_line == end_line else f"{start_line}-{end_line}"

                self._ai_chat._view.page().runJavaScript(
                    f"handleSmartPasteResult({{isMatch: true, "
                    f"filePath: {_json.dumps(file_path)}, "
                    f"fileName: {_json.dumps(file_name)}, "
                    f"lineRange: {_json.dumps(line_range)}, "
                    f"code: {_json.dumps(pasted_text)}, "
                    f"language: {_json.dumps(ext)}}});"
                )
                log.info(f"Smart paste matched: {file_name} lines {line_range}")
                # Clear after use so next paste starts fresh
                self._last_copy_info = None
                return

            self._ai_chat._view.page().runJavaScript(
                "handleSmartPasteResult({isMatch: false});"
            )

        except Exception as e:
            log.error(f"Smart paste check error: {e}")
            self._ai_chat._view.page().runJavaScript(
                "handleSmartPasteResult({isMatch: false});"
            )

    # ============================================================================
    # NEW: OpenCode Enhancement Integration Handlers
    # ============================================================================
    
    def _on_intent_classified(self, message: str, intent: str, confidence: float):
        """Handle intent classification from AI Integration Layer."""
        log.info(f"[Intent] {intent} (confidence: {confidence:.2f}): {message[:50]}...")
        # Could update UI to show detected intent
        
    def _on_agent_selected(self, agent_type: str, reason: str, confidence: float):
        """Handle agent selection from AI Integration Layer."""
        log.info(f"[Agent] Selected {agent_type} (confidence: {confidence:.2f}): {reason}")
        # Could show agent indicator in UI
        
    def _on_tools_selected(self, tool_names: list):
        """Handle tool selection from AI Integration Layer."""
        log.info(f"[Tools] Selected: {', '.join(tool_names)}")
        # Could show tool indicators in UI
        
    def _on_permission_requested(self, request_id: str, html_card: str):
        """Handle permission request - show permission card in chat."""
        log.info(f"[Permission] Request {request_id} - showing permission card")
        
        # Show permission card in AI chat
        if hasattr(self._ai_chat, 'show_permission_card'):
            self._ai_chat.show_permission_card(request_id, html_card)
        else:
            # Fallback: add as system message
            self._ai_chat.add_system_message("🔒 Permission required. Please check the chat interface.")
            
    def _on_permission_granted(self, request_id: str, scope: str):
        """Handle permission grant."""
        log.info(f"[Permission] Granted {request_id} with scope {scope}")
        
        # Retry the AI processing now that permission is granted
        # Get the last user message and retry
        if hasattr(self._ai_chat, '_last_user_message'):
            message = self._ai_chat._last_user_message
            context = []
            
            if self._project_manager.root:
                context.append(f"Project path: {self._project_manager.root}")
            
            editor = self._editor_tabs.current_editor()
            if editor:
                fp = self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
                if fp:
                    name = Path(fp).name
                    content = editor.get_all_text()
                    if len(content) > 5000:
                        content = content[:5000] + "... (truncated)"
                    context.append(f"Current file ({name}):\n```\n{content}\n```")
            
            full_context = "\n\n".join(context)
            
            # NEW: Check if we have enhancement data stored
            enhancement_data = self._ai_agent.get_last_enhancement_data()
            if enhancement_data.get("intent"):
                log.info("Retrying with chat_with_enhancement after permission grant")
                self._ai_agent.chat_with_enhancement(
                    message,
                    intent=enhancement_data["intent"],
                    route=enhancement_data["route"],
                    tools=enhancement_data["tools"],
                    code_context=full_context
                )
            else:
                if hasattr(self, '_sidebar') and self._sidebar:
                    self._sidebar._ai_active = True
                self._ai_agent.chat(message, full_context)
            
    def _on_permission_denied(self, request_id: str, reason: str):
        """Handle permission denial."""
        log.info(f"[Permission] Denied {request_id}: {reason}")
        self._ai_chat.add_system_message(f"❌ Permission denied: {reason}")
        
    def _on_chat_permission_response(self, request_id: str, approved: bool, scope: str = "session", remember: bool = False):
        """Handle permission response from chat UI.
        
        NOTE: This method is now DEPRECATED. The active permission flow uses
        permission_decided signal connected directly to bridge.on_permission_respond().
        This handler is kept for backwards compatibility but the permission_response
        signal is no longer connected.
        """
        log.info(f"[DEPRECATED] _on_chat_permission_response called: {request_id}, approved={approved}")
        # The real permission flow now goes through:
        #   _ai_chat.permission_decided -> _ai_agent.on_permission_respond(decision)
        # This method is no longer in the active path.
    
    def _on_user_denied_workflow(self, tool_name: str):
        """Handle user denying workflow twice - stop AI agent."""
        log.warning(f"User denied {tool_name} twice - stopping AI agent")
        
        # Stop the AI agent immediately
        self._ai_agent.stop()
        
        # Add system message explaining what happened
        self._ai_chat.add_system_message(
            f"⏹️ **Workflow Stopped**\n\n"
            f"You denied `{tool_name}` twice. The AI agent has stopped its current work.\n\n"
            f"If you'd like to continue with a different approach, please send a new message."
        )
        
        # Hide thinking indicator
        self._ai_chat.hide_thinking()
        
        # Reset UI - show send button again
        self._view.page().runJavaScript("if(window._onGenerationComplete) window._onGenerationComplete();")

    # ========== TESTING WORKFLOW HANDLERS (NEW) ==========
    
    def _on_testing_decision(self, decision: str, priority: str, trigger: str):
        """Handle testing decision signal."""
        log.info(f"[Testing] Decision: {decision} (priority: {priority}, trigger: {trigger})")
        
        # Show UI notification based on decision
        if decision == 'write_tests':
            self._ai_chat.add_system_message(
                f"🧪 **Testing Mode Activated**\n"
                f"Priority: {priority.upper()} | Trigger: {trigger}\n"
                f"The AI will analyze your code and suggest appropriate tests."
            )
        elif decision == 'skip_tests':
            log.debug("Testing skipped - no triggers detected")
    
    def _on_test_tools_selected(self, tools: list):
        """Handle test tools selection signal."""
        log.info(f"[Testing] Tools selected: {tools}")
        
        if tools:
            self._ai_chat.add_system_message(
                f"🔧 **Test Framework:** {', '.join(tools)}"
            )
    
    def _on_test_execution_started(self, test_type: str):
        """Handle test execution start signal."""
        log.info(f"[Testing] Execution started: {test_type}")
        self._ai_chat.add_system_message(f"▶️ Running {test_type} tests...")
    
    def _on_test_execution_completed(self, all_passed: bool, passed_count: int, failed_count: int):
        """Handle test execution completion signal."""
        log.info(f"[Testing] Execution completed: {passed_count} passed, {failed_count} failed")
        
        if all_passed:
            self._ai_chat.add_system_message(
                f"✅ **All Tests Passed!** ({passed_count} tests)"
            )
        else:
            self._ai_chat.add_system_message(
                f"⚠️ **Tests Completed:** {passed_count} passed, {failed_count} failed"
            )
    
    def _on_test_analysis_ready(self, analysis: dict):
        """Handle test analysis results signal."""
        log.info(f"[Testing] Analysis ready: {analysis.get('all_passed', False)}")
        
        # Display failure patterns if any
        patterns = analysis.get('patterns', [])
        if patterns:
            pattern_text = '\n'.join([f"- {p.get('type', 'unknown')}: {p.get('description', '')}" 
                                     for p in patterns[:3]])
            self._ai_chat.add_system_message(
                f"📊 **Failure Analysis:**\n{pattern_text}"
            )

    def _open_file_at_line_duplicate(self, filepath: str, line: int):
        self._open_file(filepath)
        editor = self._editor_tabs.current_editor()
        if editor:
            # Move cursor to line
            cursor = editor.textCursor()
            block = editor.document().findBlockByLineNumber(line - 1)
            cursor.setPosition(block.position())
            editor.setTextCursor(cursor)
            editor.centerCursor()

    def _save_current(self):
        fp = self._webview_panel.get_active_file()
        log.info(f"[Save] _save_current called, active_file={fp!r}")
        if not fp or fp == "untitled.py":
            log.info(f"[Save] _save_current early return — no active file")
            return

        # ── Race guard: set BEFORE the async get_current_content call ──────
        # The debounced onContentChanged (300ms) can fire before the async
        # callback sets _saved_content, causing the white dot to reappear
        # after save. _saving_files blocks _on_webview_content_changed
        # from re-marking the file during the save window.
        if not hasattr(self, '_saving_files'):
            self._saving_files: set = set()
        self._saving_files.add(fp)
        # Clear the white dot IMMEDIATELY (don't wait for async callback)
        self._webview_panel.mark_modified(fp, False)

        def _do_save(content: str, _retry_count: int = 0):
            # PyQt QWebChannel may convert JS empty string "" to Python None
            if content is None:
                content = ""
            log.info(f"[Save] _do_save called for {fp!r}, content_len={len(content)}, retry={_retry_count}")
            # ── SAFETY: Prevent accidental file wiping ──
            # If Monaco returns empty but file had content on disk, retry
            # up to 3 times with a delay (Monaco may still be loading).
            if len(content) == 0:
                try:
                    disk_content = self._file_manager.read(fp, lazy_load=False, use_cache=False)
                    if disk_content and len(disk_content) > 0:
                        if _retry_count < 3:
                            log.warning(f"[Save] Monaco returned empty for {fp!r} but disk has {len(disk_content)} chars — retrying ({_retry_count + 1}/3)")
                            QTimer.singleShot(500, lambda: self._webview_panel.get_current_content(fp, lambda c: _do_save(c, _retry_count + 1)))
                            return
                        else:
                            log.error(f"[Save] Monaco returned empty for {fp!r} after 3 retries — save FAILED, file NOT modified")
                            self._status_file.setText(f"  Save failed — Monaco not ready  ")
                            QTimer.singleShot(3000, lambda: self._update_status_file(fp))
                            return
                except Exception:
                    pass
            self._file_manager.write(fp, content)
            # Track saved content to prevent post-save re-mark by debounced onContentChanged
            if not hasattr(self, '_saved_content'):
                self._saved_content: dict = {}
            self._saved_content[fp] = content
            if hasattr(self, '_modified_files'):
                self._modified_files.discard(fp)
                self._sync_window_title()
            # ── CLEAR WHITE DOT: explicitly mark file as saved ──
            self._webview_panel.mark_modified(fp, False)
            # Release the save guard after the 300ms debounce window passes
            QTimer.singleShot(450, lambda fp=fp: self._saving_files.discard(fp))
            self._status_file.setText(f"  Saved \u2713  {self._status_file.text().strip()}")
            QTimer.singleShot(2000, lambda: self._update_status_file(fp))

        # ── ALWAYS fetch fresh content from Monaco (never trust Python cache) ──
        # The Python-side cache can be stale if onContentChanged (300ms debounce)
        # hasn't fired yet after a paste/delete. Monaco's model always has the
        # latest content. We fetch via async JS call with disk fallback.
        def _on_content(content: str):
            _do_save(content or "")
        self._webview_panel.get_current_content(fp, _on_content)

    def _save_all(self):
        """Save all open files — EACH file gets fresh Monaco content (fully isolated).

        Each file is fetched independently from Monaco's model via async JS.
        a.txt save NEVER touches b.txt. Empty a.txt save NEVER empties b.txt.
        """
        open_files = list(self._webview_panel._open_files.keys())
        files_to_save = [fp for fp in open_files if fp != "untitled.py"]

        if not files_to_save:
            return

        if not hasattr(self, '_saving_files'):
            self._saving_files: set = set()
        if not hasattr(self, '_saved_content'):
            self._saved_content: dict = {}

        remaining = len(files_to_save)

        def _save_one_file(fp: str):
            """Save a single file with fresh Monaco content — fully isolated."""
            def _on_content(content: str):
                nonlocal remaining
                # PyQt QWebChannel may convert JS empty string "" to Python None
                if content is None:
                    content = ""
                # ── Trust Monaco's response for this file ──
                # If user cleared a.txt → empty string → save empty a.txt.
                # b.txt is NEVER touched by this operation.
                self._file_manager.write(fp, content)
                self._webview_panel.mark_modified(fp, False)
                self._saved_content[fp] = content
                remaining -= 1
                if remaining == 0:
                    # All files saved — clear modified tracking
                    if hasattr(self, '_modified_files'):
                        self._modified_files.clear()
                        self._sync_window_title()
                    # Also save legacy editor tabs
                    for i in range(self._editor_tabs.count()):
                        editor = self._editor_tabs.widget(i)
                        legacy_fp = self._editor_tabs._files.get(i)
                        if isinstance(editor, CodeEditor) and legacy_fp:
                            text = editor.get_all_text()
                            self._file_manager.write(legacy_fp, text)
                    self._status_file.setText("  Saved ✓  All files")
                    QTimer.singleShot(2000, lambda: self._update_status_file(
                        self._webview_panel.get_active_file()))

            # ── Fetch FRESH content from Monaco for THIS file only ──
            self._saving_files.add(fp)
            self._webview_panel.get_current_content(fp, _on_content)
            # Release save guard after debounce window
            QTimer.singleShot(450, lambda fp=fp: self._saving_files.discard(fp))

        for fp in files_to_save:
            _save_one_file(fp)

    def _run_file(self):
        """Run File (Ctrl+F5) — must run the editor header's ACTIVE TAB.

        The Python-side active-file mirror can lag the editor (agent background
        reloads, image-preview branch, fast tab clicks while the GUI thread is
        busy streaming), which made Run File execute the sidebar-selected file
        instead of the visible tab. Ask the JS editor for the real active tab
        and only fall back to the mirror if the webview can't answer.
        """
        self._webview_panel.get_active_file_async(self._run_file_path)

    def _run_file_path(self, fp: str):
        fp = fp or self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
        log.info(f"[Run] _run_file resolved active file: {fp!r}")
        if not fp or not Path(fp).exists():
            return
        # HTML files → built-in Live Server (no terminal needed)
        if Path(fp).suffix.lower() in {".html", ".htm"}:
            self._save_current()
            self._run_live_server(fp)
            return
        self._save_current()

        # Ensure the integrated terminal panel is visible (VS Code style)
        if getattr(self, '_terminal_panel_hidden', False):
            self._toggle_terminal_panel(show=True)

        # Use the current terminal from the tab widget
        term = self._current_terminal()
        if not term:
            term = self._new_terminal()

        # Start the shell if it hasn't been started yet
        if not term._shell_started:
            term._start_shell()

        # Wait briefly for the shell to be ready, then execute
        def _send_cmd():
            if term._pty_process or (term._process and term._process.state() == QProcess.ProcessState.Running):
                lang = detect_language(fp)
                command = self._build_run_command(fp, lang)
                if command:
                    term.execute_command(command)
                    term.setFocus()
                else:
                    QMessageBox.information(self, "Run", f"Running {lang} is not yet supported.")
            else:
                # Shell still not ready — retry once more
                QTimer.singleShot(500, lambda: _retry_cmd(fp))

        def _retry_cmd(fp):
            if term._pty_process or (term._process and term._process.state() == QProcess.ProcessState.Running):
                lang = detect_language(fp)
                command = self._build_run_command(fp, lang)
                if command:
                    term.execute_command(command)
                    term.setFocus()
                else:
                    QMessageBox.information(self, "Run", f"Running {lang} is not yet supported.")
            else:
                log.warning("Integrated terminal shell not ready after retry")
                QMessageBox.warning(self, "Terminal", "Terminal shell is not ready yet. Please wait a moment and try again.")

        # Give the shell a moment to start if it was just initialized
        if term._shell_started and (term._pty_process or term._process):
            _send_cmd()
        else:
            QTimer.singleShot(300, _send_cmd)

    def _run_live_server(self, file_path: str):
        """Start built-in Live Server and open the HTML file in the browser."""
        import webbrowser
        from src.core.live_server import LiveServer

        # Stop any existing server
        if self._live_server and self._live_server.is_running:
            self._live_server.stop()

        root_dir = os.path.dirname(file_path)
        self._live_server = LiveServer(root_dir, file_path)
        try:
            port = self._live_server.start()
            url = self._live_server.get_url(file_path)
            webbrowser.open(url)
            log.info(f"Live Server started on port {port} → {url}")

            # Show status in status bar
            if hasattr(self, '_statusbar_label'):
                self._statusbar_label.setText(
                    f"Live Server  \u25cf  http://localhost:{port}   —   click \u25b6 to restart"
                )
        except Exception as e:
            log.error(f"Live Server failed to start: {e}")
            QMessageBox.warning(self, "Live Server", f"Failed to start Live Server:\n{e}")

    def _build_run_command(self, file_path: str, lang: str) -> str | None:
        """Build a run command for the current file based on language."""
        is_windows = platform.system() == "Windows"
        root = self._project_manager.root or str(Path(file_path).parent)
        build_dir = os.path.join(root, ".cortex_build")
        stem = Path(file_path).stem

        if is_windows:
            quote = lambda p: f'"{p}"'
            mkdir_cmd = f'New-Item -ItemType Directory -Force -Path {quote(build_dir)} | Out-Null'
        else:
            import shlex
            quote = shlex.quote
            mkdir_cmd = f'mkdir -p {quote(build_dir)}'

        if lang == "python":
            return f'python {quote(file_path)}'
        if lang in {"javascript", "jsx"}:
            return f'node {quote(file_path)}'
        if lang in {"typescript", "tsx"}:
            # Compile with tsc and run with node (faster than npx ts-node)
            js_out = os.path.join(build_dir, stem + ".js")
            if is_windows:
                return f'tsc {quote(file_path)} --outDir {quote(build_dir)}; if ($LASTEXITCODE -eq 0) {{ node {quote(js_out)} }}'
            else:
                return f'tsc {quote(file_path)} --outDir {quote(build_dir)} && node {quote(js_out)}'
        if lang == "bash":
            return f'bash {quote(file_path)}'
        if lang == "batch":
            return f'& {quote(file_path)}' if is_windows else None
        if lang == "powershell":
            return f'& {quote(file_path)}' if is_windows else f'pwsh {quote(file_path)}'
        if lang == "ruby":
            return f'ruby {quote(file_path)}'
        if lang == "php":
            return f'php {quote(file_path)}'
        if lang == "go":
            return f'go run {quote(file_path)}'
        if lang == "rust":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'rustc {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "c":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'gcc {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "cpp":
            exe_path = os.path.join(build_dir, stem + (".exe" if is_windows else ""))
            compile_cmd = f'g++ {quote(file_path)} -o {quote(exe_path)}'
            run_cmd = f'& {quote(exe_path)}' if is_windows else f'{quote(exe_path)}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        if lang == "java":
            package_name = ""
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        line = line.strip()
                        if line.startswith("package ") and line.endswith(";"):
                            package_name = line[len("package "):-1].strip()
                            break
            except Exception:
                package_name = ""
            class_name = f"{package_name}.{stem}" if package_name else stem
            compile_cmd = f'javac -d {quote(build_dir)} {quote(file_path)}'
            run_cmd = f'java -cp {quote(build_dir)} {class_name}'
            return f'{mkdir_cmd}; {compile_cmd}; if ($LASTEXITCODE -eq 0) {{ {run_cmd} }}' if is_windows else f'{mkdir_cmd} && {compile_cmd} && {run_cmd}'
        return None

    def _new_terminal(self, show_panel: bool = True) -> XTermWidget:
        # If call comes from a signal (like clicked), show_panel might be the 'checked' state (False usually)
        # So we force it to True if it's not explicitly False from our internal calls
        if not isinstance(show_panel, bool): show_panel = True
        
        term = XTermWidget()
        term.set_theme(self._theme_manager.is_dark)
        
        # Initialize with current project directory if available
        if self._project_manager.root:
            term.set_cwd(str(self._project_manager.root))

        # Use terminal's own name text for the tab title
        tab_title = term._terminal_name_text
        idx = self._terminal_tabs.addTab(term, tab_title)
        self._terminal_tabs.setCurrentIndex(idx)
        self._terminal_tabs.setTabToolTip(idx, tab_title)
        
        if show_panel:
            self._terminal_tabs.setVisible(True)
            # If terminal panel was hidden, show it
            if getattr(self, '_terminal_panel_hidden', False):
                self._toggle_terminal_panel(show=True)
            term.setFocus()
            
        # Hook up "New Terminal" button from within the terminal
        term.new_terminal_requested.connect(lambda: self._new_terminal(show_panel=True))
        
        # Hook up dropdown terminal switcher
        term.switch_to_terminal_requested.connect(self._switch_to_terminal_tab)
        
        # When this terminal's xterm.js is ready, sync the full terminal list to ALL terminals
        # (ensures dropdown gets populated even if the terminal wasn't ready during _sync_terminal_list)
        term.terminal_ready.connect(self._sync_terminal_list)
        
        # Store shell name reference for later use (used by corner shell label updates)
        term._shell_display_name = term._terminal_name_text.split()[0] if term._terminal_name_text else "PowerShell"
        
        # Connect file operations to AI chat
        term.file_operation_detected.connect(self._on_terminal_file_operation)
        
        # Sync the terminal list to all open terminals (updates their dropdowns)
        self._sync_terminal_list()
        
        return term
        

    def _on_terminal_file_operation(self, operation_type: str, file_path: str, status: str):
        """Handle file operations from terminal and show in AI chat."""
        # Map operation types to display format
        op_labels = {
            'create': 'Creating file',
            'create_dir': 'Creating directory',
            'delete': 'Deleting file',
            'delete_dir': 'Deleting directory',
            'move': 'Moving file',
            'copy': 'Copying file',
            'rename': 'Renaming file'
        }
        
        label = op_labels.get(operation_type, operation_type)
        display_info = f"{label}: {file_path}"
        
        # Show in AI chat as tool activity
        self._ai_chat.show_tool_activity('terminal_' + operation_type, display_info, status)

    def _on_terminal_tab_changed(self, index: int):
        """Update AI agent when active terminal changes."""
        term = self._terminal_tabs.widget(index)
        if isinstance(term, XTermWidget):
            self._ai_agent.set_terminal(term)
            # Update the HTML header label for the now-visible terminal
            if term._is_ready:
                term._bridge.update_terminal_label.emit(term._terminal_name_text)
        # Sync dropdown list across all terminals
        self._sync_terminal_list()

    def _close_terminal_tab(self, index: int):
        term = self._terminal_tabs.widget(index)
        if isinstance(term, XTermWidget):
            term._kill_process()
        self._terminal_tabs.removeTab(index)
        if self._terminal_tabs.count() == 0:
            self._terminal_tabs.setVisible(False)
        # Sync dropdown list across remaining terminals
        self._sync_terminal_list()

    def _kill_current_terminal(self):
        idx = self._terminal_tabs.currentIndex()
        if idx >= 0:
            self._close_terminal_tab(idx)

    def _current_terminal(self) -> XTermWidget | None:
        w = self._terminal_tabs.currentWidget()
        return w if isinstance(w, XTermWidget) else None

    def _sync_terminal_list(self):
        """Send the list of all open terminals to every terminal's dropdown."""
        terminals = []
        active_idx = self._terminal_tabs.currentIndex()
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                name = term._terminal_name_text
            else:
                name = self._terminal_tabs.tabText(i)
            terminals.append({"index": i, "name": name, "active": (i == active_idx)})
        json_str = json.dumps(terminals)
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget) and term._is_ready:
                term._bridge.update_terminal_list.emit(json_str)

    def _switch_to_terminal_tab(self, index: int):
        """Switch to a specific terminal tab (called from JS dropdown)."""
        if 0 <= index < self._terminal_tabs.count():
            self._terminal_tabs.setCurrentIndex(index)

    def _show_terminal_panel(self):
        """Show terminal panel (called from AI chat 'View in terminal' button)."""
        self._terminal_tabs.setVisible(True)
        if self._terminal_tabs.count() == 0:
            self._new_terminal()
        term = self._current_terminal()
        if term:
            term.setFocus()

    def _show_terminal_and_run(self, command: str):
        """Show terminal panel and execute command (called from 'View in terminal' with a command)."""
        self._terminal_tabs.setVisible(True)
        if self._terminal_tabs.count() == 0:
            self._new_terminal()
        # Ensure the terminal panel has a visible height
        self._terminal_tabs.setMinimumHeight(150)
        term = self._current_terminal()
        if term:
            term.setFocus()
            if command and command.strip():
                term.execute_command(command.strip())

    def _sync_splitter_handles(self, splitter: QSplitter):
        """Hide splitter handles next to hidden panels so no clipped sliver leaks.
        Visible handles get 6px. CursorSplitHandle does its own painting.
        PERFORMANCE: skips redundant setFixed calls when size is already correct."""
        is_vertical = splitter.orientation() == Qt.Orientation.Vertical
        sizes = splitter.sizes()
        for idx in range(1, splitter.count()):
            left = splitter.widget(idx - 1)
            right = splitter.widget(idx)
            handle = splitter.handle(idx)
            if handle is not None:
                left_ok = left.isVisible() and (idx - 1 < len(sizes) and sizes[idx - 1] > 0)
                right_ok = right.isVisible() and (idx < len(sizes) and sizes[idx] > 0)
                if left_ok and right_ok:
                    if not handle.isVisible():
                        handle.setVisible(True)
                    cur = handle.height() if is_vertical else handle.width()
                    if cur != 6:
                        if is_vertical:
                            handle.setFixedHeight(6)
                        else:
                            handle.setFixedWidth(6)
                elif left_ok or right_ok:
                    if not handle.isVisible():
                        handle.setVisible(True)
                    cur = handle.height() if is_vertical else handle.width()
                    if cur != 6:
                        if is_vertical:
                            handle.setFixedHeight(6)
                        else:
                            handle.setFixedWidth(6)
                else:
                    if handle.isVisible():
                        handle.setVisible(False)

    def _force_sync_all_splitter_handles(self):
        """Force re-sync ALL splitter handles after window is fully rendered.

        Called via QTimer.singleShot after showMaximized/showNormal to ensure
        CursorSplitHandle widgets get correct dimensions and cursor shapes.
        """
        if hasattr(self, '_main_splitter'):
            self._sync_splitter_handles(self._main_splitter)
            # Force update on all handles to trigger repaint
            for idx in range(1, self._main_splitter.count()):
                handle = self._main_splitter.handle(idx)
                if handle:
                    handle.update()
        if hasattr(self, '_editor_terminal_splitter'):
            self._sync_splitter_handles(self._editor_terminal_splitter)
            for idx in range(1, self._editor_terminal_splitter.count()):
                handle = self._editor_terminal_splitter.handle(idx)
                if handle:
                    handle.update()

    def _toggle_terminal_panel(self, show: bool = True):
        """Toggle integrated terminal panel via vertical splitter — VS Code style."""
        if not hasattr(self, '_editor_terminal_splitter'):
            return
        splitter = self._editor_terminal_splitter
        sizes = splitter.sizes()
        if len(sizes) < 2:
            return
        widget = splitter.widget(1)  # _terminal_tabs is bottom widget in vertical split
        if show:
            widget.setVisible(True)
            widget.setMinimumHeight(120)
            widget.setMaximumHeight(16777215)
            sizes[1] = self._terminal_panel_min_height
            # Auto-create first terminal if none exist yet
            if self._terminal_tabs.count() == 0:
                self._new_terminal(show_panel=False)
        else:
            widget.setMinimumHeight(0)
            # Don't set MaximumHeight(0) — that locks the widget to zero and
            # prevents the splitter handle from being draggable. Instead, just
            # set sizes to 0 and hide. The handle stays visible so the user
            # can drag to re-expand.
            widget.setVisible(False)
            sizes[1] = 0
        splitter.setSizes(sizes)
        self._sync_splitter_handles(splitter)
        self._terminal_panel_hidden = not show

    def _toggle_terminal(self):
        """Toggle terminal visibility via vertical splitter (Ctrl+J) — VS Code style."""
        hidden = getattr(self, '_terminal_panel_hidden', False)
        self._toggle_terminal_panel(show=hidden)

    def _toggle_live_preview(self, show: bool = True, file_path: str = "") -> None:
        """Show/hide the Live Preview panel via the horizontal editor↔preview
        splitter. Mirrors _toggle_terminal_panel's pattern exactly."""
        if not hasattr(self, '_editor_preview_splitter'):
            log.warning("[LivePreview] _editor_preview_splitter missing — panel never built")
            return
        splitter = self._editor_preview_splitter
        sizes = splitter.sizes()
        if len(sizes) < 2:
            log.warning(f"[LivePreview] splitter has {len(sizes)} panes, expected 2 — aborting")
            return
        widget = self._live_preview_panel
        log.info(f"[LivePreview] toggle show={show} file_path={file_path!r} "
                 f"current={widget.current_path()!r}")
        total = sum(sizes) or 1000
        if show:
            if file_path:
                widget.load_file(file_path)
            elif not widget.current_path():
                # Nothing to preview yet — try the currently active editor file.
                self._preview_active_file_if_html()
            widget.setVisible(True)
            # Preview takes the FULL editor area — no split view. Product
            # decision (user feedback on the 60/40 split): "browser fully
            # display it; if user needs code, close it". The code stays
            # open in Monaco underneath and reappears on close.
            #
            # The editor widget must be HIDDEN, not just sized to 0: this
            # splitter has setChildrenCollapsible(False), so setSizes([0,
            # total]) is silently clamped to the child's minimum width —
            # that clamp is exactly the phantom split the user reported
            # twice. Hidden widgets are exempt (same mechanism that keeps
            # the preview pane collapsed on startup).
            self._webview_panel.setVisible(False)
            sizes = [0, total]
        else:
            widget.setVisible(False)
            self._webview_panel.setVisible(True)
            # Restore the editor to full width. Just zeroing sizes[1] is
            # wrong here: the editor pane was collapsed to 0 while the
            # preview was showing, so that would leave BOTH panes at 0.
            sizes = [total, 0]
        splitter.setSizes(sizes)
        self._sync_splitter_handles(splitter)
        self._live_preview_hidden = not show

    def _toggle_live_preview_shortcut(self) -> None:
        """View menu / Ctrl+Shift+V entry point."""
        hidden = getattr(self, '_live_preview_hidden', True)
        self._toggle_live_preview(show=hidden)

    def _preview_active_file_if_html(self) -> None:
        """If the currently active Monaco tab is an HTML file, load it into
        the (already-visible) Live Preview panel. No-op otherwise — the
        panel just keeps showing whatever it last had, or its empty state."""
        try:
            path = self._webview_panel.get_active_file()
        except Exception as e:
            log.warning(f"[LivePreview] get_active_file() failed: {e}")
            path = None
        if path and path.lower().endswith(('.html', '.htm')):
            self._live_preview_panel.load_file(path)
        else:
            log.info(f"[LivePreview] active file {path!r} is not HTML — "
                     f"panel stays empty until a file is explicitly opened")

    def open_live_preview_for_file(self, file_path: str) -> None:
        """Public entry point (sidebar 'Open Live Preview' context menu action
        on .html files) — shows the panel loaded with this specific file."""
        log.info(f"[LivePreview] open_live_preview_for_file: {file_path}")
        self._toggle_live_preview(show=True, file_path=file_path)

    def _on_agent_live_preview(self, action: str, path: str, resp: dict) -> None:
        """Agent LivePreview tool handler — runs on the UI thread (queued
        signal from the agent bridge). Fills `resp` and sets resp['event']
        so the waiting agent thread can return the tool result.

        Contract: resp['event'] MUST be set on every path, including
        exceptions — a missed set() costs the agent a 15s timeout."""
        log.info(f"[LivePreview] agent tool: action={action!r} path={path!r}")
        try:
            panel = getattr(self, '_live_preview_panel', None)
            if panel is None:
                resp["error"] = "Live Preview panel is not available in this window."
                resp["event"].set()
                return
            if action == "open":
                self.open_live_preview_for_file(path)
                resp["result"] = (f"Live Preview opened: {path}. The panel auto-reloads on "
                                  f"file changes. Use action='console' to check for JS errors "
                                  f"and action='read' to see the rendered text.")
                resp["event"].set()
            elif action == "close":
                self._toggle_live_preview(show=False)
                resp["result"] = "Live Preview closed."
                resp["event"].set()
            elif action == "console":
                msgs = panel.get_console_messages()
                if not panel.current_path():
                    resp["result"] = "No page is loaded in Live Preview (use action='open' first)."
                elif not msgs:
                    resp["result"] = ("Console is clean — no JS errors, warnings or logs since "
                                      "the page loaded.")
                else:
                    lines = [f"[{m['level']}] line {m['line']}: {m['message']}" for m in msgs]
                    resp["result"] = f"Console output ({len(msgs)} entries):\n" + "\n".join(lines)
                resp["event"].set()
            elif action == "read":
                if not panel.current_path():
                    resp["result"] = "No page is loaded in Live Preview (use action='open' first)."
                    resp["event"].set()
                    return
                # toPlainText is async — resolve resp from its callback.
                def _deliver(text: str) -> None:
                    text = (text or "").strip()
                    if len(text) > 20000:
                        text = text[:20000] + "\n… [truncated at 20000 chars]"
                    resp["result"] = (f"Rendered page text of {panel.current_path()}:\n{text}"
                                      if text else
                                      "Page rendered but has no visible text (blank page — "
                                      "check action='console' for JS errors).")
                    resp["event"].set()
                panel.get_page_text(_deliver)
            else:
                resp["error"] = f"Unknown LivePreview action: {action!r}"
                resp["event"].set()
        except Exception as exc:
            log.error(f"[LivePreview] agent tool failed: {exc}")
            resp["error"] = str(exc)
            resp["event"].set()

    def _on_editor_term_moved(self):
        """When user drags the editor↔terminal splitter handle, auto-show the terminal
        if it was hidden and the user dragged it to give it non-zero size."""
        splitter = self._editor_terminal_splitter
        sizes = splitter.sizes()
        if len(sizes) >= 2 and sizes[1] > 20:
            # User dragged handle down → make terminal visible
            widget = splitter.widget(1)
            if widget and not widget.isVisible():
                widget.setVisible(True)
                widget.setMinimumHeight(120)
                self._terminal_panel_hidden = False
                # Auto-create first terminal if none exist yet
                if self._terminal_tabs.count() == 0:
                    self._new_terminal(show_panel=False)
        self._sync_splitter_handles(splitter)

    def _toggle_left_sidebar(self, show: bool = True):
        """Toggle left sidebar via splitter — show=True to expand, False to collapse."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 2:
            return
        widget = self._main_splitter.widget(0)
        if show:
            widget.setVisible(True)
            widget.setMinimumWidth(180)
            widget.setMaximumWidth(400)
            saved = getattr(self, '_left_sidebar_saved_size', 0)
            sizes[0] = saved if saved > 0 else self._left_sidebar_min_width
        else:
            self._left_sidebar_saved_size = sizes[0] if sizes[0] > 0 else self._left_sidebar_min_width
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            widget.setVisible(False)
            sizes[0] = 0
        self._main_splitter.setSizes(sizes)
        self._sync_splitter_handles(self._main_splitter)
        self._left_sidebar_hidden = not show

    def _toggle_sidebar(self):
        """Toggle left sidebar via splitter (Ctrl+B).
        
        Toggles the left sidebar panel embedded in the main splitter,
        matching the toolbar toggle behavior.
        """
        hidden = getattr(self, '_left_sidebar_hidden', False)
        self._toggle_left_sidebar(show=hidden)

    def _toggle_ai_chat_panel(self, show: bool = True):
        """Toggle AI chat panel via splitter."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 2:
            return
        widget = self._main_splitter.widget(1)
        if show:
            widget.setVisible(True)
            widget.setMinimumWidth(300)
            widget.setMaximumWidth(16777215)  # QWIDGETSIZE_MAX
            saved = getattr(self, '_chat_panel_saved_size', 0)
            sizes[1] = saved if saved > 0 else self._chat_panel_min_width
        else:
            self._chat_panel_saved_size = sizes[1] if sizes[1] > 0 else self._chat_panel_min_width
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            widget.setVisible(False)
            sizes[1] = 0
        self._main_splitter.setSizes(sizes)
        self._sync_splitter_handles(self._main_splitter)
        self._chat_panel_hidden = not show

    def _toggle_code_panel(self, show: bool = True):
        """Toggle Code Editor (Monaco/Webview) panel via splitter."""
        sizes = self._main_splitter.sizes()
        if len(sizes) < 3:
            return
        widget = self._main_splitter.widget(2)
        if show:
            widget.setVisible(True)
            widget.setMinimumWidth(500)
            widget.setMaximumWidth(16777215)  # QWIDGETSIZE_MAX
            saved = getattr(self, '_code_panel_saved_size', 0)
            sizes[2] = saved if saved > 0 else self._code_panel_min_width
        else:
            self._code_panel_saved_size = sizes[2] if sizes[2] > 0 else self._code_panel_min_width
            widget.setMinimumWidth(0)
            widget.setMaximumWidth(0)
            widget.setVisible(False)
            sizes[2] = 0
        self._main_splitter.setSizes(sizes)
        self._sync_splitter_handles(self._main_splitter)
        self._code_panel_hidden = not show

    def _toggle_review_panel(self, show: bool = True):
        """Toggle Git Review panel by switching sidebar to the git-review tab."""
        if show:
            self._sidebar._switch_panel(3)  # Git Review is index 3
            # Ensure sidebar is visible
            if getattr(self, '_left_sidebar_hidden', False):
                self._toggle_left_sidebar(show=True)
        else:
            self._sidebar._switch_panel(0)  # Back to Explorer
        self._review_panel_hidden = not show

    def _toggle_fullscreen(self):
        """Toggle full screen mode (F11)."""
        if self.isFullScreen():
            self.showNormal()
        else:
            self.showFullScreen()

    def _minimize_window(self):
        """Minimize window (Ctrl+M)."""
        self.showMinimized()

    def _zoom_window(self):
        """Zoom window (maximize/restore)."""
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def _close_window(self):
        """Close window (Ctrl+W)."""
        self.close()

    def _previous_chat(self):
        """Navigate to previous chat (TODO: assign shortcut)."""
        # TODO: Implement chat history navigation
        log.info("Previous chat shortcut pressed")

    def _next_chat(self):
        """Navigate to next chat (TODO: assign shortcut)."""
        # TODO: Implement chat history navigation
        log.info("Next chat shortcut pressed")

    def _navigate_back(self):
        """Navigate back (TODO: assign shortcut)."""
        # TODO: Implement navigation history
        log.info("Navigate back shortcut pressed")

    def _navigate_forward(self):
        """Navigate forward (TODO: assign shortcut)."""
        # TODO: Implement navigation history
        log.info("Navigate forward shortcut pressed")

    def _toggle_ai_chat(self):
        """Toggle AI chat panel visibility"""
        # In AI-first mode, the chat is always visible
        # This method can be used to focus it instead
        if hasattr(self, '_ai_chat'):
            self._ai_chat.setFocus()

    def _zoom_in(self):
        """Zoom in (Ctrl+=)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            current_zoom = editor.zoomIn()
            if current_zoom is not None:
                zoom = current_zoom + 1
                editor.setZoom(zoom)

    def _zoom_out(self):
        """Zoom out (Ctrl+-)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            current_zoom = editor.zoomIn()
            if current_zoom is not None:
                zoom = max(0, current_zoom - 1)
                editor.setZoom(zoom)

    def _zoom_reset(self):
        """Reset zoom (Ctrl+0)."""
        editor = self._editor_tabs.current_editor()
        if editor:
            editor.setZoom(0)

    def _focus_ai_chat(self):
        """Focus AI Chat input (Ctrl+Shift+A)."""
        self._ai_chat.focus_input()
        self._ai_chat.raise_()
        self._ai_chat.activateWindow()

    def _command_palette(self):
        """Show Command Palette (Ctrl+Shift+P) - REMOVED in AI-first mode."""
        # Command palette not implemented
        pass

    def _current_editor_action(self, action: str):
        """Focus-aware edit action handler (supports Editor, AI Chat, and Terminal)."""
        focused = QApplication.focusWidget()
        log.debug(f"Action {action} requested. Current focus: {focused}")

        # Map generic action strings to QWebEnginePage.WebAction enums
        web_action_map = {
            "copy": QWebEnginePage.WebAction.Copy,
            "paste": QWebEnginePage.WebAction.Paste,
            "cut": QWebEnginePage.WebAction.Cut,
            "selectAll": QWebEnginePage.WebAction.SelectAll,
            "undo": QWebEnginePage.WebAction.Undo,
            "redo": QWebEnginePage.WebAction.Redo
        }

        # Determine which "logical" component has focus
        logical_focused = None
        widget = focused
        max_depth = 10
        while widget and max_depth > 0:
            if hasattr(self, '_ai_chat') and (widget == self._ai_chat or
                (not getattr(self, '_is_native_chat', False) and hasattr(self._ai_chat, '_view') and widget == self._ai_chat._view)):
                logical_focused = "ai_chat"
                break
            
            # Check if this widget belongs to a terminal tab
            term = self._current_terminal()
            if term and (widget == term or widget == term._webview):
                logical_focused = "terminal"
                break
            
            widget = widget.parentWidget()
            max_depth -= 1

        # 1. Route to AI Chat
        if logical_focused == "ai_chat":
            if getattr(self, '_is_native_chat', False):
                # Native chat — let Qt handle copy/paste natively
                return
            if action in web_action_map:
                log.debug(f"Routing {action} to AI Chat WebEngineView")
                self._ai_chat._view.page().triggerAction(web_action_map[action])
                return

        # 2. Route to Terminal
        if logical_focused == "terminal":
            term = self._current_terminal()
            if action == "copy":
                term.copy()
                return
            elif action == "paste":
                term.paste()
                return
            elif action == "selectAll":
                term.select_all()
                return
            elif action == "cut":
                term.cut()
                return
            
            if action in web_action_map:
                term._webview.page().triggerAction(web_action_map[action])
                return

        # 3. Route to Sidebar explicitly (if focused)
        if hasattr(self, '_sidebar') and self._sidebar.is_explorer_focused():
            log.debug(f"Action {action} ignored globally: Sidebar handles it locally")
            return

        # 4. Fallback to Editor (current tab)
        editor = self._editor_tabs.current_editor()
        if editor:
            log.debug(f"Routing {action} to Code Editor")
            if action == "selectAll":
                if hasattr(editor, "selectAll"): editor.selectAll()
                elif hasattr(editor, "select_all"): editor.select_all()
                return
            if hasattr(editor, action):
                getattr(editor, action)()

    # ------------------------------------------------------------------
    # VS Code Style Keyboard Shortcuts
    # ------------------------------------------------------------------
    def _show_find(self):
        """Show Find widget — Monaco native (Ctrl+F) or legacy dialog."""
        # Monaco handles Ctrl+F natively — this is for menu-click fallback
        fp = self._webview_panel.get_active_file()
        if fp:
            self._webview_panel._safe_run_js("if(typeof editor!=='undefined'&&editor)editor.trigger('keyboard','actions.find',null);")
            return
        # Legacy editor fallback
        editor = self._editor_tabs.current_editor()
        if editor:
            selected = editor.get_selected_text()
            if selected:
                self._find_replace_dialog.set_find_text(selected)
            self._find_replace_dialog.show_find_only()
            self._find_replace_dialog.show()
            self._find_replace_dialog.raise_()
            self._find_replace_dialog.activateWindow()

    def _show_find_replace(self):
        """Show Find & Replace widget — Monaco native (Ctrl+H) or legacy dialog."""
        # Monaco handles Ctrl+H natively — this is for menu-click fallback
        fp = self._webview_panel.get_active_file()
        if fp:
            self._webview_panel._safe_run_js("if(typeof editor!=='undefined'&&editor)editor.trigger('keyboard','editor.action.startFindReplaceAction',null);")
            return
        # Legacy editor fallback
        editor = self._editor_tabs.current_editor()
        if editor:
            selected = editor.get_selected_text()
            if selected:
                self._find_replace_dialog.set_find_text(selected)
            self._find_replace_dialog.show_find_replace()
            self._find_replace_dialog.show()
            self._find_replace_dialog.raise_()
            self._find_replace_dialog.activateWindow()

    def _rename_file(self):
        """Rename file (F2)."""
        # Check if left sidebar explorer is focused
        if self._sidebar.is_explorer_focused():
            if self._sidebar.rename_selected_item():
                return
            return

        # Otherwise rename the currently open file in editor
        current_file = self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
        if not current_file:
            return
        
        from PyQt6.QtWidgets import QInputDialog
        from pathlib import Path
        
        old_name = Path(current_file).name
        new_name, ok = QInputDialog.getText(
            self, 
            "Rename File", 
            f"New name for '{old_name}':",
            text=old_name
        )
        
        if ok and new_name and new_name != old_name:
            try:
                old_path = Path(current_file)
                new_path = old_path.parent / new_name
                
                # Rename file on disk
                old_path.rename(new_path)
                
                # Close current tab and open renamed file
                index = self._editor_tabs.currentIndex()
                self._editor_tabs.removeTab(index)
                self._open_file(str(new_path))

                # Refresh the sidebar file tree so the rename appears immediately
                try:
                    if getattr(self, '_sidebar', None) and getattr(self._sidebar, '_bridge', None):
                        self._sidebar._bridge.refreshFileTree()
                except Exception as _re:
                    log.warning(f"[RENAME] Sidebar tree refresh failed: {_re}")

                self.statusBar().showMessage(f"Renamed to {new_name}", 3000)
            except Exception as e:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.critical(self, "Rename Failed", f"Could not rename file: {e}")

    def _rename_path(self, path: str) -> bool:
        """Rename a file or folder by path (used by Explore panel)."""
        try:
            name, ok = QInputDialog.getText(self, "Rename", "New name:", text=Path(path).name)
            if not ok or not name or name == Path(path).name:
                return False

            new_path_obj = Path(path).parent / name
            
            # Prevent overwriting existing files
            if new_path_obj.exists() and new_path_obj.resolve() != Path(path).resolve():
                QMessageBox.warning(
                    self, 
                    "Rename Failed", 
                    f"A file or folder with the name '{name}' already exists.\n\nPlease choose a different name."
                )
                return False

            new_path = str(new_path_obj)
            Path(path).rename(new_path)
            log.info(f"[RENAME] Renamed: {Path(path).name} -> {name}")
            return True
        except Exception as e:
            QMessageBox.critical(self, "Rename Failed", f"Could not rename: {e}")
            return False

    def _on_sidebar_file_renamed(self, old_path: str, new_path: str):
        old_norm = os.path.normpath(old_path)
        new_norm = os.path.normpath(new_path)
        updated = False

        for idx, fp in list(self._editor_tabs._files.items()):
            if os.path.normcase(fp) != os.path.normcase(old_norm):
                continue

            self._editor_tabs._files[idx] = new_norm
            name = Path(new_norm).name
            if fp in self._editor_tabs._modified:
                self._editor_tabs._modified.discard(fp)
                self._editor_tabs._modified.add(new_norm)
                self._editor_tabs.setTabText(idx, f"ƒ-? {name}")
                self._editor_tabs.setTabToolTip(idx, f"{new_norm} (Modified)")
            else:
                self._editor_tabs.setTabText(idx, name)
                self._editor_tabs.setTabToolTip(idx, new_norm)

            self._editor_tabs._set_tab_icon(idx, new_norm)
            updated = True

            if idx == self._editor_tabs.currentIndex():
                self._update_status_file(new_norm)
                if hasattr(self, '_ai_agent'):
                    self._ai_agent.set_active_file(new_norm)

        if updated:
            log.info(f"Updated open tabs for rename: {old_norm} -> {new_norm}")
    
    def close_editor_tabs_for_path(self, path: str):
        """PUBLIC method — close all editor tabs matching deleted file/folder path.
        
        Called both by _on_sidebar_file_deleted (via signal) AND directly
        by SidebarWidget._handle_native_modal (bypassing signals for reliability).
        
        Uses TWO matching strategies (either succeeds → tab closed):
          1. _files dict (index→filepath map)
          2. QTabWidget tabToolTip (every tab stores full filepath)
        """
        import os
        from pathlib import Path
        
        _log.info(f"close_editor_tabs_for_path CALLED with path={path!r}")
        
        norm_path = os.path.normpath(path)
        norm_path_case = os.path.normcase(norm_path)
        
        _log.info(f"close_editor_tabs_for_path norm_path_case={norm_path_case!r}")
        
        to_close = {}  # idx -> filepath (use dict to dedup across strategies)
        
        # Strategy 1: _files dict
        for idx, fp in list(self._editor_tabs._files.items()):
            fp_norm = os.path.normcase(os.path.normpath(fp))
            if fp_norm == norm_path_case or fp_norm.startswith(norm_path_case + os.sep):
                to_close[idx] = fp
                _log.info(f"close_editor_tabs_for_path STRATEGY1 match: idx={idx} fp={fp!r}")
        
        # Strategy 2: QTabWidget tabToolTip (independent fallback)
        for idx in range(self._editor_tabs.count()):
            if idx in to_close:
                continue
            tooltip = self._editor_tabs.tabToolTip(idx)
            if tooltip:
                tt_norm = os.path.normcase(os.path.normpath(tooltip))
                if tt_norm == norm_path_case or tt_norm.startswith(norm_path_case + os.sep):
                    to_close[idx] = tooltip
                    _log.info(f"close_editor_tabs_for_path STRATEGY2 match: idx={idx} tooltip={tooltip!r}")
        
        # Strategy 3: webview panel _open_files dict (MOST IMPORTANT - this is where webview-editor tabs live)
        if hasattr(self, '_webview_panel') and self._webview_panel:
            webview_open = dict(self._webview_panel._open_files)  # snapshot
            for wv_path in webview_open:
                wv_norm = os.path.normcase(os.path.normpath(wv_path))
                if wv_norm == norm_path_case or wv_norm.startswith(norm_path_case + os.sep):
                    _log.info(f"close_editor_tabs_for_path STRATEGY3 match: webview path={wv_path!r}")
                    self._webview_panel.close_file(wv_path)
        
        if not to_close:
            _log.info(f"close_editor_tabs_for_path NO matches found for {path!r}")
        
        for idx, fp in sorted(to_close.items(), key=lambda x: x[0], reverse=True):
            self._editor_tabs._modified.discard(fp)
            self._editor_tabs._files.pop(idx, None)
            widget = self._editor_tabs.widget(idx)
            if widget is not None:
                self._editor_tabs.removeTab(idx)
                widget.deleteLater()
            _log.info(f"close_editor_tabs_for_path CLOSED tab idx={idx} fp={fp!r}")
        
        if to_close:
            remaining = {}
            for old_idx, fp in sorted(self._editor_tabs._files.items()):
                new_idx = sum(1 for ci in to_close if ci < old_idx)
                remaining[old_idx - new_idx] = fp
            self._editor_tabs._files = remaining
        
        self._sidebar.refresh()
        
        if hasattr(self, '_ai_agent') and self._ai_agent:
            project_root = str(self._project_manager.root) if self._project_manager.root else None
            if project_root:
                self._ai_agent.set_project_root(project_root)
        
        log.info(f"File deleted: {path}")
    
    def _on_sidebar_file_deleted(self, path: str):
        """Signal handler for sidebar file deletion. Delegates to public method."""
        self.close_editor_tabs_for_path(path)

    def _on_agent_files_deleted(self, paths: list):
        """The AI moved these files to the Recycle Bin via a shell command —
        close any editor tabs showing them so they don't linger in editor.html."""
        for p in (paths or []):
            try:
                self.close_editor_tabs_for_path(p)
                # Auto-commit deletion if enabled
                if hasattr(self, '_git_manager') and self._git_manager:
                    try:
                        self._git_manager.auto_commit_file(p, "deleted")
                    except Exception:
                        pass
            except Exception as e:
                log.warning(f"[AGENT] Failed to close tab for deleted file {p}: {e}")
    
    def _on_file_deleted_for_undo(self, original_path: str):
        """Track file deletion for undo functionality."""
        log.debug(f"Undo tracking: File moved to trash: {original_path}")
    
    def _on_file_restored_for_redo(self, restored_path: str):
        """Track file restoration for redo functionality."""
        log.debug(f"Redo tracking: File restored: {restored_path}")
        # Refresh sidebar to show restored file
        QTimer.singleShot(100, self._sidebar.refresh)
    
    def _undo(self):
        """Handle undo - prioritize editor undo, then file restore."""
        # Try editor undo first
        editor = self._editor_tabs.current_editor()
        if editor and editor.document().isUndoAvailable():
            editor.undo()
            return
        
        # If no editor or no undo available, try file restore
        if hasattr(self, '_file_manager') and self._file_manager.can_undo():
            restored_path = self._file_manager.undo_operation()
            if restored_path:
                log.info(f"Restored file: {restored_path}")
                self.statusBar().showMessage(f"Restored: {Path(restored_path).name}", 3000)
    
    def _redo(self):
        """Handle redo - prioritize editor redo, then file re-delete."""
        # Try editor redo first
        editor = self._editor_tabs.current_editor()
        if editor and editor.document().isRedoAvailable():
            editor.redo()
            return
        
        if hasattr(self, '_file_manager') and self._file_manager.can_redo():
            deleted_path = self._file_manager.redo_operation()
            if deleted_path:
                log.info(f"Re-deleted file: {deleted_path}")
                self.statusBar().showMessage(f"Deleted: {Path(deleted_path).name}", 3000)

    def _go_to_line(self):
        """Go to line — Monaco native (Ctrl+G) or legacy dialog."""
        # Monaco handles Ctrl+G natively — this is for menu-click fallback
        fp = self._webview_panel.get_active_file()
        if fp:
            self._webview_panel._safe_run_js("if(typeof editor!=='undefined'&&editor)editor.trigger('keyboard','editor.action.gotoLine',null);")
            return
        # Legacy editor fallback
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        from PyQt6.QtWidgets import QInputDialog
        line, ok = QInputDialog.getInt(self, "Go to Line", "Line number:", min=1, max=editor.blockCount())
        if ok:
            cursor = editor.textCursor()
            cursor.movePosition(cursor.MoveOperation.Start)
            for _ in range(line - 1):
                cursor.movePosition(cursor.MoveOperation.Down)
            editor.setTextCursor(cursor)
            editor.setFocus()

    def _format_code(self):
        """Format current file (Shift+Alt+F).

        Routes to Monaco's editor.action.formatDocument when the webview
        has an active file; falls back to the legacy editor's formatter.
        """
        # Prefer Monaco webview
        fp = self._webview_panel.get_active_file() if hasattr(self, '_webview_panel') else None
        if fp:
            self._webview_panel._safe_run_js(
                "if(typeof editor!=='undefined'&&editor)"
                "editor.getAction('editor.action.formatDocument').run();"
            )
            return

        # Legacy editor fallback
        editor = self._editor_tabs.current_editor()
        if editor and hasattr(editor, '_format_current_code'):
            editor._format_current_code()

    def _quick_open(self):
        """Quick open file (Ctrl+P) - Opens file dialog."""
        filepath, _ = QFileDialog.getOpenFileName(
            self, "Open File", 
            str(self._project_manager.root) if self._project_manager.root else "",
            "All Files (*.*)"
        )
        if filepath:
            self._open_file(filepath)

    def _go_to_symbol(self):
        """Go to symbol in file (Ctrl+Shift+O)."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        text = editor.get_all_text()
        # Find function/class definitions
        import re
        symbols = []
        for match in re.finditer(r'^(def|class|function|const|let|var)\s+(\w+)', text, re.MULTILINE):
            line_num = text[:match.start()].count('\n') + 1
            symbols.append(f"{match.group(1)} {match.group(2)} (line {line_num})")
        
        if symbols:
            self._ai_chat.add_system_message("📍 Symbols in file:\n" + "\n".join(symbols[:20]))
        else:
            self._ai_chat.add_system_message("No symbols found in current file.")

    def _close_current_tab(self):
        """Close current tab (Ctrl+W) — works for both webview and legacy tabs."""
        # Close webview tab first if one is active
        fp = self._webview_panel.get_active_file()
        if fp:
            self._webview_panel.close_file(fp)
            return
        # Fall back to legacy editor tabs
        if hasattr(self, '_editor_tabs'):
            self._editor_tabs.close_current_tab()

    def _close_all_tabs(self):
        """Close all tabs (Ctrl+Shift+W) — works for both webview and legacy tabs."""
        # Close all webview tabs
        if hasattr(self, '_webview_panel'):
            self._webview_panel.close_all_files()
        # Close all legacy editor tabs
        if hasattr(self, '_editor_tabs'):
            self._editor_tabs.close_all_tabs()

    def _next_tab(self):
        """Go to next tab (Ctrl+Tab) — webview or legacy."""
        # Try webview tab navigation first
        if hasattr(self, '_webview_panel'):
            self._webview_panel.next_tab()
            return
        # Legacy tabs
        current = self._editor_tabs.currentIndex()
        count = self._editor_tabs.count()
        if count > 0:
            self._editor_tabs.setCurrentIndex((current + 1) % count)

    def _prev_tab(self):
        """Go to previous tab (Ctrl+Shift+Tab) — webview or legacy."""
        # Try webview tab navigation first
        if hasattr(self, '_webview_panel'):
            self._webview_panel.prev_tab()
            return
        # Legacy tabs
        current = self._editor_tabs.currentIndex()
        count = self._editor_tabs.count()
        if count > 0:
            self._editor_tabs.setCurrentIndex((current - 1) % count)

    def _show_shortcuts_help(self):
        """Show keyboard shortcuts help dialog."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QScrollArea
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Keyboard Shortcuts Reference")
        dialog.setMinimumSize(700, 500)
        
        layout = QVBoxLayout(dialog)
        
        shortcuts_html = """
        <html>
        <head>
            <style>
                body { font-family: 'Segoe UI', sans-serif; background: #1e1e1e; color: #f5f5f5; }
                h2 { color: #3b82f6; margin-top: 20px; }
                table { width: 100%; border-collapse: collapse; margin: 10px 0; }
                th, td { padding: 10px; text-align: left; border-bottom: 1px solid #3a3a3a; }
                th { background: #2d2d2d; color: #3b82f6; font-weight: 600; }
                tr:hover { background: #2a2a2a; }
                kbd { 
                    background: #2d2d2d; 
                    border: 1px solid #3a3a3a; 
                    border-radius: 4px; 
                    padding: 2px 6px; 
                    font-family: 'Consolas', monospace;
                    color: #3b82f6;
                }
            </style>
        </head>
        <body>
            <h2>📝 Editing</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Tab</kbd></td><td>Indent (inserts 4 spaces)</td></tr>
                <tr><td><kbd>Shift</kbd>+<kbd>Tab</kbd></td><td>Outdent selected lines</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Z</kbd></td><td>Undo (current file only)</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Y</kbd></td><td>Redo (current file only)</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>A</kbd></td><td>Select All</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>C</kbd></td><td>Copy</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>X</kbd></td><td>Cut</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>V</kbd></td><td>Paste</td></tr>
            </table>
            
            <h2>🔍 Find & Replace</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>F</kbd></td><td>Find</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>H</kbd></td><td>Find and Replace</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>F</kbd></td><td>Find in Files</td></tr>
                <tr><td><kbd>F3</kbd></td><td>Find Next</td></tr>
                <tr><td><kbd>Shift</kbd>+<kbd>F3</kbd></td><td>Find Previous</td></tr>
            </table>
            
            <h2>📑 File & Tab Navigation</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Tab</kbd></td><td>Next Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>Tab</kbd></td><td>Previous Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>W</kbd></td><td>Close Current Tab</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>W</kbd></td><td>Close All Tabs</td></tr>
                <tr><td><kbd>F2</kbd></td><td>Rename File</td></tr>
            </table>
            
            <h2>🚀 Quick Open</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>P</kbd></td><td>Quick Open File</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>O</kbd></td><td>Go to Symbol</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>G</kbd></td><td>Go to Line</td></tr>
            </table>
            
            <h2>🎨 View & Tools</h2>
            <table>
                <tr><th>Shortcut</th><th>Action</th></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>B</kbd></td><td>Toggle Sidebar</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>`</kbd></td><td>Toggle Terminal</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>=</kbd></td><td>Zoom In</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>-</kbd></td><td>Zoom Out</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>0</kbd></td><td>Reset Zoom</td></tr>
                <tr><td><kbd>Ctrl</kbd>+<kbd>Shift</kbd>+<kbd>P</kbd></td><td>Command Palette</td></tr>
            </table>
        </body>
        </html>
        """
        
        label = QLabel(shortcuts_html)
        label.setWordWrap(True)
        layout.addWidget(label)
        
        dialog.exec()

    # ------------------------------------------------------------------
    # Find/Replace Handlers
    # ------------------------------------------------------------------
    def _on_find_requested(self, text: str, options: dict):
        """Handle find request from dialog."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        
        # Search options
        case_sensitive = options.get('case_sensitive', False)
        whole_word = options.get('whole_word', False)
        use_regex = options.get('use_regex', False)
        wrap_around = options.get('wrap_around', True)
        search_forward = options.get('forward', True)
        
        # Build search flags
        flags = 0
        if case_sensitive:
            flags |= 0x00010  # QTextDocument.FindFlag.FindCaseSensitively
        
        # Perform search
        find_flags = QTextDocument.FindFlag(flags)
        
        if search_forward:
            found = editor.find(text, find_flags)
        else:
            found = editor.find(text, find_flags | QTextDocument.FindFlag.FindBackward)
        
        if not found and wrap_around:
            # Wrap around
            cursor.movePosition(cursor.MoveOperation.Start if search_forward else cursor.MoveOperation.End)
            editor.setTextCursor(cursor)
            if search_forward:
                found = editor.find(text, find_flags)
            else:
                found = editor.find(text, find_flags | QTextDocument.FindFlag.FindBackward)
    
    def _on_replace_requested(self, find_text: str, replace_text: str, options: dict):
        """Handle replace request from dialog."""
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        cursor = editor.textCursor()
        
        # Check if there's selected text matching find_text
        selected = cursor.selectedText()
        if selected == find_text:
            cursor.insertText(replace_text)
        
        # Find next
        self._on_find_requested(find_text, options)
    
    def _on_replace_all_requested(self, find_text: str, replace_text: str, options: dict):
        """Handle replace all request from dialog."""
        import re
        editor = self._editor_tabs.current_editor()
        if not editor:
            return
        
        document = editor.document()
        cursor = editor.textCursor()
        
        count = 0
        case_sensitive = options.get('case_sensitive', False)
        whole_word = options.get('whole_word', False)
        use_regex = options.get('use_regex', False)
        
        if use_regex:
            # Regex replace all
            try:
                flags = 0 if case_sensitive else re.IGNORECASE
                pattern = re.compile(find_text, flags)
                content = editor.toPlainText()
                new_content, count = pattern.subn(replace_text, content)
                
                if count > 0:
                    cursor.beginEditBlock()
                    cursor.select(cursor.SelectionType.Document)
                    cursor.insertText(new_content)
                    cursor.endEditBlock()
                    
            except re.error as e:
                from PyQt6.QtWidgets import QMessageBox
                QMessageBox.warning(self, "Regex Error", f"Invalid regular expression: {e}")
                return
        else:
            # Simple replace all
            cursor.beginEditBlock()
            
            # Save original position
            original_position = cursor.position()
            
            # Move to start of document
            cursor.setPosition(0)
            
            # Find and replace all occurrences
            while True:
                found = document.find(find_text, cursor, 
                                    QTextDocument.FindFlag.FindCaseSensitively if case_sensitive else QTextDocument.FindFlag(0))
                
                if found.isNull():
                    break
                
                # Check whole word if needed
                if whole_word:
                    # Verify it's a whole word
                    start = found.selectionStart()
                    end = found.selectionEnd()
                    text = editor.toPlainText()
                    
                    # Check character before
                    if start > 0 and (text[start-1].isalnum() or text[start-1] == '_'):
                        cursor.setPosition(end)
                        continue
                    
                    # Check character after
                    if end < len(text) and (text[end].isalnum() or text[end] == '_'):
                        cursor.setPosition(end)
                        continue
                
                # Replace
                found.insertText(replace_text)
                count += 1
            
            cursor.endEditBlock()
        
        from PyQt6.QtWidgets import QMessageBox
        if count > 0:
            QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s).")
        else:
            QMessageBox.information(self, "Replace All", f"No occurrences of '{find_text}' found.")
        
        # Restore cursor position
        cursor.setPosition(original_position)
        editor.setTextCursor(cursor)

    # ------------------------------------------------------------------
    # AI Actions
    # ------------------------------------------------------------------
    def _on_ai_run_command(self, command: str):
        """Execute command from AI in active terminal."""
        self._terminal_tabs.setVisible(True)
        term = self._current_terminal()
        if not term:
            term = self._new_terminal()
        term.setFocus()
        term.execute_command(command)
        # Ensure terminal is visible
        self._terminal_tabs.setCurrentIndex(self._terminal_tabs.indexOf(term))

    def _on_directory_contents_for_tree(self, path: str, contents: str):
        """Handle directory contents for project tree card display."""
        self._ai_chat.emit_directory_tree(path, contents)

    def _on_image_pasted(self, base64_data: str, index: int):
        """Handle pasted image — inject into agent bridge for vision-capable models."""
        try:
            if hasattr(self, '_ai_agent') and self._ai_agent:
                # Inject image into agent's vision history
                if hasattr(self._ai_agent, 'inject_image'):
                    self._ai_agent.inject_image(base64_data, f"Image {index}")
                    log.info(f"[Image] Injected image {index} into agent vision")
                else:
                    log.info(f"[Image] Agent does not support image injection")
        except Exception as e:
            log.warning(f"[Image] Failed to inject image: {e}")

    def _on_ai_chat_message(self, message: str):
        """Handle user message from AI chat with project context."""
        # Inject any attached images into agent vision.
        # Images are stored on the InputArea (input_area), not the ChatPanel itself.
        # _emit_send() saves them to input_area._pending_send_images and _on_send()
        # mirrors them to input_area._pending_images. Read whichever is populated so
        # the bridge's chat() can run the Mistral OCR fallback before the selected model.
        _input_area = getattr(self._ai_chat, 'input_area', None)
        _imgs = []
        if _input_area is not None:
            _imgs = (getattr(_input_area, '_pending_images', None)
                     or getattr(_input_area, '_pending_send_images', None)
                     or [])
        # Legacy/fallback: some paths stored directly on the panel.
        if not _imgs:
            _imgs = getattr(self._ai_chat, '_pending_images', None) or []
        if _imgs:
            for img in _imgs:
                try:
                    self._on_image_pasted(img["data"], img["index"])
                except Exception as _img_e:
                    log.warning(f"[Image] Failed to inject attached image: {_img_e}")
            # Clear every place images may have been stashed so they fire once.
            if _input_area is not None:
                _input_area._pending_images = []
                _input_area._pending_send_images = []
            if hasattr(self._ai_chat, '_pending_images'):
                self._ai_chat._pending_images = []
        # TODO: POINTS SYSTEM - Disabled for development
        # Will be enabled when connecting to https://logic-practice.com backend
        # Points tracking will be handled via API authentication with production server
        """
        # Check points balance before processing
        try:
            points_mgr = get_points_manager()
            perf_mode = self._get_current_performance_mode()
            estimated_tokens = len(message) // 4  # Rough estimate: 4 chars per token
            
            # Check if user can afford this request
            if not points_mgr.can_afford(estimated_tokens, perf_mode):
                cost = points_mgr.estimate_cost(estimated_tokens, perf_mode)
                balance = points_mgr.get_balance()
                log.warning(
                    f"[MainWindow] Insufficient points: need {cost:,}, have {balance:,}. "
                    f"Mode: {perf_mode}, estimated tokens: {estimated_tokens}"
                )
                # Send error to UI
                self._ai_chat.on_error(
                    f"Insufficient points. This request needs ~{cost:,} points but you have {balance:,} points.\n\n"
                    f"Please purchase more points or switch to Efficient mode (0.3x cost)."
                )
                return
            
            log.info(
                f"[MainWindow] Points check passed: balance={points_mgr.get_balance():,}, "
                f"mode={perf_mode}, estimated_cost={points_mgr.estimate_cost(estimated_tokens, perf_mode):,}"
            )
        except Exception as e:
            log.warning(f"[MainWindow] Points check failed: {e} - proceeding anyway")
        """
        
        # Build context for ALL messages - let agent_bridge decide
        # what to do with simple vs complex queries
        context = []

        # 1. Project Root Info
        if self._project_manager.root:
            context.append(f"Project path: {self._project_manager.root}")

        # 2. Current File Context
        editor = self._editor_tabs.current_editor()
        if editor:
            fp = self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
            if fp:
                name = Path(fp).name
                content = editor.get_all_text()
                # Limit context size
                if len(content) > 5000:
                    content = content[:5000] + "... (truncated)"
                context.append(f"Current file ({name}):\n```\n{content}\n```")

        full_context = "\n\n".join(context)
        
        # NOTE: _ai_integration was removed - using _ai_agent directly
        # session_id = getattr(self._ai_chat, '_current_conversation_id', 'default-session')
        # self._ai_integration.set_session(...)  # Removed - not needed
        
        # Start AI processing immediately for responsiveness
        # But wait - we need to see if enhancement layer (Intent/Routing) is fast enough
        # Optimization: start AI chat, and if enhancement layer finds a special tool/route, 
        # we can inject that context later or stop/restart. 
        # For now, let's just fix the DUPLICATE problem by only calling it ONCE.
        
        conv_id = getattr(self._ai_chat, '_current_conversation_id', None) if hasattr(self._ai_chat, '_current_conversation_id') else None
        
        if conv_id and not self._title_generator.get_cached_title(conv_id):
            title = self._ai_agent.generate_chat_title(message, conv_id)
            if title:
                log.info(f"Generated chat title: {title}")
                if hasattr(self._ai_chat, 'update_conversation_title'):
                    self._ai_chat.update_conversation_title(conv_id, title)
        
        if conv_id:
            try:
                sessions = self._session_db.list_sessions(self._project_manager.root, limit=1)
                if not sessions:
                    title = self._title_generator.get_cached_title(conv_id) or "New Chat"
                    self._session_db.create_session(
                        conv_id, 
                        title, 
                        self._project_manager.root or "", 
                        self._ai_agent._settings.get("ai", "model", default="mistral-large-latest")
                    )
                import uuid
                self._session_db.add_message(
                    conv_id,
                    str(uuid.uuid4()),
                    "user",
                    message,
                    len(message) // 4
                )
            except Exception as e:
                log.debug(f"Could not store message in database: {e}")
        
        # FIX: Initialize enhancement_result to None (not yet implemented in async flow)
        enhancement_result = None
        if enhancement_result and enhancement_result.get("intent"):
            if enhancement_result.get("testing_decision") and \
               enhancement_result["testing_decision"].decision == 'write_tests':
                log.info("Using chat_with_testing with testing workflow")
                self._ai_agent.chat_with_testing(
                    message,
                    code_changes=[],
                    code_context=full_context
                )
            else:
                log.info("Using chat_with_enhancement with intent classification data")
                self._ai_agent.chat_with_enhancement(
                    message, 
                    intent=enhancement_result["intent"],
                    route=enhancement_result["route"],
                    tools=enhancement_result["tools"],
                    code_context=full_context
                )
        else:
            # Suppress sidebar file-tree refreshes during AI work —
            # os.walk on every tool call blocks the main thread.
            if hasattr(self, '_sidebar') and self._sidebar:
                self._sidebar._ai_active = True
            self._ai_agent.chat(message, full_context)

    def _on_model_changed(self, model_id: str, perf: str = None, cost: str = None):
        """Handle model selection change from AI chat."""
        log.info(f"[MainWindow] DEBUG: model_id='{model_id}'")
        
        # Check if this is a performance mode (not an actual model)
        # These are kept for backward compatibility with settings but
        # are no longer shown in the dropdown. They silently map to
        # the default Auto mode which uses smart routing.
        performance_modes = ["efficient", "auto", "performance", "ultimate"]
        if model_id.lower() in performance_modes:
            try:
                from src.config.settings import get_settings
                settings = get_settings()
                settings.set("ai", "performance_mode", model_id.lower())
                # All modes now default to Auto smart-routing (DeepSeek)
                _provider, _model = "deepseek", "deepseek-chat"
                settings.set("ai", "model_id", _model)
                settings.set("ai", "model", _model)
                settings.set("ai", "provider", _provider)
                settings.set("ai", "token_multiplier", "1.0")
                log.info(
                    "[MainWindow] Performance mode '%s' → %s/%s (auto-routing)",
                    model_id, _provider, _model,
                )
                self._ai_agent.update_settings(provider=_provider, model_id=_model)
                return
            except Exception as e:
                log.error(f"[MainWindow] Failed to save performance mode: {e}")
                return
        
        # This is an actual model ID - determine provider from model_id
        try:
            if model_id.startswith("mistral-") or model_id.startswith("codestral-"):
                provider = "mistral"
            elif model_id.startswith("deepseek"):
                provider = "deepseek"
            elif model_id.startswith("mimo-"):
                provider = "mimo"
            elif model_id.startswith(("gpt-", "o1", "o3", "codex")):
                if model_id.startswith(("gpt-5", "o1", "o3", "codex")):
                    provider = "openai_responses"
                else:
                    provider = "openai"
            elif "/" in model_id:
                provider = "openrouter"
            elif model_id.startswith("qwen") or model_id.startswith("qwq"):
                provider = "alibaba"
            else:
                provider = "deepseek"

            # Save model_id and provider to settings
            try:
                from src.config.settings import get_settings
                settings = get_settings()
                settings.set("ai", "model_id", model_id)
                settings.set("ai", "model", model_id)
                settings.set("ai", "provider", provider)
                settings.set("ai", "token_multiplier", "1.0")
                log.info(f"[MainWindow] Saved model: {model_id} (provider: {provider})")
            except Exception as e:
                log.warning(f"[MainWindow] Failed to save model to settings: {e}")

            log.info(f"[MainWindow] Model changed to: {model_id} (provider: {provider})")
            self._ai_agent.update_settings(provider=provider, model_id=model_id)
        except Exception as e:
            log.error(f"[MainWindow] Model change failed for '{model_id}': {e}")

    def _on_settings_model_changed(self, model_id: str, model_label: str):
        """Handle model change from Settings page → sync chat panel button."""
        try:
            if hasattr(self, '_ai_chat') and self._ai_chat:
                input_area = getattr(self._ai_chat, 'input_area', None)
                if input_area:
                    input_area.model = model_id
                    input_area.model_label = model_label
                    input_area.model_btn.setText(model_label)
                    log.info(f"[MainWindow] Settings synced model button to: {model_label}")
        except Exception as e:
            log.warning(f"[MainWindow] _on_settings_model_changed failed: {e}")

    def _on_mode_changed(self, mode: str):
        """Handle mode change (Agent/Ask/Plan) → update agent tool access."""
        try:
            if hasattr(self, '_ai_agent') and self._ai_agent:
                self._ai_agent.set_interaction_mode(mode)
                log.info(f"[MainWindow] Mode changed to: {mode}")
                # Show toast in chat
                if hasattr(self, '_ai_chat') and self._ai_chat:
                    if mode == 'Ask':
                        self._ai_chat.add_system_message("🔒 Ask mode — read-only. Write/Edit/Bash tools are disabled.")
                    elif mode == 'Plan':
                        self._ai_chat.add_system_message("📋 Plan mode — planning only. Use Agent mode to execute code.")
                    else:
                        self._ai_chat.add_system_message("⚡ Agent mode — full access to all tools.")
        except Exception as e:
            log.warning(f"[MainWindow] _on_mode_changed failed: {e}")

    def _on_ai_stop_requested(self):
        """Handle stop request from AI (via web bridge)."""
        self._ai_agent.stop()
    
    def _on_toggle_autogen(self):
        """Toggle AutoGen multi-agent mode.

        When enabled AND the currently-selected model is Mistral/Codestral,
        switches the performance mode to "performance" which activates the
        CoordinationEngine multi-agent path.
        
        For non-Mistral models (DeepSeek, MiMo, etc.), the toggle state is
        stored in the bridge but does NOT change performance_mode — the
        model routing in ai_chat.py already bypasses performance-mode for
        non-Mistral models, so forcing it would have no effect.
        """
        # Get current status
        status = self._ai_agent.get_autogen_status()
        current_enabled = status.get('enabled', False)

        # Toggle bridge state
        new_state = not current_enabled
        self._ai_agent.enable_autogen(new_state)

        # ── Sync performance mode with autogen toggle ────────────
        # Only set performance_mode when using a Mistral model, because
        # the performance-mode pipeline hardcodes Mistral as the provider.
        # For non-Mistral models (DeepSeek, MiMo), the routing in ai_chat.py
        # bypasses performance-mode regardless, so we leave it unchanged.
        try:
            from src.config.settings import get_settings
            settings = get_settings()
            user_model = (settings.get("ai", "model_id", default="") or "").strip()
            # Always persist the autogen flag so ai_chat.py can detect it
            # even when performance_mode is left unchanged for non-Mistral models.
            settings.set("ai", "autogen_enabled", new_state)
            _is_mistral = user_model.startswith("mistral") or user_model.startswith("codestral") or not user_model
            if _is_mistral:
                _new_perf_mode = "performance" if new_state else "auto"
                settings.set("ai", "performance_mode", _new_perf_mode)
                log.info(
                    f"[AutoGen] Performance mode → {_new_perf_mode} "
                    f"({'multi-agent' if new_state else 'single-agent'}) "
                    f"(model={user_model or 'default'})"
                )
            else:
                log.info(
                    f"[AutoGen] Toggle {'' if new_state else 'DIS'}ABLED "
                    f"— performance_mode left unchanged (non-Mistral model: {user_model})"
                )
        except Exception as _e:
            log.warning(f"[AutoGen] Failed to sync performance mode: {_e}")

        log.info(f"AutoGen {'enabled' if new_state else 'disabled'} via UI toggle")
    
    def _on_generate_plan(self):
        log.info("MainWindow: Automated plan generation triggered")
        self._ai_agent.chat("__GENERATE_PLAN__")

    def _ai_action(self, action: str):
        """Execute an AI tool action on the currently active file."""
        # Normalize sidebar action names to internal keys
        action_map = {
            'explain-code':    'explain',
            'explain':         'explain',
            'refactor':        'refactor',
            'debug':           'debug',
            'generate-tests':  'tests',
            'tests':           'tests',
            'optimize':        'optimize',
            'document':        'docstring',
            'docstring':       'docstring',
        }
        internal_action = action_map.get(action, action)

        code = ""
        language = "plaintext"

        # Prefer Monaco webview file if open
        fp = self._webview_panel.get_active_file() if hasattr(self, '_webview_panel') else None
        if fp:
            code = self._webview_panel.get_content(fp)
            from src.utils.helpers import detect_language
            language = detect_language(fp)
        else:
            # Fall back to legacy editor
            editor = self._editor_tabs.current_editor()
            if not editor:
                self._ai_chat.add_system_message("Open a file first to use AI actions.")
                return
            code = editor.get_selected_text() or editor.get_all_text()
            language = editor.language

        if not code or not code.strip():
            self._ai_chat.add_system_message("The file appears to be empty.")
            return

        analyzer = CodeAnalyzer()

        prompts = {
            "explain":   analyzer.build_explain_prompt(code, language),
            "refactor":  analyzer.build_refactor_prompt(code, language),
            "tests":     analyzer.build_test_prompt(code, language),
            "debug":     analyzer.build_debug_prompt(code, "unknown error", language),
            "optimize":  analyzer.build_optimize_prompt(code, language),
            "docstring": f"Add comprehensive docstrings to this {language} code:\n\n```{language}\n{code}\n```",
        }

        prompt = prompts.get(internal_action, f"Help me with this {language} code:\n\n```{language}\n{code}\n```")
        self._ai_chat.add_system_message(f"🤖 Running: {internal_action.title()} Code…")
        self._ai_chat._add_ai_bubble_streaming()
        # Toggle send → stop button in chat UI while AI is working
        self._ai_chat._view.page().runJavaScript("if(window.onAiActionStarted) window.onAiActionStarted();")
        self._ai_agent.chat(prompt)

    def _get_selected_code(self) -> str:
        editor = self._editor_tabs.current_editor()
        if editor:
            return editor.get_selected_text() or editor.get_all_text()
        return ""

    def _get_code_context(self) -> str:
        """Alias for _get_selected_code to satisfy WebAIChatWidget callback."""
        return self._get_selected_code()

    # ------------------------------------------------------------------
    # Project & Events
    # ------------------------------------------------------------------

    def _cleanup_old_project(self):
        """Clean up all state from the old project before opening a new one."""
        # Only cleanup if we have a previous project loaded
        if not hasattr(self, '_current_project_path') or not self._current_project_path:
            log.info("🆕 First project load - skipping cleanup")
            return
            
        log.info("🧹 Cleaning up old project state...")
        
        # 1. Close all editor tabs
        self._editor_tabs.close_all_tabs()
        log.info("   ✓ Closed all editor tabs")
        
        # 2. Clear file snapshots (diff data)
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
            log.info("   ✓ Cleared file snapshots")
        
        # 3. Clear diff cache
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
            log.info("   ✓ Cleared diff data store")
        
        # 4. Clear file tracker
        if hasattr(self, '_file_tracker'):
            if hasattr(self._file_tracker, '_edits'):
                self._file_tracker._edits.clear()
            log.info("   ✓ Cleared file edit history")
        
        # 5. Clear AI agent context
        self._ai_agent.clear_active_file()
        log.info("   ✓ Cleared AI agent active file")
        
        # 5b. CRITICAL: Clear conversation history to prevent cross-project leakage.
        # Without this, the old project's chat context (tool calls, reasoning,
        # file operations like dark.py) leaks into the new project, causing the
        # AI to hallucinate files from previous projects.
        self._ai_agent.clear_conversation()
        log.info("   ✓ Cleared AI agent conversation history")
        # Reset the guard flag so set_project_root() re-loads DB context
        # for the NEW project (instead of keeping stale old-project data).
        self._ai_agent._chat_context_restored = False
        log.info("   ✓ Reset chat context restore flag for new project")
        
        # 6. Clear codebase index
        self._codebase_index = None
        log.info("   ✓ Cleared codebase index")
        
        # 7. Prepare terminals for new project (will set CWD after cleanup)
        log.info("   ✓ Prepared terminals for new project")
        
        # 8. Todo list is session-based, no need to clear
        # Todos are tied to chat sessions, not projects
        log.info("   ✓ Skipped todo cleanup (session-based)")
        
        # 9. Clear search results
        if hasattr(self, '_search_results'):
            self._search_results.clear()
            log.info("   ✓ Cleared search results")
        
        log.info("✅ Old project cleanup complete!")
    
    def _on_project_opened(self, folder_path: str):
        log.info(f"Project opened: {folder_path}")

        # ── Guard: warn if AI is mid-turn ──
        # If the AI is actively streaming or running tool calls, switching
        # projects mid-turn can cause tool calls to execute against the
        # wrong project root. Cancel pending work first.
        #
        # BUG FIX: this checked '_is_generating' and called '.cancel()' —
        # neither exists on CortexAgentBridge (the real flag is
        # '_agentic_turn_active', the real stop method is '.stop()', which
        # calls stop_generation() and actually cancels the asyncio task).
        # getattr()'s default silently made this guard permanently inert —
        # a project switch NEVER cancelled an in-flight turn on the OLD
        # project. That turn kept running in the background, and its late
        # streamed tokens/tool results still delivered to the same shared
        # chat panel widget — now showing the NEW project — which is what
        # made the previous project's chat appear to "leak" into it after
        # the user started chatting there.
        _is_generating = getattr(getattr(self, '_ai_agent', None), '_agentic_turn_active', False)
        if _is_generating:
            log.warning("[PROJECT-SWITCH] AI is mid-turn — cancelling before switch")
            try:
                self._ai_agent.stop()
            except Exception:
                pass
            # Small delay to let cancel propagate
            from PyQt6.QtCore import QTimer as _QT
            _QT.singleShot(200, lambda: self._do_project_switch(folder_path))
            return

        self._do_project_switch(folder_path)

    def _do_project_switch(self, folder_path: str):
        """Actual project switch logic — called after AI turn is cancelled."""
        log.info(f"[PROJECT-SWITCH] Switching to: {folder_path}")
        
        # Clean up old project state BEFORE loading new one (only if switching)
        self._cleanup_old_project()
        
        # Set project root FIRST (this loads project-specific context)
        self._ai_agent.set_project_root(folder_path)
        self._current_project_path = folder_path  # Track current project
        if hasattr(self, '_git_manager'):
            repo_ok = self._git_manager.set_repository(folder_path)
            log.info(f"[GIT] Repository set on project-opened: {folder_path} (ok={repo_ok})")
        
        # Reset codebase index for new project
        self._codebase_index = None
        
        # Clear old chat and load new project's chat history
        self._ai_chat.clear_messages()
        self._setup_chat_persistence()
        
        self._ai_agent.clear_active_file()
        
        self._sidebar.set_project(folder_path)

        # Update all current terminal tabs to the new project directory
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term.set_cwd(folder_path)
                
        project_name = Path(folder_path).name
        self._sync_window_title()
        
        # Load existing chats for this project from SQLite (non-native mode only)
        chats_json = "[]"
        if not getattr(self, '_is_native_chat', False):
            try:
                chats_json = self._ai_chat.load_chats_for_project(folder_path)
            except Exception as e:
                log.warning(f"Failed to load chats for project {folder_path}: {e}")
        
        # Update project indicator in AI chat (this triggers project-specific chat loading)
        self._ai_chat.set_project_info(project_name, folder_path, chats_json)
        
        # Auto-detect and activate virtual environment
        self._check_and_activate_venv(folder_path)
        
        # Start background project context scan
        self._start_project_context_scan(folder_path)
    
    def _start_project_context_scan(self, folder_path: str):
        """Start background project scanning for instant AI awareness."""
        from src.ai.project_context import build_project_context
        
        # Show indexing status in chat
        self._ai_chat.show_indexing_status("Indexing project...")
        
        def on_context_ready(ctx):
            """Called when background scan finishes."""
            from PyQt6.QtCore import QMetaObject, Qt
            # Must update UI from main thread
            QMetaObject.invokeMethod(
                self, "_on_project_context_ready",
                Qt.ConnectionType.QueuedConnection,
            )
            # Store context in agent
            self._ai_agent.set_project_context(ctx)
        
        self._context_builder = build_project_context(folder_path, on_context_ready)

        # Start background semantic search indexing — but only once the
        # startup storm has passed. Its tree walk + embedding calls compete
        # for CPU/IO/RAM with the GUI thread, and a FIXED delay wasn't
        # enough: on a memory-constrained machine chat restore alone takes
        # ~14s, so a 4s timer landed indexing right in the middle of it
        # (measured: restore batches jumping to 5s each while indexing ran).
        # Instead, poll the stability engine and start indexing only when
        # the machine is NOT under memory/CPU pressure — or after a hard cap
        # so it still runs on a chronically-pressured machine, just well
        # after the visible startup work is done.
        _INDEX_POLL_MS = 3000
        _INDEX_MAX_WAIT_MS = 45000  # start anyway after this, pressure or not

        def _start_indexing_when_calm(waited_ms: int = 0):
            try:
                from src.core.stability_engine import get_stability_engine
                _under_pressure = get_stability_engine().should_defer()
            except Exception:
                _under_pressure = False  # can't tell → don't block indexing
            if _under_pressure and waited_ms < _INDEX_MAX_WAIT_MS:
                QTimer.singleShot(
                    _INDEX_POLL_MS,
                    lambda: _start_indexing_when_calm(waited_ms + _INDEX_POLL_MS),
                )
                return
            log.info(f"[SemanticSearch] Starting background indexing "
                     f"(waited {waited_ms}ms for calm, under_pressure={_under_pressure})")
            try:
                from src.core.semantic_search import start_background_indexing, set_indexing_progress_callback
                set_indexing_progress_callback(self._on_semantic_indexing_progress)
                start_background_indexing(folder_path)
            except Exception as e:
                log.debug(f"[SemanticSearch] Background indexing trigger skipped: {e}")
        # Initial 4s settle before the first pressure check.
        QTimer.singleShot(4000, lambda: _start_indexing_when_calm(4000))
    
    def _on_semantic_indexing_progress(self, status: str, files_indexed: int = 0, total_files: int = 0):
        """Callback for semantic search background indexing progress."""
        from PyQt6.QtCore import QMetaObject, Qt
        # Must update UI from main thread
        if status == 'indexing':
            QMetaObject.invokeMethod(
                self, "_show_semantic_indexing_status",
                Qt.ConnectionType.QueuedConnection,
            )
        elif status == 'done':
            QMetaObject.invokeMethod(
                self, "_hide_semantic_indexing_status",
                Qt.ConnectionType.QueuedConnection,
            )
    
    @pyqtSlot()
    def _show_semantic_indexing_status(self):
        """Show semantic indexing status in chat panel."""
        if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, 'show_semantic_indexing_status'):
            self._ai_chat.show_semantic_indexing_status()
    
    @pyqtSlot()
    def _hide_semantic_indexing_status(self):
        """Hide semantic indexing status in chat panel."""
        if hasattr(self, '_ai_chat') and hasattr(self._ai_chat, 'hide_semantic_indexing_status'):
            self._ai_chat.hide_semantic_indexing_status()
    
    @pyqtSlot()
    def _on_project_context_ready(self):
        """Called on main thread when project indexing finishes."""
        from src.ai.project_context import get_project_context
        ctx = get_project_context(self._ai_agent._project_root)
        if ctx:
            self._ai_chat.hide_indexing_status()
            # Show ready indicator  
            self._ai_chat.show_indexing_status(
                f"✓ Indexed {ctx.source_file_count} files ({ctx.build_time_ms:.0f}ms)",
                auto_hide=True
            )

    def _on_project_closed(self):
        """Handle project close - show welcome page and clean up."""
        log.info("Project closed - showing welcome page")

        # Save MEMORY.md before closing so context survives project switch
        try:
            from src.ai.agent_bridge import get_agent_bridge
            bridge = get_agent_bridge()
            if bridge is not None:
                bridge.save_session_to_memory()
                log.info("[PROJECT_CLOSE] MEMORY.md saved before project close")
        except Exception as _mem_exc:
            log.warning(f"[PROJECT_CLOSE] Failed to save MEMORY.md: {_mem_exc}")

        # Stop Live Server (removed in AI-first mode)
        # if self._live_server and self._live_server.is_running:
        #     self._live_server.stop()
        self._live_server = None

        # Clear all state
        self._current_project_path = None
        if hasattr(self, '_modified_files'):
            self._modified_files.clear()
        self._sync_window_title()
        self._ai_chat.clear_project_info()
        
        # Close all editor tabs
        if hasattr(self, '_editor_tabs'):
            self._editor_tabs.close_all_tabs()
        
        # Clear file snapshots and caches
        if hasattr(self, '_file_snapshots'):
            self._file_snapshots.clear()
        if hasattr(self, '_diff_data_store'):
            self._diff_data_store.clear()
        if hasattr(self, '_file_tracker') and hasattr(self._file_tracker, '_edits'):
            self._file_tracker._edits.clear()
        if hasattr(self, '_search_results'):
            self._search_results.clear()
        
        # Clear AI agent context
        self._ai_agent.clear_active_file()
        self._codebase_index = None

    def _check_and_activate_venv(self, project_path: str):
        """Check for and activate Python virtual environment."""
        import os
        
        venv_names = ["venv", ".venv", "env", ".env", "virtualenv"]
        
        for venv_name in venv_names:
            venv_path = os.path.join(project_path, venv_name)
            if os.path.exists(venv_path):
                # Activate in current terminal if available
                term = self._current_terminal()
                if term and hasattr(term, 'activate_virtual_env'):
                    term.activate_virtual_env(venv_path)

                self.statusBar().showMessage(f"Virtual environment: {venv_name}", 5000)
                break

    def _on_tab_changed(self, index: int):
        """Legacy tab changed (welcome/PDF/image tabs only)."""
        fp = self._editor_tabs._files.get(index)
        self._update_status_file(fp)
        
        # Update AI agent with active file for context injection
        if fp and hasattr(self, '_ai_agent'):
            try:
                editor = self._editor_tabs.widget(index)
                cursor_pos = None
                if hasattr(editor, 'getCursorPosition'):
                    line, col = editor.getCursorPosition()
                    cursor_pos = (line + 1, col)  # Convert to 1-indexed
                self._ai_agent.set_active_file(fp, cursor_pos)
                log.debug(f"Active file updated for AI: {fp} at {cursor_pos}")
            except Exception as e:
                log.warning(f"Could not update active file for AI: {e}")

    def _on_webview_file_changed(self, file_path: str):
        """Webview editor active file changed (tab switch)."""
        self._update_status_file(file_path if file_path else None)
        # NOTE: deliberately no splitter resizing here. Live Preview is
        # full-area by design ("browser fully display it; close it to get
        # the code back") — resizing on tab changes would shrink the
        # preview every time the agent opens a file mid-edit.
        if file_path and hasattr(self, '_ai_agent'):
            try:
                def _on_cursor(line, col):
                    self._ai_agent.set_active_file(file_path, (line, col))
                self._webview_panel.get_cursor_position(_on_cursor)
            except Exception as e:
                log.warning(f"Could not update active file for AI: {e}")

    def _on_webview_content_changed(self, file_path: str, content: str):
        """Webview editor content changed — mark file modified (VS Code-style white dot)."""
        # Race guard: if file is being saved right now, suppress re-marking.
        # The _saving_files flag is set BEFORE the async get_current_content call
        # in _save_current, preventing the debounced onContentChanged (300ms)
        # from re-adding the white dot after Ctrl+S.
        if hasattr(self, '_saving_files') and file_path in self._saving_files:
            return
        # Race guard: skip if content matches the last-saved version (debounced
        # onContentChanged can fire after Ctrl+S, re-marking a saved file).
        saved = getattr(self, '_saved_content', {}).get(file_path)
        if saved is not None and saved == content:
            return  # Content unchanged since last save — not actually modified
        # Mark the file as modified in the webview tab (triggers white dot on tab)
        if hasattr(self, '_webview_panel') and self._webview_panel._page_loaded:
            self._webview_panel.mark_modified(file_path, True)
        # Update window title with modified indicator (VS Code-style)
        if not hasattr(self, '_modified_files'):
            self._modified_files: set = set()
        self._modified_files.add(file_path)
        self._sync_window_title()

    def _on_webview_save_requested(self, file_path: str, content: str):
        """Ctrl+S from Monaco — write directly to disk (no async callback)."""
        log.info(f"[Save] _on_webview_save_requested called for {file_path!r} ({len(content) if content else 0} chars)")
        if not file_path or file_path == "untitled.py":
            return
        # PyQt QWebChannel converts JS empty string "" to Python None
        if content is None:
            content = ""
        # ── Save content to disk — even if empty ──
        # User may have intentionally cleared all content (Ctrl+A → Delete → Ctrl+S).
        # Monaco sends content directly via QWebChannel — if it's empty, the user cleared it.
        try:
            # ── Race guard: block the post-save debounced onContentChanged ──
            # A format/edit schedules a 300ms-debounced onContentChanged in
            # editor.html. If Ctrl+S fires within that window, the debounced
            # call lands AFTER this save and re-adds the white dot. Hold the
            # file in _saving_files for the debounce window (+margin) so
            # _on_webview_content_changed suppresses the stale re-mark.
            if not hasattr(self, '_saving_files'):
                self._saving_files: set = set()
            self._saving_files.add(file_path)
            # Log save for debugging empty content issue
            log.info(f"[Save] Saving {file_path} ({len(content) if content else 0} chars, content={repr(content[:50]) if content else 'None'})")
            self._file_manager.write(file_path, content)
            # Track saved content for race guard
            if not hasattr(self, '_saved_content'):
                self._saved_content: dict = {}
            self._saved_content[file_path] = content
            # Clear modified state
            self._webview_panel.mark_modified(file_path, False)
            if hasattr(self, '_modified_files'):
                self._modified_files.discard(file_path)
                self._sync_window_title()
            # Release the save guard after the 300ms debounce window passes
            QTimer.singleShot(450, lambda fp=file_path: self._saving_files.discard(fp))
            # Status bar feedback
            self._status_file.setText(f"  Saved ✓  {Path(file_path).name}")
            QTimer.singleShot(2000, lambda: self._update_status_file(file_path))
        except Exception as e:
            log.error(f"[Save] Failed to save {file_path}: {e}")

    def _on_webview_file_closed(self, file_path: str):
        """Webview editor file tab closed."""
        if hasattr(self, '_modified_files'):
            self._modified_files.discard(file_path)
            self._sync_window_title()
        if hasattr(self, '_saved_content'):
            self._saved_content.pop(file_path, None)
        self._update_status_file(None)

    def _restore_session(self):
        # Check if session restore is enabled in settings
        _settings = getattr(self, '_settings', None)
        if _settings and not _settings.get("memory", "restore_session", default=True):
            log.info("Session restore disabled in settings — skipping.")
            return

        # Restore last project — skip if it no longer exists (blank state)
        log.info("Restoring last session project...")
        restored = self._project_manager.restore_last()

        if not restored:
            # No valid project to restore — show clean blank state
            log.info("No project to restore — showing blank start state.")
            return

        # Initialize Git repository for restored project
        if hasattr(self, '_git_manager') and self._project_manager.root:
            project_path = str(self._project_manager.root)
            self._git_manager.set_repository(project_path)
            log.info(f"[GIT] Repository restored to: {project_path}")

        # Populate left sidebar file tree from restored project
        if self._project_manager.root:
            self._sidebar.set_project(str(self._project_manager.root))

        # Skip file restoration: bulk-restoring tabs from previous sessions
        # shows wrong/empty content due to single-model editor architecture.
        # Files open on demand when clicked in the sidebar.
        log.info("Session restore: skipping file restoration (open on demand).")



    @property
    def codebase_index(self):
        """Get the codebase index, creating it if needed."""
        if self._codebase_index is None:
            if self._project_manager.root:
                self._codebase_index = get_codebase_index(str(self._project_manager.root))
                self._codebase_index.index_project()
            else:
                # No project open, return a dummy index?
                raise RuntimeError("No project open for indexing")
        return self._codebase_index

    def _show_about(self):
        _app_version = QApplication.instance().applicationVersion() or "2.8.1"
        QMessageBox.about(self, "About Cortex AI IDE",
                          "<h2>Cortex AI IDE</h2>"
                          "<p>A modern AI-powered IDE built with Python and PyQt6.</p>"
                          "<p>Features: Multi-file editor · Syntax highlighting · "
                          "AI chat · File explorer · Terminal</p>"
                          f"<p><b>Version:</b> {_app_version}</p>")

    def _open_documentation(self):
        """Open Cortex documentation in browser."""
        import webbrowser
        webbrowser.open("https://github.com/cortex-ai/docs")
        log.info("Opening documentation")

    def _show_whats_new(self):
        """Show what's new dialog."""
        log.info("What's new dialog requested")

    def _check_for_updates(self):
        """Check for Cortex IDE updates from the server.
        
        Runs on startup (deferred 5s). If an update is available,
        shows the UpdateDialog. Force updates block the IDE.
        """
        try:
            from src.services.update_checker import UpdateChecker
            from src.ui.dialogs.update_dialog import UpdateDialog

            checker = UpdateChecker()
            if not checker.is_enabled():
                return

            # Connect once: signal delivery is auto-queued onto the GUI
            # thread, so the dialog is created where Qt requires it.
            if not getattr(self, '_update_signal_connected', False):
                self._update_check_ready.connect(self._show_update_dialog)
                self._update_signal_connected = True

            # Run check in background thread to avoid blocking UI
            import threading
            def _bg_check():
                try:
                    result = checker.check()
                    if result is None or not result.update_available:
                        return
                    # Show dialog on main thread (queued signal emission)
                    self._update_check_ready.emit(result)
                except Exception as e:
                    log.warning("[Update] Background check failed: %s", e)

            thread = threading.Thread(target=_bg_check, daemon=True)
            thread.start()
        except Exception as e:
            log.warning("[Update] Update checker setup failed: %s", e)

    def _show_update_dialog(self, result):
        """Show the update notification dialog."""
        from src.ui.dialogs.update_dialog import UpdateDialog
        import subprocess
        import os

        dlg = UpdateDialog(
            latest_version=result.latest_version,
            current_version=result.current_version,
            force=result.force_update,
            download_url=result.download_url,
            file_size=result.file_size,
            release_notes=result.release_notes,
            parent=self,
        )

        def _on_install(path):
            """Launch the downloaded installer."""
            dlg.close()
            if path and os.path.exists(path):
                try:
                    subprocess.Popen([path], shell=True)
                    log.info("[Update] Installer launched: %s", path)
                    # If force update, quit to let installer replace files
                    if result.force_update:
                        QApplication.instance().quit()
                except Exception as e:
                    log.error("[Update] Failed to launch installer: %s", e)

        dlg.install_requested.connect(_on_install)
        dlg.exec()

    def _show_automations(self):
        """Show automations help."""
        log.info("Automations help requested")

    def _show_local_envs(self):
        """Show local environments help."""
        log.info("Local environments help requested")

    def _show_worktrees(self):
        """Show worktrees help."""
        log.info("Worktrees help requested")

    def _show_skills_help(self):
        """Show skills help."""
        log.info("Skills help requested")

    def _show_mcp_help(self):
        """Show Model Context Protocol help."""
        log.info("MCP help requested")

    def _show_troubleshooting(self):
        """Show troubleshooting guide."""
        log.info("Troubleshooting guide requested")

    def _send_feedback(self):
        """Send feedback."""
        import webbrowser
        webbrowser.open("https://github.com/cortex-ai/feedback")
        log.info("Opening feedback page")

    def _start_trace(self):
        """Start trace recording for debugging."""
        log.info("Trace recording started")

    def _show_keyboard_shortcuts(self):
        """Show keyboard shortcuts reference dialog (F1)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QTextEdit
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Keyboard Shortcuts Reference")
        dialog.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(dialog)
        
        # Create a text edit for displaying shortcuts
        shortcuts_text = QTextEdit()
        shortcuts_text.setReadOnly(True)
        shortcuts_text.setFontFamily("Consolas")
        shortcuts_text.setFontPointSize(10)
        
        # Build shortcuts HTML - Simple dark theme
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { 
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                    background-color: #1e1e1e;
                    color: #ffffff;
                    padding: 30px;
                    margin: 0;
                }
                h1 {
                    color: #4CAF50;
                    text-align: center;
                    font-size: 32px;
                    margin-bottom: 10px;
                }
                .subtitle {
                    text-align: center;
                    color: #9cdcfe;
                    margin-bottom: 30px;
                    font-size: 16px;
                }
                .section {
                    background-color: #252526;
                    border-left: 4px solid #4CAF50;
                    padding: 20px;
                    margin-bottom: 25px;
                    border-radius: 5px;
                }
                .section h2 {
                    color: #4CAF50;
                    margin: 0 0 15px 0;
                    font-size: 20px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    background-color: #2d2d30;
                }
                th {
                    background-color: #3e3e42;
                    color: #ffffff;
                    padding: 12px;
                    text-align: left;
                    font-weight: 600;
                    border: 1px solid #555;
                }
                td {
                    padding: 10px 12px;
                    border: 1px solid #555;
                    color: #cccccc;
                }
                tr:nth-child(even) {
                    background-color: #333337;
                }
                tr:hover {
                    background-color: #3e3e42;
                }
                .shortcut {
                    font-family: 'Consolas', monospace;
                    background-color: #1e1e1e;
                    padding: 5px 10px;
                    border-radius: 4px;
                    font-weight: bold;
                    color: #9cdcfe;
                    display: inline-block;
                    min-width: 110px;
                    text-align: center;
                }
                .status {
                    color: #4CAF50;
                    font-weight: bold;
                }
                .tip {
                    background-color: #264f78;
                    padding: 15px;
                    border-radius: 5px;
                    margin-top: 25px;
                    border-left: 4px solid #007acc;
                }
                .tip strong {
                    color: #9cdcfe;
                }
                .tip p {
                    margin: 8px 0 0 0;
                    color: #cccccc;
                }
            </style>
        </head>
        <body>
            <h1>⌨️ Keyboard Shortcuts Reference</h1>
            <p class="subtitle">Quick reference for all Cortex IDE shortcuts</p>
            
            <!-- File Operations -->
            <div class="section">
                <h2>📁 File Operations</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Menu</th></tr>
                    <tr><td>New Window</td><td><span class="shortcut">Ctrl+Shift+N</span></td><td>File</td></tr>
                    <tr><td>New Chat</td><td><span class="shortcut">Ctrl+N</span></td><td>File</td></tr>
                    <tr><td>Quick Chat</td><td><span class="shortcut">Alt+Ctrl+N</span></td><td>File</td></tr>
                    <tr><td>Open File…</td><td><span class="shortcut">Ctrl+Shift+O</span></td><td>File</td></tr>
                    <tr><td>Open Folder…</td><td><span class="shortcut">Ctrl+O</span></td><td>File</td></tr>
                    <tr><td>Save</td><td><span class="shortcut">Ctrl+S</span></td><td>File</td></tr>
                    <tr><td>Save All</td><td><span class="shortcut">Ctrl+Shift+S</span></td><td>File</td></tr>
                    <tr><td>Settings…</td><td><span class="shortcut">Ctrl+,</span></td><td>File</td></tr>
                    <tr><td>Exit</td><td><span class="shortcut">Alt+F4</span></td><td>File</td></tr>
                    <tr><td>Close Tab</td><td><span class="shortcut">Ctrl+W</span></td><td>Global</td></tr>
                    <tr><td>Close All Tabs</td><td><span class="shortcut">Ctrl+Shift+W</span></td><td>Global</td></tr>
                    <tr><td>Rename File</td><td><span class="shortcut">F2</span></td><td>Edit</td></tr>
                </table>
            </div>
            
            <!-- Edit Operations -->
            <div class="section">
                <h2>✏️ Edit Operations</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Menu</th></tr>
                    <tr><td>Undo</td><td><span class="shortcut">Ctrl+Z</span></td><td>Edit</td></tr>
                    <tr><td>Redo</td><td><span class="shortcut">Ctrl+Y</span></td><td>Edit</td></tr>
                    <tr><td>Cut</td><td><span class="shortcut">Ctrl+X</span></td><td>Edit</td></tr>
                    <tr><td>Copy</td><td><span class="shortcut">Ctrl+C</span></td><td>Edit</td></tr>
                    <tr><td>Paste</td><td><span class="shortcut">Ctrl+V</span></td><td>Edit</td></tr>
                    <tr><td>Select All</td><td><span class="shortcut">Ctrl+A</span></td><td>Edit</td></tr>
                    <tr><td>Find…</td><td><span class="shortcut">Ctrl+F</span></td><td>Monaco</td></tr>
                    <tr><td>Find & Replace…</td><td><span class="shortcut">Ctrl+H</span></td><td>Monaco</td></tr>
                    <tr><td>Find in Files…</td><td><span class="shortcut">Ctrl+Shift+F</span></td><td>Edit</td></tr>
                    <tr><td>Go to Line…</td><td><span class="shortcut">Ctrl+G</span></td><td>Monaco</td></tr>
                    <tr><td>Toggle Comment</td><td><span class="shortcut">Ctrl+/</span></td><td>Edit</td></tr>
                    <tr><td>Delete Line</td><td><span class="shortcut">Ctrl+Shift+K</span></td><td>Edit</td></tr>
                    <tr><td>Duplicate Line</td><td><span class="shortcut">Ctrl+Shift+D</span></td><td>Edit</td></tr>
                    <tr><td>Indent</td><td><span class="shortcut">Ctrl+]</span></td><td>Edit</td></tr>
                    <tr><td>Outdent</td><td><span class="shortcut">Ctrl+[</span></td><td>Edit</td></tr>
                    <tr><td>Move Line Up</td><td><span class="shortcut">Alt+Up</span></td><td>Edit</td></tr>
                    <tr><td>Move Line Down</td><td><span class="shortcut">Alt+Down</span></td><td>Edit</td></tr>
                </table>
            </div>
            
            <!-- View & Navigation -->
            <div class="section">
                <h2>👁️ View & Navigation</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Menu</th></tr>
                    <tr><td>Toggle Sidebar</td><td><span class="shortcut">Ctrl+B</span></td><td>View</td></tr>
                    <tr><td>Toggle Review Panel</td><td><span class="shortcut">Alt+Ctrl+B</span></td><td>View</td></tr>
                    <tr><td>Toggle Full Screen</td><td><span class="shortcut">F11</span></td><td>View</td></tr>
                    <tr><td>Next Tab</td><td><span class="shortcut">Ctrl+Tab</span></td><td>Global</td></tr>
                    <tr><td>Previous Tab</td><td><span class="shortcut">Ctrl+Shift+Tab</span></td><td>Global</td></tr>
                    <tr><td>Toggle Terminal</td><td><span class="shortcut">Ctrl+J</span></td><td>Terminal</td></tr>
                    <tr><td>New Terminal</td><td><span class="shortcut">Ctrl+Shift+`</span></td><td>Terminal</td></tr>
                    <tr><td>Command Palette…</td><td><span class="shortcut">Ctrl+K</span></td><td>File</td></tr>
                    <tr><td>Minimize</td><td><span class="shortcut">Ctrl+M</span></td><td>Window</td></tr>
                    <tr><td>Close Window</td><td><span class="shortcut">Ctrl+F4</span></td><td>Window</td></tr>
                </table>
            </div>
            
            <!-- AI & Tools -->
            <div class="section">
                <h2>🤖 AI & Tools</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th><th>Menu</th></tr>
                    <tr><td>Explain Code</td><td><span class="shortcut">Ctrl+Shift+E</span></td><td>AI</td></tr>
                    <tr><td>Refactor Code</td><td><span class="shortcut">Ctrl+Shift+R</span></td><td>AI</td></tr>
                    <tr><td>Write Tests</td><td><span class="shortcut">Ctrl+Shift+U</span></td><td>AI</td></tr>
                    <tr><td>Debug Help</td><td><span class="shortcut">Ctrl+Shift+H</span></td><td>AI</td></tr>
                    <tr><td>Memory Manager…</td><td><span class="shortcut">Ctrl+Shift+M</span></td><td>AI</td></tr>
                    <tr><td>AI Chat Focus</td><td><span class="shortcut">Ctrl+Shift+A</span></td><td>AI</td></tr>
                    <tr><td>Format Code</td><td><span class="shortcut">Shift+Alt+F</span></td><td>Editor</td></tr>
                    <tr><td>Collapse Region</td><td><span class="shortcut">Ctrl+Shift+[</span></td><td>Editor</td></tr>
                    <tr><td>Expand Region</td><td><span class="shortcut">Ctrl+Shift+]</span></td><td>Editor</td></tr>
                    <tr><td>Collapse All</td><td><span class="shortcut">Ctrl+K Ctrl+0</span></td><td>Editor</td></tr>
                    <tr><td>Expand All</td><td><span class="shortcut">Ctrl+K Ctrl+J</span></td><td>Editor</td></tr>
                    <tr><td>IntelliSense</td><td><span class="shortcut">Ctrl+Space</span></td><td>Editor</td></tr>
                    <tr><td>Debug Console</td><td><span class="shortcut">Ctrl+Alt+D</span></td><td>Chat</td></tr>
                </table>
            </div>
            
            <!-- Monaco Editor -->
            <div class="section">
                <h2>📝 Monaco Editor</h2>
                <table>
                    <tr><th>Action</th><th>Shortcut</th></tr>
                    <tr><td>Close Editor Tab</td><td><span class="shortcut">Ctrl+W</span></td></tr>
                    <tr><td>Send Chat Message</td><td><span class="shortcut">Enter</span></td></tr>
                    <tr><td>New Line (Chat)</td><td><span class="shortcut">Ctrl+Enter</span> / <span class="shortcut">Shift+Enter</span></td></tr>
                    <tr><td>Completion (Chat)</td><td><span class="shortcut">Ctrl+Space</span></td></tr>
                </table>
            </div>
            
            <!-- Tip -->
            <div class="tip">
                <strong>💡 Tip:</strong>
                <p>Press <span class="shortcut" style="min-width: 50px;">F1</span> anytime to open this reference!</p>
            </div>
        </body>
        </html>
        """
        
        shortcuts_text.setHtml(html_content)
        layout.addWidget(shortcuts_text)
        
        dialog.exec()

    def _emergency_shutdown_save(self) -> None:
        """Save agent state + trigger chat persistence immediately.
        
        Called from closeEvent, nativeEvent (WM_ENDSESSION), and signal handlers.
        Idempotent — safe to call multiple times.
        
        FIXED v2: Three-layer save strategy to eliminate chat loss on IDE restart:
        1. JS bridge save (fast path — works if WebChannel is alive)
        2. Python direct save (reliable — works even if WebChannel is torn down)
        3. DB flush (guarantees SQLite write queue is emptied)
        """
        if getattr(self, '_shutdown_save_done', False):
            return
        self._shutdown_save_done = True
        
        # 1. Save agent session state
        try:
            from src.core.agent_session_manager import save_snapshot
            from src.ai.agent_bridge import get_agent_bridge
            bridge = get_agent_bridge()
            if bridge is not None:
                save_snapshot(bridge)
                log.info("[SESSION] Agent state saved on shutdown")
                # Auto-apply any pending deferred edits so AI work is never lost
                try:
                    applied = bridge.apply_all_deferred_edits()
                    if applied > 0:
                        log.info(f"[SHUTDOWN] Auto-applied {applied} deferred edit(s) — AI changes saved to disk")
                except Exception as _def_exc:
                    log.warning(f"[SHUTDOWN] Failed to auto-apply deferred edits: {_def_exc}")
                # Save MEMORY.md on shutdown so context survives app restart
                # WARNING: Call save_memory_sync() NOT save_session_to_memory().
                # The latter spawns a DAEMON thread that calls the LLM (5-15s).
                # Daemon threads are killed when the process exits, so the
                # summary would NEVER be written. save_memory_sync() is fully
                # synchronous and uses _build_brief_done_summary() — no LLM.
                try:
                    bridge.save_memory_sync()
                    log.info("[SHUTDOWN] MEMORY.md saved synchronously on shutdown")
                except Exception as _mem_exc:
                    log.warning(f"[SHUTDOWN] Failed to save MEMORY.md: {_mem_exc}")
        except Exception as _ses_exc:
            log.warning(f"[SESSION] Failed to save agent state on shutdown: {_ses_exc}")
        
        # 2. Save chat history - SYNCHRONOUS wait for completion
        try:
            if hasattr(self, '_ai_chat') and self._ai_chat:
                if getattr(self, '_is_native_chat', False):
                    # Native mode: chat is persisted via ChatStore, no JS needed
                    log.info("[SHUTDOWN] Native chat — no JS save needed")
                else:
                    # Step A: Set shutdown flag BEFORE any JS executes
                    self._ai_chat.run_javascript("window._shutdownInProgress = true;")

                    # Step A2: Flush partial response + timeline
                    self._ai_chat.run_javascript(
                        "if(window._flushTimelineEntriesToChat) _flushTimelineEntriesToChat();"
                        "if(window._savePartialResponse) _savePartialResponse();"
                    )

                    # Step B: Trigger the save and wait for confirmation
                    save_confirmed = [False]

                    def _on_save_done(status):
                        save_confirmed[0] = True
                        log.info(f"[SHUTDOWN] Chat save confirmed: {status}")

                    if hasattr(self._ai_chat, 'save_finished'):
                        self._ai_chat.save_finished.connect(_on_save_done)

                    self._ai_chat.run_javascript(
                        "if(window.flushScheduledSaveChats) { "
                        "  flushScheduledSaveChats(); "
                        "  if(window.bridge && bridge.on_save_finished) bridge.on_save_finished('OK'); "
                        "} else if(window.saveProjectChats) { "
                        "  if(window.chats) saveProjectChats(window.chats); "
                        "  if(window.bridge && bridge.on_save_finished) bridge.on_save_finished('OK'); "
                        "}"
                    )

                    # Step C: Wait up to 3000ms for JS save
                    from PyQt6.QtCore import QDeadlineTimer, QEventLoop as _QEventLoop
                    from PyQt6.QtWidgets import QApplication
                    deadline = QDeadlineTimer(3000)
                    while not save_confirmed[0] and not deadline.hasExpired():
                        QApplication.processEvents(_QEventLoop.ProcessEventsFlag.AllEvents, 50)

                    if hasattr(self._ai_chat, 'save_finished'):
                        try:
                            self._ai_chat.save_finished.disconnect(_on_save_done)
                        except Exception:
                            pass

                    if not save_confirmed[0]:
                        log.warning("[SHUTDOWN] JS bridge chat save timed out — Python fallback")
                        try:
                            from src.core.chat_history import get_chat_history
                            from src.core.database import get_database
                            history = get_chat_history()
                            db = get_database()
                            project_path = getattr(self, '_current_project_path', None) or "shutdown_fallback"
                            db.flush_write_queue(force=True)
                            conversations = history.get_conversations(project_path)
                            if conversations:
                                log.info(f"[SHUTDOWN] Python fallback: found {len(conversations)} conversations")
                        except Exception as py_exc:
                            log.warning(f"[SHUTDOWN] Python fallback save error: {py_exc}")
        except Exception as save_exc:
            log.warning(f"[SHUTDOWN] Chat save error: {save_exc}")
        
        # 3. Force-flush DB write queue (guaranteed synchronous)
        try:
            from src.core.database import get_database
            db = get_database()
            db.flush_write_queue(force=True)
            log.info("[SHUTDOWN] DB write queue flushed")
        except Exception as db_exc:
            log.warning(f"[SHUTDOWN] DB flush error: {db_exc}")



    def closeEvent(self, event: QCloseEvent):
        """Save session on close, prompt for unsaved files, and kill terminals."""
        # ════════════════════════════════════════════════════════════════
        # INSTANT HIDE — Prevent the "frozen/dark window" glitch during
        # close. The heavy shutdown work (file saves, agent stop, DB flush)
        # can take 2-5 seconds. Without hiding first, the user sees the
        # window freeze and go dark during that time.
        # ════════════════════════════════════════════════════════════════
        self.hide()
        QApplication.processEvents()

        # 0. Save agent session state for resume after restart
        self._emergency_shutdown_save()

        # 0b. Mark clean shutdown so crash recovery doesn't fire on next startup
        try:
            conv_id = getattr(self._ai_chat, '_conversation_id', None)
            if conv_id:
                from src.core.crash_persistence import get_crash_store
                get_crash_store().mark_clean_shutdown(conv_id)
        except Exception:
            pass

        # 0c. STOP AGENT WORKER — CRITICAL: Must happen BEFORE the Qt event
        # loop exits. The asyncio thread runs its own event loop (IOCP on
        # Windows). If we don't stop it, the thread keeps running after
        # closeEvent returns and crashes with access violation when the
        # process teardown destroys socket handles.
        try:
            from src.ai.agent_bridge import get_agent_bridge
            bridge = get_agent_bridge()
            if bridge is not None:
                worker = getattr(bridge, '_worker', None)
                if worker is not None and worker.isRunning():
                    worker.stop()  # Sets _is_running=False, breaks the loop
                    # Force-close the asyncio event loop to interrupt any
                    # blocked I/O (network reads, socket polls). Without this,
                    # the worker thread can hang on PySocket_Read for 30+ seconds.
                    _w_loop = getattr(worker, '_loop', None)
                    if _w_loop is not None and not _w_loop.is_closed():
                        try:
                            _w_loop.call_soon_threadsafe(_w_loop.stop)
                        except Exception:
                            pass
                        try:
                            import time as _wt
                            _wt.sleep(0.1)  # Let loop settle
                            if not _w_loop.is_closed():
                                _w_loop.close()
                        except Exception:
                            pass
                    worker.wait(3000)  # Wait up to 3s for clean shutdown
                    if worker.isRunning():
                        log.warning("[SHUTDOWN] Agent worker did not stop in 3s — terminating")
                        worker.terminate()
                        worker.wait(1000)
                    log.info("[SHUTDOWN] Agent worker stopped")
                # Also stop the asyncio event loop if still running
                loop = getattr(bridge, '_loop', None) or (getattr(worker, '_loop', None) if worker else None)
                if loop is not None and loop.is_running():
                    loop.call_soon_threadsafe(loop.stop)
        except Exception as _w_exc:
            log.warning(f"[SHUTDOWN] Agent worker stop error: {_w_exc}")

        # 1. Check for unsaved files — use _modified_files (tracked by
        #    _on_webview_content_changed) which accurately reflects Monaco
        #    content changes, plus legacy editor_tabs as fallback.
        modified_files = getattr(self, '_modified_files', set()).copy()
        # Fallback: if _modified_files is empty (e.g., tracker not yet initialized),
        # check _open_files as a safety net for any unsaved changes.
        if not modified_files and hasattr(self, '_webview_panel'):
            modified_files = set(
                fp for fp in self._webview_panel._open_files
                if fp and fp != "untitled.py"
            )
        # Also include legacy editor tab modified files
        legacy_modified = getattr(self._editor_tabs, '_modified', set())
        if legacy_modified:
            modified_files.update(legacy_modified)

        if modified_files:
            from PyQt6.QtWidgets import QMessageBox
            file_names = [os.path.basename(f) for f in modified_files]
            files_str = ", ".join(file_names[:3]) + ("..." if len(file_names) > 3 else "")

            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                f"You have unsaved changes in: {files_str}\n\nDo you want to save them before closing?",
                QMessageBox.StandardButton.SaveAll |
                QMessageBox.StandardButton.Discard |
                QMessageBox.StandardButton.Cancel
            )

            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.SaveAll:
                # ── Save all modified files using Monaco's fresh content ──
                # Tell Monaco to save every modified file; Monaco will call
                # bridge.saveFile() for each, which triggers _on_webview_save_requested
                # that writes fresh content to disk (no stale cache).
                from PyQt6.QtCore import QDeadlineTimer, QEventLoop as _QEventLoop

                saved_count = [0]
                failed = []

                if hasattr(self, '_webview_panel') and self._webview_panel._page_loaded:
                    for fp in list(modified_files):
                        try:
                            safe_path = json.dumps(fp)
                            # saveFile bridge: Monaco sends fresh model content → Python writes to disk
                            self._webview_panel._safe_run_js(
                                f"(function(){{"
                                f"var models=monaco.editor.getModels();"
                                f"for(var i=0;i<models.length;i++){{"
                                f"var u=models[i].uri.toString();"
                                f"var p=u.replace('file:///','');"
                                f"if(p.indexOf('/')!==-1)p=p.replace(/\\//g,'\\\\');"
                                f"if(p===decodeURI({safe_path}).replace(/\\//g,'\\\\')){{"
                                f"bridge.saveFile({safe_path},models[i].getValue());"
                                f"break;"
                                f"}}"
                                f"}}"
                                f"}})();"
                            )
                            saved_count[0] += 1
                        except Exception as e:
                            failed.append((fp, str(e)))
                            log.error(f"[SHUTDOWN] Failed to trigger save for {fp}: {e}")

                    # Process events to let JS bridge.saveFile() callbacks fire
                    deadline = QDeadlineTimer(2000)
                    while not deadline.hasExpired():
                        QApplication.processEvents(_QEventLoop.ProcessEventsFlag.AllEvents, 50)

                # Also save any legacy editor tab files
                for fp in legacy_modified:
                    idx = -1
                    for i, tfp in self._editor_tabs._files.items():
                        if tfp == fp:
                            idx = i
                            break
                    if idx >= 0:
                        editor = self._editor_tabs.widget(idx)
                        if isinstance(editor, CodeEditor):
                            try:
                                with open(fp, 'w', encoding='utf-8') as f:
                                    f.write(editor.toPlainText())
                                self._editor_tabs._mark_saved(fp)
                            except Exception as e:
                                log.error(f"[SHUTDOWN] Failed to auto-save {fp}: {e}")

                if failed:
                    log.warning(f"[SHUTDOWN] {len(failed)} file(s) failed to save: {failed}")
                log.info(f"[SHUTDOWN] Saved {saved_count[0]} modified file(s) on close")
        
        # 2. Save IDE UI state
        # Only save from the active webview editor — NOT the hidden legacy _editor_tabs
        # which accumulates stale entries from old sessions and never gets cleaned.
        fps = []
        if hasattr(self, '_webview_panel'):
            for fp in self._webview_panel._open_files:
                if fp and fp != "untitled.py":
                    fps.append(fp)
            log.info(f"[SESSION] Saving {len(fps)} open files from webview: {[Path(f).name for f in fps]}")
        active = self._webview_panel.get_active_file() or self._editor_tabs.current_filepath()
        expanded = self._sidebar.get_expanded_paths()
        self._session_manager.save(fps, active, {"expanded_paths": expanded})
        self._settings.set("window", "maximized", self.isMaximized())
        if not self.isMaximized():
            self._settings.set("window", "width", self.width())
            self._settings.set("window", "height", self.height())
        
        # Save panel widths (Codex 4-panel layout)
        # Left sidebar: 220px, Review: 380px, File tree: 280px are fixed
        # Only chat panel width varies and is not saved (flexible)
            
        # 4. Clean up terminals
        for i in range(self._terminal_tabs.count()):
            term = self._terminal_tabs.widget(i)
            if isinstance(term, XTermWidget):
                term._kill_process()

        # 5. Stop Live Server if running
        # Stop Live Server (removed in AI-first mode)
        if False and self._live_server and self._live_server.is_running:
            self._live_server.stop()

    
        # 6. FINAL DB FLUSH — Guarantee all chat history is persisted to SQLite.
        # This is the last line of defense: after JS bridge save and Python fallback
        # in _emergency_shutdown_save, force the SQLite write queue to empty.
        # Without this, chats may appear as empty metadata on next IDE restart.
        try:
            from src.core.database import get_database as _get_db
            _db = _get_db()
            _db.flush_write_queue(force=True)
            log.info("[SHUTDOWN] Final DB write queue flush complete")
        except Exception as _db_exc:
            log.warning(f"[SHUTDOWN] Final DB flush error: {_db_exc}")

        # 7. Force-hide any visible spinner overlays to ensure clean shutdown.
        try:
            if hasattr(self, '_ai_chat') and self._ai_chat:
                if hasattr(self._ai_chat, '_spinner_overlay'):
                    self._ai_chat._spinner_overlay.force_hide()
                    log.info("[SHUTDOWN] Spinner overlay force-hidden")
        except Exception:
            pass

        event.accept()

    def dragEnterEvent(self, event):
        """Accept drag of folders or files from Explorer onto the window."""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        """Handle drop of folders/files — open as project or file."""
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if os.path.isdir(path):
                self._open_folder_programmatic(path)
                event.acceptProposedAction()
                return
            elif os.path.isfile(path):
                self._open_file(path)
                event.acceptProposedAction()
                return

    # ── Custom title bar: maximize toggle ─────────────────────────────

    def _toggle_maximize(self):
        """Toggle between maximized and normal window state.

        Called by the custom title bar's green maximize button and
        double-click on the title bar area.
        """
        if self.isMaximized():
            self.showNormal()
        else:
            self.showMaximized()

    def resizeEvent(self, event):
        """Handle window resize — responsive layout + startup overlay."""
        super().resizeEvent(event)
        # Keep startup overlay covering full window during resize
        if hasattr(self, '_startup_overlay') and self._startup_overlay:
            self._startup_overlay.setGeometry(self.rect())

    # Phase 1, 2, 3 Integration Methods
    def _set_agent_mode(self, mode: str):
        """
        Set AI agent mode (Phase 1 Integration).
        
        Args:
            mode: One of 'build', 'explore', 'debug', 'plan'
        """
        self._ai_agent.set_mode(mode)
        mode_names = {
            'build': '🏗️ Build',
            'explore': '🔍 Explore', 
            'debug': '🐛 Debug',
            'plan': '📋 Plan'
        }
        self._statusbar.showMessage(f"Agent mode: {mode_names.get(mode, mode)}", 3000)
        log.info(f"Agent mode switched to: {mode}")
    
    def _show_skills_browser(self):
        """Show skills browser dialog (Phase 3 Integration)."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QListWidget, QLabel, QPushButton, QTextEdit
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Skills Browser")
        dialog.setMinimumSize(600, 400)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🛠️ Available Skills</h2>")
        layout.addWidget(title)
        
        # Skills list
        skills_list = QListWidget()
        skills = self._ai_agent.get_available_skills()
        
        for skill in skills:
            item_text = f"{skill['name']} ({skill['id']})"
            skills_list.addItem(item_text)
        
        layout.addWidget(skills_list)
        
        # Description
        desc_label = QLabel("Select a skill to view capabilities")
        layout.addWidget(desc_label)
        
        # Capability display
        capability_text = QTextEdit()
        capability_text.setReadOnly(True)
        capability_text.setPlaceholderText("Skill capabilities will appear here...")
        layout.addWidget(capability_text)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("Skills browser opened")

    # Phase 4 Integration Methods
    def _show_todo_manager(self):
        """Show TODO manager dialog.
        
        NOTE: The legacy _todo_manager has been removed. This dialog now shows
        the current todos from the bridge/JS UI instead.
        """
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QListWidget
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Tasks & TODOs")
        dialog.setMinimumSize(500, 350)
        
        layout = QVBoxLayout(dialog)
        
        title = QLabel("<h2>Task Manager</h2>")
        layout.addWidget(title)
        
        info = QLabel("Todos are now managed by the AI agent via TodoWrite.\nUse the chat sidebar to view and track task progress.")
        info.setWordWrap(True)
        layout.addWidget(info)
        
        layout.addWidget(QLabel("Current session todos are shown in the sidebar."))
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("TODO manager dialog opened")
    
    def _add_todo_task(self):
        """Add a new TODO task. (Legacy - todo manager removed)"""
        self._statusbar.showMessage("Todo manager has been removed. Ask the AI to create tasks via TodoWrite.", 5000)
        log.info("[TODO] _add_todo_task called - todo manager removed")
    
    def _complete_todo_task(self):
        """Complete a TODO task. (Legacy - todo manager removed)"""
        self._statusbar.showMessage("Todo manager has been removed. Toggle tasks in the sidebar instead.", 5000)
        log.info("[TODO] _complete_todo_task called - todo manager removed")
    
    def _show_permission_settings(self):
        """Show permission settings dialog with notification toggles."""
        from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel,
                                     QPushButton, QCheckBox, QGroupBox)
        from src.config.settings import get_settings

        settings = get_settings()
        dialog = QDialog(self)
        dialog.setWindowTitle("Permission and Notification Settings")
        dialog.setMinimumSize(520, 420)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>🔒 Permission System</h2>")
        layout.addWidget(title)
        
        # Info
        info = QLabel("Permission system is active and monitoring tool usage.")
        layout.addWidget(info)
        
        # Cache info
        cache_size = len(self._permission_evaluator._permission_cache)
        cache_label = QLabel(f"Cached decisions: {cache_size}")
        layout.addWidget(cache_label)
        
        # Clear cache button
        def clear_cache():
            self._permission_evaluator.clear_cache()
            cache_label.setText("Cached decisions: 0")
        
        clear_btn = QPushButton("Clear Permission Cache")
        clear_btn.clicked.connect(clear_cache)
        layout.addWidget(clear_btn)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.close)
        layout.addWidget(close_btn)
        
        dialog.exec()
        log.info("Permission settings dialog opened")
    
    def _show_memory_manager(self):
        """Open the Memory Manager dialog (AI → Memory Manager...).
        
        Deferred via QTimer.singleShot(0, ...) to break out of the QWebChannel
        JS callback stack when triggered from the sidebar settings gear.
        Without this, the nested event loop from dlg.exec() prevents the
        dialog's QWebEngineView from completing its async page load.
        """
        import logging
        from PyQt6.QtCore import QTimer
        logging.getLogger("main_window").info("[MainWindow] _show_memory_manager called")
        QTimer.singleShot(0, self._do_show_memory_manager)

    def _do_show_memory_manager(self):
        """Actual dialog creation — deferred to next event loop iteration."""
        import logging
        from src.ui.dialogs.memory_manager import MemoryManagerDialog
        try:
            from src.config.settings import get_settings
            settings = get_settings()
        except Exception:
            settings = None
        project_root = getattr(self, '_current_project_path', None) or os.getcwd()
        try:
            # CACHE the dialog across opens. Building it fresh every time
            # spawns a whole new Chromium renderer + reloads ~3k lines of
            # HTML/CSS/JS — measured as a multi-second white-page delay on
            # every open in compiled builds. First open still pays the
            # boot cost; every reopen is instant. Cache is invalidated on
            # project switch (memory dirs are per-project).
            dlg = getattr(self, '_memory_manager_dlg', None)
            cached_root = getattr(self, '_memory_manager_dlg_root', None)
            if dlg is not None and cached_root != project_root:
                try:
                    dlg.deleteLater()
                except RuntimeError:
                    pass
                dlg = None
            if dlg is None:
                dlg = MemoryManagerDialog(project_root, settings=settings, parent=self)
                # Sync settings model change → chat panel model button
                if hasattr(dlg, '_bridge') and hasattr(self, '_ai_chat'):
                    dlg._bridge.model_changed.connect(self._on_settings_model_changed)
                self._memory_manager_dlg = dlg
                self._memory_manager_dlg_root = project_root
            else:
                # Reopening the live page — push fresh memory state into it
                try:
                    dlg._bridge._emit_refresh()
                except Exception:
                    pass
            dlg.exec()
        except Exception as exc:
            logging.getLogger("main_window").error(f"[MainWindow] Memory manager failed: {exc}", exc_info=True)

    # Phase 4: Real-time UI Update Handlers
    # NOTE: The following handlers are LEGACY stubs. The _todo_manager has been removed.
    # Todo state is now managed entirely by the bridge (CortexAgentBridge) via TodoWrite.
    
    def _on_todo_task_added(self, task_id: str):
        """Handle new todo task - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_added called (legacy): {task_id}")

    def _on_toggle_todo(self, task_id: str, completed: bool):
        """Handle todo toggle from UI.
        
        NOTE: The legacy _todo_manager has been removed. Todo state is now managed
        entirely by the bridge (CortexAgentBridge) and the JS UI. This handler
        logs the toggle but does not persist it - the UI already shows the toggled
        state visually and the bridge will emit updated todos on its next turn.
        """
        status = "completed" if completed else "reopened"
        log.info(f"[TODO] UI toggle: {task_id} -> {status}")
        try:
            self._ai_agent.toggle_todo_status(task_id, completed)
        except Exception as exc:
            log.warning(f"[TODO] Failed to sync toggle to backend state: {exc}")

    def _on_todo_task_completed(self, task_id: str):
        """Handle completed todo - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_completed called (legacy): {task_id}")

    def _on_todo_task_updated(self, task_id: str):
        """Handle updated todo - LEGACY (todo manager removed)."""
        log.info(f"[TODO] _on_todo_task_updated called (legacy): {task_id}")

    def _on_task_progress_update(self, completed: int, total: int, pct: int, message: str):
        """Update UI with real-time AI task progress."""
        try:
            self._last_task_progress_msg = message
            self._last_task_progress_counts = (completed, total, pct)
            if hasattr(self, '_statusbar') and self._statusbar:
                self._statusbar.showMessage(f'AI Progress: {message}', 5000)
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')

    def _on_ai_task_complete(self, response: str):
        """Response stored on bridge._pending_notification_response BEFORE emit.
        The _on_chat_rendering_done callback reads it after rendering completes.
        Fallback: 500ms timer if callback isn't registered."""
        if not hasattr(self._ai_chat, '_rendering_done_cb') or not self._ai_chat._rendering_done_cb:
            from PyQt6.QtCore import QTimer
            QTimer.singleShot(500, self._on_chat_rendering_done)

    def _on_chat_rendering_done(self):
        """Show notification AFTER the ChatPanel has finished rendering."""
        # Read from bridge (set BEFORE response_complete.emit) to avoid
        # race between _on_response_complete and _on_ai_task_complete order.
        response = getattr(self._ai_agent, '_pending_notification_response', None)
        if response is None:
            return
        self._ai_agent._pending_notification_response = None
        self._do_ai_task_complete(response)

    def _do_ai_task_complete(self, response: str):
        """Actual notification logic — runs after event queue drains."""
        try:
            # Re-enable sidebar file-tree refreshes now that AI is done
            if hasattr(self, '_sidebar') and self._sidebar:
                self._sidebar._ai_active = False
                # Only refresh sidebar if agent actually made file changes
                try:
                    _bridge = self._ai_agent
                    _mutations_now = getattr(_bridge, '_session_mutation_count', 0)
                    _mutations_before = getattr(_bridge, '_last_cycle_session_mutations', 0)
                    _had_file_changes = _mutations_now > _mutations_before
                except Exception:
                    _had_file_changes = False
                if _had_file_changes:
                    from PyQt6.QtCore import QTimer
                    QTimer.singleShot(2000, lambda: self._sidebar._bridge.refreshFileTree() if hasattr(self, '_sidebar') and self._sidebar and self._sidebar._bridge else None)

            # ── Check if task complete notifications are enabled ──
            from src.utils.notifications import _get_notif_setting
            if not _get_notif_setting("task_complete_enabled", True):
                log.info('[MainWindow] Task complete notification DISABLED by user settings')
                return

            # Build an up-to-date progress summary from todos when available
            progress_msg = getattr(self, '_last_task_progress_msg', '')
            try:
                todos = getattr(self._ai_agent, '_current_todos', None) or []
                if isinstance(todos, list) and todos:
                    total = len(todos)
                    completed = sum(
                        1 for t in todos
                        if str(getattr(t, 'get', lambda *_: '')('status', '')).upper() in ('COMPLETE', 'CANCELLED')
                    )
                    pending = max(0, total - completed)
                    pct = int((completed / total) * 100) if total else 0
                    progress_msg = f"{completed}/{total} tasks complete ({pct}%)"
                    if pending:
                        progress_msg += f", {pending} pending"
            except Exception as e:
                log.debug(f'[MainWindow] Suppressed error: {e}')

            # If the bridge auto-cancelled remaining todos, never claim completion
            if getattr(self._ai_agent, '_todos_auto_cancelled', False):
                show_toast_notification(
                    'Cortex AI IDE — Needs Attention',
                    'Tasks were auto-cancelled after repeated attempts without progress.'
                )
                log.info('[MainWindow] Notification shown: todos auto-cancelled')
                return

            # INDUSTRY-STANDARD: Check if notification should actually show
            _should_notify = True
            if hasattr(self._ai_agent, '_allow_notification'):
                _should_notify = self._ai_agent._allow_notification

            # Play alert sound if enabled (before checking focus)
            try:
                _sound_on = self._settings.get("notifications", "sound_alerts", default=False) if hasattr(self, '_settings') and self._settings else False
                log.info(f"[SOUND] Setting={_sound_on}")
                if _sound_on:
                    import winsound, threading as _st
                    def _beep():
                        try:
                            winsound.Beep(523, 120)
                            winsound.Beep(659, 150)
                        except Exception:
                            pass
                    _st.Thread(target=_beep, daemon=True, name="AlertSound").start()
                    log.info("[SOUND] Beep thread started")
            except Exception as _snd_err:
                log.warning(f"[SOUND] Error: {_snd_err}")

            if not _should_notify:
                if progress_msg:
                    show_toast_notification('Cortex AI IDE — In Progress', progress_msg)
                log.info('[MainWindow] Completion notification suppressed: tasks not genuinely complete')
                return

            # Only show completion notification if response is meaningful
            if response and len(response) > 10:
                # ── Extract CLEAN, MINIMAL summary — not raw AI text ──
                summary = _extract_notification_summary(response)
                if progress_msg:
                    msg = f"{progress_msg}  •  {summary}" if summary else progress_msg
                else:
                    msg = summary or "Task completed successfully."
                show_toast_notification('Cortex AI IDE — Task Complete', msg)
                log.info(f"[MainWindow] Windows notification shown: {msg[:80]}...")
            else:
                # Response was empty or very short — task still genuinely completed.
                # Show a simple notification so the user knows the AI is done.
                msg = progress_msg or "Task completed successfully."
                show_toast_notification('Cortex AI IDE — Task Complete', msg)
                log.info(f"[MainWindow] Windows notification shown (short/empty response): {msg}")

            # TODO: POINTS SYSTEM - Disabled for development
            # self._consume_points_for_response(response)

            # ── AUTO-APPLY REMAINING DEFERRED EDITS ON TASK COMPLETE ──
            # Defer to next event loop so response_complete.emit() returns immediately.
            _deferred_keys = self._ai_agent.get_deferred_edit_keys()
            if _deferred_keys:
                QTimer.singleShot(0, lambda: self._apply_deferred_edits_on_complete(_deferred_keys))
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')

    def _apply_deferred_edits_on_complete(self, _deferred_keys):
        """Apply deferred edits after task completes (runs on next event loop)."""
        try:
            applied = self._ai_agent.apply_all_deferred_edits()
            log.info(
                f"[MainWindow] Auto-applied {applied} deferred edit(s) "
                f"on AI task complete"
            )
            for _key in _deferred_keys:
                _is_open = (
                    hasattr(self, '_webview_panel')
                    and _key in self._webview_panel._open_files
                )
                if _is_open:
                    try:
                        self._on_accept_file_edit(_key)
                    except Exception as _rfe:
                        log.warning(
                            f"[MainWindow] Auto-accept refresh failed "
                            f"for {_key}: {_rfe}"
                        )
        except Exception as e:
            log.debug(f'[MainWindow] Suppressed error: {e}')